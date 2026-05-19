from __future__ import annotations

import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from lxml import etree


VENDOR_LATEX2MATHML = Path(__file__).resolve().parents[1] / "_vendor" / "latex2mathml"
if VENDOR_LATEX2MATHML.exists():
    sys.path.insert(0, str(VENDOR_LATEX2MATHML))

try:
    import latex2mathml.converter as latex_converter
except Exception:  # pragma: no cover - fallback keeps document generation usable.
    latex_converter = None


ROOT = Path(__file__).resolve().parents[1]
THESIS_TITLE = "基于OpenGL的CAE软件数据预处理方法研究与实践"
SOURCE_MD = ROOT / "wendang" / f"{THESIS_TITLE}_本科毕业论文.md"
TEMPLATE_DOCX = Path.home() / "Desktop" / "本科生毕业论文撰写规范.docx"
OUTPUT_DOCX = ROOT / "wendang" / f"{THESIS_TITLE}_本科毕业论文.docx"
FORMULA_OUTPUT_DOCX = ROOT / "wendang" / f"{THESIS_TITLE}_本科毕业论文_公式版.docx"
WORD_FORMULA_SCRIPT = ROOT / "scripts" / "render_docx_formulas.vbs"

CN_FONT = "宋体"
HEADING_FONT = "黑体"
EN_FONT = "Times New Roman"
MATH_FONT = "Cambria Math"
FORMULA_PLACEHOLDERS: list[tuple[str, str]] = []


def set_run_font(run, cn_font: str = CN_FONT, en_font: str = EN_FONT, size: float | None = 12):
    if size is not None:
        run.font.size = Pt(size)
    run.font.name = en_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), cn_font)


def set_style_font(document: Document, style_name: str, cn_font: str, en_font: str, size: float | None):
    if style_name not in [style.name for style in document.styles]:
        return
    style = document.styles[style_name]
    if size is not None:
        style.font.size = Pt(size)
    style.font.name = en_font
    style._element.rPr.rFonts.set(qn("w:eastAsia"), cn_font)


def configure_document(document: Document):
    set_style_font(document, "Normal", CN_FONT, EN_FONT, 12)
    set_style_font(document, "Heading 1", HEADING_FONT, EN_FONT, 16)
    set_style_font(document, "Heading 2", HEADING_FONT, EN_FONT, 14)
    set_style_font(document, "Heading 3", HEADING_FONT, EN_FONT, 12)

    for section in document.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)


def clear_body(document: Document):
    body = document._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


LATEX_REPLACEMENTS = {
    r"\nabla": "∇",
    r"\Delta": "Δ",
    r"\delta": "δ",
    r"\xi": "ξ",
    r"\eta": "η",
    r"\zeta": "ζ",
    r"\varepsilon": "ε",
    r"\epsilon": "ε",
    r"\sigma": "σ",
    r"\alpha": "α",
    r"\ell": "l",
    r"\in": "∈",
    r"\mid": "|",
    r"\sum": "∑",
    r"\frac": "frac",
    r"\dfrac": "frac",
    r"\partial": "∂",
    r"\approx": "≈",
    r"\times": "×",
    r"\cdot": "·",
    r"\le": "≤",
    r"\ge": "≥",
    r"\neq": "≠",
    r"\left": "",
    r"\right": "",
    r"\qquad": "    ",
    r"\quad": "  ",
    r"\mathrm": "",
    r"\mathbf": "",
    r"\boldsymbol": "",
    r"\mathcal": "",
    r"\begin{aligned}": "",
    r"\end{aligned}": "",
    r"\begin{bmatrix}": "[",
    r"\end{bmatrix}": "]",
}


def simplify_latex(text: str) -> str:
    text = text.replace(r"\(", "").replace(r"\)", "")
    text = text.replace(r"\[", "").replace(r"\]", "")
    text = text.replace("`", "")

    previous = None
    frac_pattern = re.compile(r"\\d?frac\{([^{}]+)\}\{([^{}]+)\}")
    while previous != text:
        previous = text
        text = frac_pattern.sub(r"(\1)/(\2)", text)

    for command, value in LATEX_REPLACEMENTS.items():
        text = text.replace(command, value)

    text = re.sub(r"\\(?:mathrm|mathbf|boldsymbol|mathcal)\{([^{}]+)\}", r"\1", text)
    text = text.replace(r"\\", "\n")
    text = text.replace("&", " ")
    text = text.replace("{", "").replace("}", "")
    text = re.sub(r"\s+\n", "\n", text)
    text = re.sub(r"\n\s+", "\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


def m_el(name: str):
    return OxmlElement(f"m:{name}")


def m_text_run(text: str):
    run = m_el("r")
    text_el = m_el("t")
    if text.startswith(" ") or text.endswith(" "):
        text_el.set(qn("xml:space"), "preserve")
    text_el.text = text
    run.append(text_el)
    return run


def append_math_nodes(parent, nodes):
    for node in nodes:
        parent.append(node)


def math_arg(name: str, nodes):
    arg = m_el(name)
    if nodes:
        append_math_nodes(arg, nodes)
    else:
        arg.append(m_text_run(" "))
    return arg


def mathml_node_text(node) -> str:
    return "".join(node.itertext()).strip()


def mathml_to_omml_nodes(node):
    tag = etree.QName(node).localname

    if tag in {"math", "mrow"}:
        nodes = []
        if node.text and node.text.strip():
            nodes.append(m_text_run(node.text))
        for child in node:
            nodes.extend(mathml_to_omml_nodes(child))
            if child.tail and child.tail.strip():
                nodes.append(m_text_run(child.tail))
        return nodes

    if tag in {"mi", "mn", "mo", "mtext"}:
        text = "".join(node.itertext())
        return [m_text_run(text)] if text else []

    if tag == "mspace":
        return [m_text_run(" ")]

    if tag == "msub":
        children = list(node)
        if len(children) >= 2:
            el = m_el("sSub")
            el.append(math_arg("e", mathml_to_omml_nodes(children[0])))
            el.append(math_arg("sub", mathml_to_omml_nodes(children[1])))
            return [el]

    if tag == "msup":
        children = list(node)
        if len(children) >= 2:
            el = m_el("sSup")
            el.append(math_arg("e", mathml_to_omml_nodes(children[0])))
            el.append(math_arg("sup", mathml_to_omml_nodes(children[1])))
            return [el]

    if tag == "msubsup":
        children = list(node)
        if len(children) >= 3:
            el = m_el("sSubSup")
            el.append(math_arg("e", mathml_to_omml_nodes(children[0])))
            el.append(math_arg("sub", mathml_to_omml_nodes(children[1])))
            el.append(math_arg("sup", mathml_to_omml_nodes(children[2])))
            return [el]

    if tag == "mfrac":
        children = list(node)
        if len(children) >= 2:
            el = m_el("f")
            el.append(math_arg("num", mathml_to_omml_nodes(children[0])))
            el.append(math_arg("den", mathml_to_omml_nodes(children[1])))
            return [el]

    if tag == "mover":
        children = list(node)
        if len(children) >= 2:
            accent_char = mathml_node_text(children[1]) or "¯"
            el = m_el("acc")
            pr = m_el("accPr")
            chr_el = m_el("chr")
            chr_el.set(qn("m:val"), accent_char)
            pr.append(chr_el)
            el.append(pr)
            el.append(math_arg("e", mathml_to_omml_nodes(children[0])))
            return [el]

    # Fallback for uncommon MathML elements in the generated thesis.
    nodes = []
    if node.text and node.text.strip():
        nodes.append(m_text_run(node.text))
    for child in node:
        nodes.extend(mathml_to_omml_nodes(child))
        if child.tail and child.tail.strip():
            nodes.append(m_text_run(child.tail))
    return nodes


def latex_to_omath(latex: str):
    if latex_converter is None:
        raise RuntimeError("latex2mathml is not available")
    mathml = latex_converter.convert(latex.strip())
    root = etree.fromstring(mathml.encode("utf-8"))
    omath = m_el("oMath")
    append_math_nodes(omath, mathml_to_omml_nodes(root))
    return omath


def linearize_latex(latex: str) -> str:
    lines = latex.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    return " ".join(line.strip() for line in lines if line.strip())


def register_formula_placeholder(latex: str) -> str:
    token = f"__EQ_PLACEHOLDER_{len(FORMULA_PLACEHOLDERS) + 1:04d}__"
    FORMULA_PLACEHOLDERS.append((token, linearize_latex(latex)))
    return token


def write_formula_map(output_path: Path) -> Path:
    map_path = output_path.with_suffix(".formula-map.txt")
    lines = []
    for token, latex in FORMULA_PLACEHOLDERS:
        safe_latex = latex.replace("\t", " ").replace("\n", " ")
        lines.append(f"{token}\t{safe_latex}\n")
    map_path.write_text("".join(lines), encoding="utf-16")
    return map_path


def render_formulas_with_word(output_path: Path) -> bool:
    if not FORMULA_PLACEHOLDERS or not WORD_FORMULA_SCRIPT.exists():
        return False
    map_path = write_formula_map(output_path)
    try:
        subprocess.run(
            ["cscript", "//nologo", str(WORD_FORMULA_SCRIPT), str(output_path), str(map_path)],
            check=True,
            capture_output=True,
            text=True,
        )
        try:
            map_path.unlink()
        except OSError:
            pass
        return True
    except Exception as exc:  # pragma: no cover - local Word automation is environment-specific.
        print(f"Word formula post-process failed: {exc}", file=sys.stderr)
        return False


def add_inline_math(paragraph, latex: str):
    try:
        paragraph._p.append(latex_to_omath(latex))
        return True
    except Exception:
        if sys.platform.startswith("win"):
            run = paragraph.add_run(register_formula_placeholder(latex))
            set_run_font(run, cn_font=MATH_FONT, en_font=MATH_FONT, size=11)
            return False
        run = paragraph.add_run(simplify_latex(latex))
        set_run_font(run, cn_font=MATH_FONT, en_font=MATH_FONT, size=11)
        return False


def render_inline(paragraph, text: str, size: float = 12):
    # Supports the limited Markdown syntax used in the thesis source:
    # bold keywords, inline code identifiers and LaTeX inline formulas.
    token_re = re.compile(r"(\\\(.+?\\\)|\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for match in token_re.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=size)
        token = match.group(0)
        if token.startswith(r"\("):
            add_inline_math(paragraph, token[2:-2])
        elif token.startswith("**"):
            run = paragraph.add_run(simplify_latex(token[2:-2]))
            run.bold = True
            set_run_font(run, size=size)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, cn_font="Consolas", en_font="Consolas", size=size)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size)


def body_paragraph(document: Document, text: str):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    render_inline(paragraph, text)
    return paragraph


def reference_paragraph(document: Document, text: str):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.74)
    paragraph.paragraph_format.first_line_indent = Cm(-0.74)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    render_inline(paragraph, text)
    return paragraph


def add_heading(document: Document, text: str, level: int):
    paragraph = document.add_paragraph(style=f"Heading {level}")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 6)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run("摘  要" if text == "摘要" else text)
    run.bold = True
    set_run_font(run, cn_font=HEADING_FONT, en_font=EN_FONT, size={1: 16, 2: 14, 3: 12}.get(level, 12))
    return paragraph


def add_formula(document: Document, formula: str):
    formula = formula.strip()
    if not formula:
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    try:
        paragraph._p.append(latex_to_omath(formula))
    except Exception:
        if sys.platform.startswith("win"):
            run = paragraph.add_run(register_formula_placeholder(formula))
            set_run_font(run, cn_font=MATH_FONT, en_font=MATH_FONT, size=11)
            return
        run = paragraph.add_run(simplify_latex(formula))
        set_run_font(run, cn_font=MATH_FONT, en_font=MATH_FONT, size=11)


def split_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [cell.strip() for cell in line.split("|")]


def is_table_separator(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)


def add_table(document: Document, table_lines: list[str]):
    rows = [split_table_row(line) for line in table_lines if not is_table_separator(line)]
    rows = [row for row in rows if row]
    if not rows:
        return

    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True

    font_size = 8.5 if column_count >= 6 else 10.5
    for r_idx, row in enumerate(rows):
        for c_idx in range(column_count):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            text = row[c_idx] if c_idx < len(row) else ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.line_spacing = 1.0
            render_inline(paragraph, text, size=font_size)
            for run in paragraph.runs:
                run.bold = r_idx == 0

    document.add_paragraph()


def add_caption(document: Document, text: str):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    render_inline(paragraph, text, size=10.5)
    return paragraph


def image_path_from_markdown(src: str) -> Path:
    src = src.strip().strip('"').strip("'")
    src = src.replace("/", "\\")
    path = Path(src)
    if path.is_absolute():
        return path
    return (SOURCE_MD.parent / path).resolve()


def add_image(document: Document, alt: str, src: str):
    path = image_path_from_markdown(src)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    if not path.exists():
        run = paragraph.add_run(f"[图片缺失：{alt}，路径：{src}]")
        set_run_font(run, size=10.5)
        return
    run = paragraph.add_run()
    run.add_picture(str(path), width=Cm(15.2))
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Cm(0)
    caption.paragraph_format.space_before = Pt(0)
    caption.paragraph_format.space_after = Pt(6)
    caption_run = caption.add_run(alt)
    set_run_font(caption_run, size=10.5)


def add_title_page(document: Document):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(120)
    run = paragraph.add_run("本科毕业论文")
    run.bold = True
    set_run_font(run, cn_font=HEADING_FONT, en_font=EN_FONT, size=24)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(48)
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(THESIS_TITLE)
    run.bold = True
    set_run_font(run, cn_font=HEADING_FONT, en_font=EN_FONT, size=22)

    document.add_page_break()


def add_toc_field(document: Document):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(12)
    title_run = paragraph.add_run("目  录")
    title_run.bold = True
    set_run_font(title_run, cn_font=HEADING_FONT, en_font=EN_FONT, size=16)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run()

    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    run._r.append(instr)

    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_char)

    text = OxmlElement("w:t")
    text.text = "打开 Word 后右键更新域可生成目录页码。"
    run._r.append(text)

    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    set_run_font(run, size=12)
    document.add_page_break()


def enable_field_update_prompt(document: Document):
    settings = document.settings._element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def generate_docx():
    if not SOURCE_MD.exists():
        raise FileNotFoundError(SOURCE_MD)
    if not TEMPLATE_DOCX.exists():
        raise FileNotFoundError(TEMPLATE_DOCX)
    FORMULA_PLACEHOLDERS.clear()

    document = Document(str(TEMPLATE_DOCX))
    clear_body(document)
    configure_document(document)
    enable_field_update_prompt(document)
    document.core_properties.title = THESIS_TITLE

    add_title_page(document)

    lines = SOURCE_MD.read_text(encoding="utf-8-sig").splitlines()
    i = 0
    toc_inserted = False
    in_references = False

    while i < len(lines):
        raw = lines[i].rstrip()
        stripped = raw.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "$$":
            formula_lines: list[str] = []
            i += 1
            while i < len(lines) and lines[i].strip() != "$$":
                formula_lines.append(lines[i].rstrip())
                i += 1
            add_formula(document, "\n".join(formula_lines))
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].rstrip())
                i += 1
            add_table(document, table_lines)
            continue

        if re.match(r"^[表图]\s*\d+-\d+\s+[^，。；：]+$", stripped):
            add_caption(document, stripped)
            i += 1
            continue

        image_match = re.fullmatch(r"!\[(.*?)\]\((.*?)\)", stripped)
        if image_match:
            add_image(document, image_match.group(1), image_match.group(2))
            i += 1
            continue

        html_image_match = re.search(r'<img\s+[^>]*src="([^"]+)"[^>]*>', stripped)
        if html_image_match:
            add_image(document, "图片", html_image_match.group(1))
            i += 1
            continue

        heading_match = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading_match:
            level = min(len(heading_match.group(1)), 3)
            text = heading_match.group(2).strip()
            if level == 1:
                if text.startswith("第一章") and not toc_inserted:
                    document.add_page_break()
                    add_toc_field(document)
                    toc_inserted = True
                elif text not in {"摘要"}:
                    document.add_page_break()
                in_references = text == "参考文献"
            add_heading(document, text, level)
            i += 1
            continue

        if in_references or re.match(r"^\[\d+\]", stripped):
            reference_paragraph(document, stripped)
        else:
            body_paragraph(document, stripped)
        i += 1

    try:
        document.save(str(OUTPUT_DOCX))
        output_path = OUTPUT_DOCX
    except PermissionError:
        document.save(str(FORMULA_OUTPUT_DOCX))
        output_path = FORMULA_OUTPUT_DOCX

    render_formulas_with_word(output_path)
    return output_path


if __name__ == "__main__":
    output = generate_docx()
    print(output)
