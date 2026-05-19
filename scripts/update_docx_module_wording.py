from __future__ import annotations

import sys
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]


REPLACEMENTS = {
    "核心功能层负责系统主要处理流程，包括数据转换、流程调度、梯度计算、数据优化、OpenGL 计算管理和结果导出。该层是系统功能的主要承载部分，向上连接界面操作，向下读取和写回内部数据对象。":
    "核心功能层负责系统主要处理流程，包括数据转换、梯度计算、数据优化、OpenGL 计算管理和结果导出。该层是系统功能的主要承载部分，向上连接界面操作，向下读取和写回内部数据对象。界面层或测试程序可通过统一调用接口组织具体处理流程，但该调用关系不单独作为系统功能模块划分。",

    "流程调度模块负责连接数据准备、算法计算、结果写回和文件导出。该模块接收来自界面层或测试程序的请求，检查数据集、字段名称和参数有效性，然后调用对应功能模块，并记录输出字段名称、字段关联方式、分量数和耗时信息。":
    "结果导出模块负责将新增字段写回内部数据对象，并重建为可保存的 VTK 数据集。该模块需要保留输出字段名称、字段关联方式和分量数等语义信息，保证导出结果能够被 ParaView 等工具继续读取和观察。在具体实现中，界面层或测试程序通过统一调用接口组织数据准备、算法计算和结果导出流程；该统一调用关系属于实现层封装，不单独作为系统功能模块划分。",

    "在系统设计方面，本文将系统划分为数据准备模块、梯度计算模块、数据优化模块和 OpenGL 计算管理模块。数据准备模块负责读取 VTK 数据并构造内部 DataObject；梯度计算模块根据网格类型和字段关联方式选择有限差分路径或形函数导数路径；数据优化模块基于图邻域执行双边滤波和多尺度融合；OpenGL 计算管理模块负责上下文、计算着色器、SSBO 和 GPU 计时。通过这种划分，系统能够在统一数据结构上组织不同算法，并保持结果写回和导出过程中的字段语义一致。":
    "在系统设计方面，本文将系统划分为数据准备模块、梯度计算模块、数据优化模块、OpenGL 计算管理模块和结果导出模块。数据准备模块负责读取 VTK 数据并构造内部 DataObject；梯度计算模块根据网格类型和字段关联方式选择有限差分路径或形函数导数路径；数据优化模块基于图邻域执行双边滤波和多尺度融合；OpenGL 计算管理模块负责上下文、计算着色器、SSBO 和 GPU 计时；结果导出模块负责将新增字段写回内部数据对象并重建 VTK 数据集。通过这种划分，系统能够在统一数据结构上组织不同算法，并保持结果写回和导出过程中的字段语义一致。",
}


def replace_paragraph_text(paragraph, new_text: str) -> None:
    runs = list(paragraph.runs)
    template = runs[0] if runs else None
    for run in runs:
        run._element.getparent().remove(run._element)
    new_run = paragraph.add_run(new_text)
    if template is not None:
        new_run.bold = template.bold
        new_run.italic = template.italic
        new_run.underline = template.underline
        if template.font is not None:
            new_run.font.name = template.font.name
            new_run.font.size = template.font.size
            new_run.font.bold = template.font.bold
            new_run.font.italic = template.font.italic
            new_run.font.underline = template.font.underline
            try:
                new_run._element.rPr.rFonts.set(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia",
                    template._element.rPr.rFonts.get(
                        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia"
                    ),
                )
            except Exception:
                pass


def update_docx(path: Path) -> int:
    doc = Document(path)
    changed = 0
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if text in REPLACEMENTS:
            replace_paragraph_text(paragraph, REPLACEMENTS[text])
            changed += 1
    doc.save(path)
    return changed


def main() -> int:
    if len(sys.argv) > 1:
        target = Path(sys.argv[1])
    else:
        target = next(
            p for p in (ROOT / "wendang").glob("*修订版_数据优化_结构化网格.docx")
        )

    changed = update_docx(target)
    print(f"updated: {target}")
    print(f"changed_paragraphs: {changed}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
