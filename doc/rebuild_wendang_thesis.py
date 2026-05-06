from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
WENDANG = ROOT / "wendang"
DOC = ROOT / "doc"
TITLE = "基于OpenGL的CAE软件数据预处理方法研究与实践"
OUT_NAME = f"{TITLE}_本科毕业论文"
CHAPTERS = [
    "00_摘要.md",
    "01_绪论.md",
    "02_相关技术与理论基础.md",
    "03_系统需求分析与总体设计.md",
    "04_核心算法设计与实现.md",
    "05_实验设计与结果分析.md",
    "06_总结与展望.md",
    "07_参考文献.md",
]


def read_chapters(image_prefix: str = "assets") -> str:
    parts = [f"# {TITLE}"]
    for name in CHAPTERS:
        text = (WENDANG / name).read_text(encoding="utf-8").strip()
        if image_prefix != "assets":
            text = text.replace("](assets/", f"]({image_prefix}/")
        parts.append(text)
    return "\n\n".join(parts) + "\n"


def set_run_font(run, size: float = 10.5, font_name: str = "宋体", bold: bool = False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_doc_defaults(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.6)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size in [
        ("Title", 18),
        ("Heading 1", 16),
        ("Heading 2", 14),
        ("Heading 3", 12),
    ]:
        style = doc.styles[style_name]
        style.font.name = "黑体"
        style.font.size = Pt(size)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_paragraph(doc: Document, text: str, style: str | None = None, center: bool = False, indent: bool = True):
    p = doc.add_paragraph(style=style)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if indent and style is None:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    if style in {"Title", "Heading 1", "Heading 2", "Heading 3"}:
        set_run_font(run, 18 if style == "Title" else 16 if style == "Heading 1" else 14 if style == "Heading 2" else 12, "黑体", True)
    else:
        set_run_font(run)
    return p


def add_formula(doc: Document, lines: list[str]):
    text = "\n".join(lines).strip()
    if not text:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, 10.5, "Times New Roman")


def is_table_separator(line: str) -> bool:
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", c or "") for c in cells)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
        if not is_table_separator(lines[i]):
            rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
        i += 1
    return rows, i


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(cell_text)
            set_run_font(run, 10 if r_idx else 10.5, "宋体", r_idx == 0)


def image_path_from_markdown(src: str) -> Path:
    src = src.replace("/", "\\")
    if src.startswith("assets\\"):
        return WENDANG / src
    return (WENDANG / src).resolve()


def add_image(doc: Document, alt: str, src: str):
    path = image_path_from_markdown(src)
    if not path.exists():
        add_paragraph(doc, f"[图片缺失：{alt}，路径：{src}]", indent=False)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(15.2))


def markdown_to_docx(md_text: str, out_path: Path):
    doc = Document()
    set_doc_defaults(doc)
    lines = md_text.splitlines()
    i = 0
    in_formula = False
    formula_lines: list[str] = []

    while i < len(lines):
        raw = lines[i]
        line = raw.strip()

        if line == "$$":
            if in_formula:
                add_formula(doc, formula_lines)
                formula_lines = []
                in_formula = False
            else:
                in_formula = True
            i += 1
            continue

        if in_formula:
            formula_lines.append(raw)
            i += 1
            continue

        if not line:
            i += 1
            continue

        image_match = re.fullmatch(r"!\[(.*?)\]\((.*?)\)", line)
        if image_match:
            add_image(doc, image_match.group(1), image_match.group(2))
            i += 1
            continue

        if line.startswith("|") and line.endswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue

        if line.startswith("# "):
            title = line[2:].strip()
            if title == TITLE:
                add_paragraph(doc, title, style="Title", center=True, indent=False)
                add_paragraph(doc, "本科毕业论文", center=True, indent=False)
                doc.add_page_break()
            else:
                if title.startswith(("第一章", "第二章", "第三章", "第四章", "第五章", "第六章", "参考文献")):
                    doc.add_page_break()
                add_paragraph(doc, title, style="Heading 1", center=True, indent=False)
            i += 1
            continue

        if line.startswith("## "):
            add_paragraph(doc, line[3:].strip(), style="Heading 2", indent=False)
            i += 1
            continue

        if line.startswith("### "):
            add_paragraph(doc, line[4:].strip(), style="Heading 3", indent=False)
            i += 1
            continue

        if line.startswith("> "):
            add_paragraph(doc, line[2:].strip(), indent=False)
            i += 1
            continue

        add_paragraph(doc, line)
        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def main():
    wendang_md = read_chapters("assets")
    wendang_md_path = WENDANG / f"{OUT_NAME}.md"
    wendang_docx_path = WENDANG / f"{OUT_NAME}.docx"
    wendang_md_path.write_text(wendang_md, encoding="utf-8")
    markdown_to_docx(wendang_md, wendang_docx_path)

    if DOC.exists():
        doc_md = read_chapters("../wendang/assets")
        doc_md_path = DOC / f"{OUT_NAME}.md"
        doc_docx_path = DOC / f"{OUT_NAME}.docx"
        doc_md_path.write_text(doc_md, encoding="utf-8")
        markdown_to_docx(wendang_md, doc_docx_path)

    print(f"rebuilt {wendang_md_path}")
    print(f"rebuilt {wendang_docx_path}")


if __name__ == "__main__":
    main()
