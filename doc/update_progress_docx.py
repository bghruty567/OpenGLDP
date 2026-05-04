from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "doc" / "进展.docx"


def set_run_font(run, size_pt: float = 12, bold: bool = False, name: str = "Times New Roman") -> None:
    run.font.name = name
    run.bold = bold
    run.font.size = Pt(size_pt)
    rpr = run._element.rPr
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        run._element.insert(0, rpr)
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), "宋体")
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rpr.append(rfonts)


def setup_document(doc: Document) -> None:
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.6)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.2)


def add_paragraph(doc: Document, text: str, *, bold: bool = False, first_indent: bool = True) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, 12, bold)


def add_label(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, 12, True, "黑体")


def add_table(doc: Document, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header = table.rows[0].cells
    for i, value in enumerate(rows[0]):
        p = header[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(value)
        set_run_font(run, 10.5, True)
        header[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for row in rows[1:]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(value)
            set_run_font(run, 10.5)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def save_with_fallback(doc: Document, path: Path) -> Path:
    try:
        doc.save(path)
        return path
    except PermissionError:
        fallback = path.with_name(f"{path.stem}_更新版{path.suffix}")
        try:
            doc.save(fallback)
            return fallback
        except PermissionError:
            stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            fallback2 = path.with_name(f"{path.stem}_{stamp}{path.suffix}")
            doc.save(fallback2)
            return fallback2


def build_doc() -> Document:
    doc = Document()
    setup_document(doc)

    add_paragraph(doc, "最近主要是对上次汇报中老师提出的两个主要问题进行了改进。")

    add_paragraph(
        doc,
        "1. 首先是关于非结构化网格梯度计算功能的精度问题以及整个梯度计算模块的实验内容问题。目前已将方法调整为与 vtkGradientFilter 更一致的基于形函数导数的梯度计算方法，支持一阶三角形、四边形、四面体、六面体四种单元类型。",
    )

    add_paragraph(doc, "实验内容上目前做了三类实验：")

    add_label(doc, "解析场验证实验")
    add_paragraph(
        doc,
        "这组实验的目标是证明通过将系统计算结果与解析真值进行对比，验证算法本身的正确性。数据集使用 SampleStructGrid（规则网格）、hexa（非结构化三维网格）和 1_0（非结构化二维曲面网格）。其中 1_0 采用曲面切向解析场，并以内在梯度指标作为主口径，以避免法向分量对结论的干扰。",
    )

    add_label(doc, "线性标量场")
    add_paragraph(
        doc,
        "三维数据集在线性标量场实验中，均在归一化局部坐标上构造一次标量函数；1_0 数据集则在曲面局部切向坐标上构造一次标量函数。实验结果如下。",
    )
    add_table(
        doc,
        [
            ["数据集", "场函数说明", "NRMSE"],
            ["SampleStructGrid", "三维线性标量场", "2.82038e-07"],
            ["hexa", "三维线性标量场", "1.57998e-07"],
            ["1_0", "曲面切向线性标量场", "1.68088e-07"],
        ],
    )

    add_label(doc, "线性向量场")
    add_paragraph(
        doc,
        "三维数据集在线性向量场实验中，分别为三个分量构造一次函数；1_0 数据集则在曲面局部切向坐标上分别构造各分量，并投影到曲面局部支撑空间。实验结果如下。",
    )
    add_table(
        doc,
        [
            ["数据集", "场函数说明", "NRMSE"],
            ["SampleStructGrid", "三维线性向量场", "3.30256e-07"],
            ["hexa", "三维线性向量场", "2.36149e-07"],
            ["1_0", "曲面切向线性向量场", "1.32859e-07"],
        ],
    )
    add_paragraph(
        doc,
        "从以上结果可以看出，当前系统在规则网格、三维六面体非结构化网格以及二维曲面四边形网格上的点数据解析场验证结果均达到 1e-7 量级，可以作为算法正确性的直接证据。",
    )

    add_label(doc, "与 VTK 结果一致性对比实验")
    add_paragraph(
        doc,
        "这组实验的目标是证明：在真实数据字段上，当前系统与 vtkGradientFilter 的工程输出是一致的。该实验主要说明工程一致性，不单独承担算法数学正确性的证明任务。",
    )

    add_label(doc, "点数据")
    add_table(
        doc,
        [
            ["数据集", "字段", "NRMSE"],
            ["SampleStructGrid", "scalars", "8.22437e-08"],
            ["hexa", "scalars", "6.27285e-08"],
            ["1_0", "RF", "5.63477e-08"],
        ],
    )

    add_label(doc, "单元数据")
    add_table(
        doc,
        [
            ["数据集", "字段", "NRMSE"],
            ["SampleStructGrid", "scalars", "7.62786e-07"],
            ["limb", "chem_0", "3.89516e-07"],
            ["1_0", "S_Mises", "7.77806e-08"],
        ],
    )
    add_paragraph(
        doc,
        "这组结果表明，在当前展示的规则网格、二维曲面四边形网格和三维六面体网格案例上，系统结果与 vtkGradientFilter 保持了较高一致性，主要误差指标处于 1e-7 到 1e-6 量级。",
    )

    add_label(doc, "与 VTK 的时间对比实验")
    add_paragraph(
        doc,
        "这组实验的目标是在控制变量成立的前提下，对比系统实现与 VTK 并行实现的计算时间，分析当前 OpenGL 实现的工程效率。该部分不再直接采用异构真实模型，而是采用统一生成的同构网格族进行测试，以保证实验结论具有可比性和可解释性。实验中统一采用 POINT 路径、统一字段名 scalars、统一 VTK 并行 backend 为 STDThread，并固定 VTK 并行线程数为 16。",
    )

    add_label(doc, "规则网格")
    add_table(
        doc,
        [
            ["数据集", "点数/单元数", "VTK 并行线程数", "VTK 并行时间/ms", "系统总时间/ms", "GPU 时间/ms"],
            ["timing_struct_20x20x20", "8000 / 6859", "16", "2.7573", "0.9036", "0.025272"],
            ["timing_struct_32x32x32", "32768 / 29791", "16", "3.5564", "1.4336", "0.043732"],
            ["timing_struct_48x48x48", "110592 / 103823", "16", "8.2716", "4.0996", "0.107692"],
        ],
    )

    add_label(doc, "非结构化网格")
    add_table(
        doc,
        [
            ["数据集", "点数/单元数", "VTK 并行线程数", "VTK 并行时间/ms", "系统总时间/ms", "GPU 时间/ms"],
            ["timing_uhex_20x20x20", "8000 / 6859", "16", "163.799", "2.8527", "0.623896"],
            ["timing_uhex_32x32x32", "32768 / 29791", "16", "708.650", "5.3835", "2.53781"],
            ["timing_uhex_48x48x48", "110592 / 103823", "16", "3349.36", "44.8039", "7.91991"],
        ],
    )
    add_paragraph(
        doc,
        "从时间实验结果可以看出，在当前测试条件下，系统的系统总时间与 GPU 时间均明显低于 VTK 并行时间，且随数据规模增大呈现稳定增长趋势。这说明当前基于 OpenGL 的梯度计算实现具备较稳定的 GPU 计算性能。需要说明的是，系统总时间反映的是算法调用级总耗时，GPU 时间反映的是 OpenGL 侧纯计算时间，而 VTK 并行时间对应的是 vtkGradientFilter::Update() 的执行时间，因此三者在统计口径上并不完全相同。",
    )

    add_paragraph(
        doc,
        "2. 对于之前数据优化模块实验与噪声原因不够对应的问题，我打算将表述收敛如下：",
    )
    add_paragraph(
        doc,
        "CAE 仿真数据会因离散化误差、迭代收敛误差与舍入误差等因素产生局部数值扰动。对于其中一类表现为局部随机高频波动的扰动，可采用高斯扰动作为代理模型进行近似描述。因此，本模块面向这类局部数值扰动，实验上通过在干净场上叠加高斯扰动并比较优化前后的误差与粗糙度指标，验证模块对局部随机高频扰动的抑制能力及其边缘保持特性。",
    )
    add_paragraph(
        doc,
        "当前实验在 ShipHull_0 数据集上构造多类干净场，并分别添加高斯扰动。下面给出点数据与单元数据路径上的平均结果。",
    )
    add_table(
        doc,
        [
            ["关联方式", "输入 NRMSE", "输出 NRMSE", "RMSE 改进比", "粗糙度比"],
            ["POINT", "0.296642", "0.180858", "0.610772", "0.406530"],
            ["CELL", "0.294558", "0.121921", "0.414806", "0.387052"],
        ],
    )
    add_paragraph(
        doc,
        "从实验结果可以看出，无论是点数据还是单元数据，经过数据优化后，输出误差均低于输入误差，说明该模块对局部随机高频扰动具有稳定的抑制效果。进一步结合粗糙度指标进行分析，还可以说明该模块在抑制局部波动的同时，能够较好地保持场中主要结构特征。",
    )

    add_paragraph(doc, "基于以上工作，当前我想请老师重点帮我确认两个问题。")
    add_paragraph(
        doc,
        "第一，梯度计算模块是否可以按当前口径收敛为：规则网格采用有限差分法，非结构化网格采用基于形函数导数的方法，并以解析场验证、与 VTK 结果一致性对比、与 VTK 时间对比三组实验作为论文主实验设计。",
    )
    add_paragraph(
        doc,
        "第二，数据优化模块是否可以按“面向一类局部随机高频数值扰动的抑制与边缘保持处理”这一表述进行收敛，而不再泛化为对所有仿真噪声问题的统一处理。",
    )

    return doc


def main() -> None:
    doc = build_doc()
    saved = save_with_fallback(doc, OUT_DOCX)
    print(saved)


if __name__ == "__main__":
    main()
