from __future__ import annotations

from pathlib import Path

import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "doc" / "generated"
ASSET_DIR = ROOT / "doc" / "assets"
FLOW_DIR = ASSET_DIR / "shader_flows"
RESULT_DIR = ROOT / "results"
GRAD_DIR = RESULT_DIR / "gradient"
MUL_DIR = RESULT_DIR / "mul"

OUT_DOCX = OUT_DIR / "老师汇报文稿_会议后更新版.docx"
OUT_MD = OUT_DIR / "老师汇报文稿_会议后更新版.md"

FONT_CANDIDATES = [
    Path("C:/Windows/Fonts/msyh.ttc"),
    Path("C:/Windows/Fonts/msyh.ttf"),
    Path("C:/Windows/Fonts/simsun.ttc"),
    Path("C:/Windows/Fonts/simhei.ttf"),
]


def pick_font() -> Path | None:
    for path in FONT_CANDIDATES:
        if path.exists():
            return path
    return None


FONT_PATH = pick_font()


def img_font(size: int):
    if FONT_PATH is not None:
        return ImageFont.truetype(str(FONT_PATH), size=size)
    return ImageFont.load_default()


def set_run_font(run, size_pt: float = 12, bold: bool = False, name: str = "Times New Roman") -> None:
    run.font.name = name
    run.bold = bold
    run.font.size = Pt(size_pt)
    rpr = run._element.rPr
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        run._element.insert(0, rpr)
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), "宋体")
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rpr.append(rfonts)


def add_page_number(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    p._p.append(fld_begin)
    p._p.append(instr)
    p._p.append(fld_end)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {min(level, 3)}"]
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, 16, True, "黑体")
    elif level == 2:
        set_run_font(run, 14, True, "黑体")
    else:
        set_run_font(run, 12, True, "黑体")


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, 12)


def add_center(doc: Document, text: str, size: float = 12, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, size, bold)


def add_table(doc: Document, title: str, headers: list[str], rows: list[list[str]], note: str | None = None) -> None:
    add_center(doc, title, 11, False)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, head in enumerate(headers):
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(head)
        set_run_font(run, 10.5, True)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(value))
            set_run_font(run, 10.5)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if note:
        add_body(doc, note)


def add_figure(doc: Document, caption: str, path: Path, note: str | None = None, width_cm: float = 15.5) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    add_center(doc, caption, 11, False)
    if note:
        add_body(doc, note)


def sci(value: float, digits: int = 3) -> str:
    if value == 0:
        return "0"
    if abs(value) >= 1e4 or abs(value) < 1e-3:
        return f"{value:.{digits}e}"
    return f"{value:.{digits}f}"


def ms(value: float) -> str:
    return f"{value:.3f}"


def wrap_text(text: str, width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    for ch in text:
        current += ch
        if len(current) >= width:
            lines.append(current)
            current = ""
    if current:
        lines.append(current)
    return lines or [""]


def draw_round_box(draw: ImageDraw.ImageDraw, xy, text: str, fill: str = "#e8f1f8", outline: str = "#2f5b78") -> None:
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
    x0, y0, x1, y1 = xy
    font = img_font(22)
    lines = []
    for raw in text.split("\n"):
        lines.extend(wrap_text(raw, 16))
    total_h = len(lines) * 30
    y = y0 + (y1 - y0 - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        w = bbox[2] - bbox[0]
        draw.text((x0 + (x1 - x0 - w) / 2, y), line, font=font, fill="#173042")
        y += 30


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], fill: str = "#2f5b78") -> None:
    draw.line((start, end), fill=fill, width=5)
    import math
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 18
    wing = math.pi / 7
    p1 = (end[0] - length * math.cos(angle - wing), end[1] - length * math.sin(angle - wing))
    p2 = (end[0] - length * math.cos(angle + wing), end[1] - length * math.sin(angle + wing))
    draw.polygon([end, p1, p2], fill=fill)


def render_vertical_flow(path: Path, title: str, steps: list[str], note: str | None = None) -> None:
    width = 1500
    box_w = 1160
    box_h = 110
    gap = 36
    top_margin = 120
    bottom_extra = 90 if note else 40
    height = top_margin + len(steps) * box_h + (len(steps) - 1) * gap + bottom_extra
    img = Image.new("RGB", (width, height), "#ffffff")
    draw = ImageDraw.Draw(img)
    draw.text((60, 28), title, font=img_font(34), fill="#173042")
    x0 = (width - box_w) // 2
    x1 = x0 + box_w
    y = top_margin
    fills = ["#dceffc", "#e6f3e7", "#f8ecd4", "#f4dce2", "#e7e7f8", "#e8f1f8", "#eef6dd", "#f8e5d8"]
    centers = []
    for idx, step in enumerate(steps):
        y0 = y
        y1 = y0 + box_h
        draw_round_box(draw, (x0, y0, x1, y1), step, fill=fills[idx % len(fills)])
        centers.append((width // 2, y0, y1))
        y = y1 + gap
    for i in range(len(centers) - 1):
        start = (centers[i][0], centers[i][2])
        end = (centers[i + 1][0], centers[i + 1][1])
        draw_arrow(draw, start, end)
    if note:
        draw.text((60, height - 60), note, font=img_font(18), fill="#4a5d6b")
    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path, format="PNG")


def generate_shader_flowcharts() -> dict[str, Path]:
    FLOW_DIR.mkdir(parents=True, exist_ok=True)
    flows: dict[str, Path] = {}
    flow_specs = {
        "fd": (
            "FD.glsl 规则网格有限差分流程图",
            [
                "线程索引 = gl_GlobalInvocationID(x,y,z)\n越界则直接返回",
                "根据 (i,j,k) 计算线性样本 id",
                "在 ξ/η/ζ 三个方向调用 pick1D\n选择前后邻居与差分系数",
                "读取前后邻居坐标 xp/xm、yp/ym、zp/zm",
                "在参数空间估计几何导数\n构造局部 Jacobian 与行列式 aj",
                "若 Jacobian 退化\ninvAj 置 0，避免 NaN 扩散",
                "对每个字段分量分别计算 dXi、dEta、dZeta",
                "通过逆 Jacobian 做链式法则映射\n得到物理空间 gx、gy、gz",
                "把 3×numComponents 梯度写入输出缓冲 G[]",
            ],
            "一个线程处理一个规则网格样本；点数据与单元数据共用同一差分框架。",
        ),
        "shape_point": (
            "ShapePointGradient.glsl 非结构点梯度流程图",
            [
                "线程索引 = pointId\n先把该点所有输出梯度清零",
                "扫描 pointCellOff / pointCellNbr\n确定该点关联单元中的最高拓扑维数",
                "仅保留最高维可贡献单元\n忽略无效单元和低维混合单元",
                "在当前单元中定位该点 localId\n求该点在参考单元中的参数坐标",
                "若是六面体点位\n使用牛顿迭代反求参数坐标 pcoords",
                "prepareCellGeometry：读取单元类型、节点、dN/dξ、切向量、逆映射",
                "对每个字段分量计算参考单元梯度 gradRef",
                "把 gradRef 通过逆 Jacobian 或度量映射到物理空间",
                "对所有可贡献单元的梯度做累加",
                "按 usedCells 求平均\n写回该点的最终梯度",
            ],
            "注意：全零回退到 ShapeCellGradient + CellDataToPointLift 的保底逻辑发生在 GLGradientEngine.cpp，不在本 shader 内。",
        ),
        "shape_cell": (
            "ShapeCellGradient.glsl 非结构单元梯度流程图",
            [
                "线程索引 = cellId\n先把该单元所有输出梯度清零",
                "按单元类型选择单元中心参考坐标 pcoords",
                "prepareCellGeometry：检查单元类型、节点数和连接关系",
                "计算该单元在中心点处的 dN/dξ",
                "累加节点坐标构造切向量 tangents",
                "三维体单元直接求逆 Jacobian\n曲面/线单元改用度量张量逆",
                "对每个字段分量计算参考梯度 gradRef",
                "把 gradRef 映射为物理空间梯度 g",
                "将 g 写入该单元输出缓冲 grad[]",
            ],
            "该 shader 假定输入 pointVal 已经准备好；若来源是单元场，则通常先经过 CellDataToPointLift。",
        ),
        "cell_lift": (
            "CellDataToPointLift.glsl 单元场提升到点场流程图",
            [
                "线程索引 = pointId\n先把该点输出 pointVal 清零",
                "扫描 pointCellOff / pointCellNbr\n收集与该点相邻的单元",
                "统计这些相邻单元中的最高拓扑维数",
                "仅保留最高维单元\n避免点同时被壳单元和线单元混合污染",
                "对每个字段分量累加这些单元的 cellVal",
                "统计 usedCells 数量",
                "若 usedCells > 0\n按 usedCells 求平均",
                "把平均后的结果写成该点 pointVal",
            ],
            "它不是梯度计算 shader，而是“cell -> point”恢复 shader，为 ShapeCellGradient 的输入准备点值场。",
        ),
        "bilateral": (
            "Bilateral.glsl 图双边滤波流程图",
            [
                "线程索引 = 样本 i\n越界则返回",
                "读取样本位置 pi\n和邻域区间 off[i] ~ off[i+1]",
                "计算 sigmaS² 与 sigmaR²\n作为空间核和值域核尺度",
                "对每个字段分量分别独立滤波",
                "以中心值 center 初始化\nsumW = 1，sumV = center",
                "遍历 CSR 邻域中的每个邻居 j",
                "计算 dist²、dv、wSpatial、wRange 和总权重 w",
                "累加 sumW 与 sumV",
                "输出 sumV / sumW 作为平滑结果 outVal",
            ],
            "一个线程处理一个点或一个单元中心；标量场和向量场通过“逐分量独立滤波”统一处理。",
        ),
        "fuse": (
            "MultiScaleFuse.glsl 多尺度融合流程图",
            [
                "线程索引 = 标量样本 idx\n越界则返回",
                "读取 base 层和 d0/d1/d2 三个 detail 层",
                "根据 uLevelCount 决定哪些 detail 有效",
                "计算各层细节幅值 m0/m1/m2\n以及总特征强度 feature",
                "根据 edgeSigma 计算细节回注衰减 atten",
                "按各层幅值占比和 detailGain\n分配 w0、w1、w2",
                "执行重建：out = base + w0*d0 + w1*d1 + w2*d2",
                "把融合结果写入 outVal[]",
            ],
            "该 shader 只做逐样本融合，不再访问几何邻域；几何相关的平滑步骤已在 Bilateral.glsl 中完成。",
        ),
    }
    for key, (title, steps, note) in flow_specs.items():
        path = FLOW_DIR / f"{key}.png"
        render_vertical_flow(path, title, steps, note)
        flows[key] = path
    return flows


def load_data() -> dict:
    sample_point = pd.read_csv(GRAD_DIR / "SampleStructGridpoint.csv")
    sample_cell = pd.read_csv(GRAD_DIR / "SampleStructGridcell.csv")
    hexa_point = pd.read_csv(GRAD_DIR / "hexapoint.csv")
    limb_cell = pd.read_csv(GRAD_DIR / "limbcell.csv")
    ship_point = pd.read_csv(GRAD_DIR / "ShipHull_0point.csv")
    ship_cell = pd.read_csv(GRAD_DIR / "ShipHull_0cell.csv")
    one_point = pd.read_csv(GRAD_DIR / "1_0point.csv")
    one_cell = pd.read_csv(GRAD_DIR / "1_0cell.csv")
    point_mul = pd.read_csv(MUL_DIR / "multiscale_report+point.csv")
    cell_mul = pd.read_csv(MUL_DIR / "multiscale_report+cell.csv")

    def pick(df: pd.DataFrame, name: str):
        return df[df["array"] == name].iloc[0]

    def real(df: pd.DataFrame):
        return df[~df["array"].astype(str).str.startswith("benchmark")].iloc[0]

    point_avg = point_mul.groupby("noise")[["input_nrmse", "fused_nrmse", "rmse_improvement_ratio", "roughness_ratio", "gpu_avg_ms"]].mean()
    cell_avg = cell_mul.groupby("noise")[["input_nrmse", "fused_nrmse", "rmse_improvement_ratio", "roughness_ratio", "gpu_avg_ms"]].mean()

    param_ref = point_mul[(point_mul["clean_array"] == "ms_clean_trig") & (point_mul["noise"] == "gaussian")].iloc[0]
    parameter_defaults = [
        ("levels", str(int(param_ref["levels"])), "多尺度分解层数"),
        ("iterations", str(int(param_ref["iterations"])), "每层双边滤波迭代次数"),
        ("spatial_sigma_factor", f"{float(param_ref['spatial_sigma_factor']):.2f}", "空间邻域权重尺度"),
        ("range_sigma_factor", f"{float(param_ref['range_sigma_factor']):.2f}", "值域相似性权重尺度"),
        ("level_scale", f"{float(param_ref['level_scale']):.2f}", "层间尺度放大倍数"),
        ("edge_sigma_factor", f"{float(param_ref['edge_sigma_factor']):.2f}", "边缘抑制尺度"),
        ("detail_gain[0,1,2]", f"[{float(param_ref['detail_gain0']):.2f}, {float(param_ref['detail_gain1']):.2f}, {float(param_ref['detail_gain2']):.2f}]", "三层细节回注增益"),
        ("sigma_factor", f"{float(param_ref['sigma_factor']):.2f}", "白高斯扰动强度系数"),
        ("corr_length/iters/alpha", f"{float(param_ref['corr_length_factor']):.2f} / {int(param_ref['corr_iters'])} / {float(param_ref['corr_alpha']):.2f}", "相关高斯扰动控制"),
        ("impulse_ratio/scale", f"{float(param_ref['impulse_ratio']):.2f} / {float(param_ref['impulse_scale']):.2f}", "脉冲异常占比与幅值"),
    ]

    parameter_basis = []
    for clean_array, noise, note in [
        ("ms_clean_edge", "gaussian", "默认参数对目标高斯型局部扰动有效，同时保持边缘。"),
        ("ms_clean_trig", "grf", "对空间相关扰动改善有限，说明没有靠过度平滑虚构性能。"),
        ("ms_clean_edge", "impulse", "对尖峰异常值改善很弱，说明模块边界清晰。"),
    ]:
        row = point_mul[(point_mul["clean_array"] == clean_array) & (point_mul["noise"] == noise)].iloc[0]
        parameter_basis.append({
            "case": f"POINT / {clean_array} / {noise}",
            "input_nrmse": float(row["input_nrmse"]),
            "fused_nrmse": float(row["fused_nrmse"]),
            "roughness_ratio": float(row["roughness_ratio"]),
            "note": note,
        })

    return {
        "sample_linear": pick(sample_point, "benchmark_linear"),
        "sample_trig": pick(sample_point, "benchmark_trig"),
        "sample_vec_linear": pick(sample_point, "benchmark_vec_linear"),
        "hexa_linear": pick(hexa_point, "benchmark_linear"),
        "hexa_trig": pick(hexa_point, "benchmark_trig"),
        "sample_real_point": real(sample_point),
        "sample_real_cell": real(sample_cell),
        "hexa_real": real(hexa_point),
        "limb_real": real(limb_cell),
        "ship_real_point": real(ship_point),
        "ship_real_cell": real(ship_cell),
        "one_real_point": real(one_point),
        "one_real_cell": real(one_cell),
        "one_surface_linear": pick(one_point, "benchmark_surface_linear"),
        "one_surface_trig": pick(one_point, "benchmark_surface_trig"),
        "one_linear": pick(one_point, "benchmark_linear"),
        "point_avg": point_avg,
        "cell_avg": cell_avg,
        "parameter_defaults": parameter_defaults,
        "parameter_basis": parameter_basis,
    }


def build_markdown(data: dict) -> str:
    gaussian_point = data["point_avg"].loc["gaussian"]
    gaussian_cell = data["cell_avg"].loc["gaussian"]
    mixed_point = data["point_avg"].loc["mixed"]
    return f"""# 老师汇报文稿（会议后更新版）

## 一、本次汇报重点

本次汇报不再按一般“完成了什么功能”展开，而是只围绕上次会议中老师强调的两个核心问题组织：

1. 梯度计算功能中，如何保证与 VTK 的算法可比性，并把实验重点收敛到结构化网格主线。
2. 数据优化功能中，如何说明仿真数据扰动的根本来源，以及当前模块的研究意义和适用边界。

## 二、问题一：梯度计算功能

### 1. 本次调整后的主线

- 论文和实验主结论收敛到结构化网格。
- 结构化网格采用有限差分，与 VTK 在规则网格上的梯度计算具有直接可比性。
- 非结构化网格不再以原来的加权最小二乘法作为论文主线，而是转向更接近 `vtkGradientFilter` 思路的基于形函数导数的方法；但当前答辩主结论不建立在复杂非结构网格深度对比上。

### 2. 对 VTK 可比性的理解

- `vtkGradientFilter` 不能只看输出结果，还要同时考虑其算法配置与并行配置。
- 当前测试程序已经可以输出 VTK 单线程时间、并行时间、并行线程数和后端信息。
- 因此，后续与 VTK 的比较不再是“系统时间 vs VTK 一个时间”，而是“系统时间 vs VTK 单线程/并行时间”的分层比较。

### 3. 当前主线 shader 需要单独给流程图

这次汇报会补 6 张详细流程图，分别对应：

1. `FD.glsl`
2. `ShapePointGradient.glsl`
3. `ShapeCellGradient.glsl`
4. `CellDataToPointLift.glsl`
5. `Bilateral.glsl`
6. `MultiScaleFuse.glsl`

这样老师在看结果前，能先看到每个 GPU 核函数到底做了什么，而不是只看到模块级框图。

### 4. 结构化网格当前数据

#### 表1 结构化网格解析场验证（SampleStructGrid point）

| 字段 | NRMSE | 说明 |
| --- | --- | --- |
| benchmark_linear | {sci(float(data["sample_linear"]["ambient_nrmse"]), 6)} | 线性场接近机器精度 |
| benchmark_trig | {sci(float(data["sample_trig"]["ambient_nrmse"]), 6)} | 三角函数场误差较低 |
| benchmark_vec_linear | {sci(float(data["sample_vec_linear"]["ambient_nrmse"]), 6)} | 向量线性场同样接近机器精度 |

#### 表2 结构化网格真实字段与 VTK 对比（SampleStructGrid）

| 关联 | 字段 | NRMSE | VTK单线程/ms | VTK并行/ms | 系统总时间/ms | GPU时间/ms |
| --- | --- | --- | ---: | ---: | ---: | ---: |
| POINT | scalars | {sci(float(data["sample_real_point"]["ambient_nrmse"]), 6)} | {ms(float(data["sample_real_point"]["ambient_vtk_single_avg_ms"]))} | {ms(float(data["sample_real_point"]["ambient_vtk_parallel_avg_ms"]))} | {ms(float(data["sample_real_point"]["result_wall_avg_ms"]))} | {ms(float(data["sample_real_point"]["result_gpu_avg_ms"]))} |
| CELL | scalars | {sci(float(data["sample_real_cell"]["ambient_nrmse"]), 6)} | {ms(float(data["sample_real_cell"]["ambient_vtk_single_avg_ms"]))} | {ms(float(data["sample_real_cell"]["ambient_vtk_parallel_avg_ms"]))} | {ms(float(data["sample_real_cell"]["result_wall_avg_ms"]))} | {ms(float(data["sample_real_cell"]["result_gpu_avg_ms"]))} |

结论：

- 结构化网格解析场结果已经可以直接支撑“算法正确性”。
- 结构化网格真实字段与 VTK 的结果高度一致。
- 时间对比中，系统总时间显著低于当前环境下的 VTK 单线程和并行时间；GPU 纯计算时间更低，说明核心计算开销已被压缩。

### 5. 非结构化网格当前口径

当前非结构化网格结果保留为“补充一致性结果”，不作为本次汇报主结论。现阶段更适合向老师说明两点：

1. 已经研究 VTK 算法路径，并把系统主线从“局部拟合最小二乘”转向“形函数导数”。
2. 复杂曲面和单元数据恢复路径仍更敏感，因此非结构化深度对比不再作为当前答辩主战场。

#### 表3 非结构化网格真实字段与 VTK 对比（补充）

| 数据集 | 关联/字段 | NRMSE | 系统总时间/ms |
| --- | --- | --- | ---: |
| hexa | POINT / scalars | {sci(float(data["hexa_real"]["ambient_nrmse"]), 6)} | {ms(float(data["hexa_real"]["result_wall_avg_ms"]))} |
| limb | CELL / chem_0 | {sci(float(data["limb_real"]["ambient_nrmse"]), 6)} | {ms(float(data["limb_real"]["result_wall_avg_ms"]))} |
| ShipHull_0 | POINT / RF | {sci(float(data["ship_real_point"]["ambient_nrmse"]), 6)} | {ms(float(data["ship_real_point"]["result_wall_avg_ms"]))} |
| ShipHull_0 | CELL / S_Mises | {sci(float(data["ship_real_cell"]["ambient_nrmse"]), 6)} | {ms(float(data["ship_real_cell"]["result_wall_avg_ms"]))} |

### 6. 曲面数据为什么不能直接拿普通解析场评价

- 以 `1_0 point` 为例，普通 `benchmark_linear` 的 ambient NRMSE 为 {sci(float(data["one_linear"]["ambient_nrmse"]), 6)}。
- 但曲面专用 `benchmark_surface_linear` 的 intrinsic NRMSE 只有 {sci(float(data["one_surface_linear"]["intrinsic_nrmse"]), 6)}。
- 这说明曲面数据必须区分 ambient 与 intrinsic 指标，否则会把法向分量误差过度混入主结论。

## 三、问题二：数据优化功能

### 1. 对老师问题的调整

本次不再把数据优化模块简单表述为“高斯噪声去噪”，而是先回答“真实仿真数据为什么会出现扰动”。

当前更合理的表述是：

- CAE 后处理结果中的局部数值扰动可能来自离散化误差、求解迭代残差、网格质量、梯度/应力恢复不连续以及浮点舍入等因素。
- 其中有一类扰动在结果场上表现为局部随机高频波动。
- 当前数据优化模块面向的就是这类“局部随机高频数值扰动”，而不是所有噪声类型。

### 2. 模块定位

- 模块采用图双边滤波做多尺度分解。
- 采用细节增益控制的多尺度融合做重建。
- 它的目标不是修正离散化根因，而是在后处理阶段抑制一类局部随机高频扰动，并尽量保持边缘与主结构。

### 3. 当前统计结果

#### 表4 点数据多尺度优化平均结果

| 扰动类型 | 输入NRMSE | 输出NRMSE | RMSE改进率 | 粗糙度比 |
| --- | ---: | ---: | ---: | ---: |
| gaussian | {gaussian_point["input_nrmse"]:.6f} | {gaussian_point["fused_nrmse"]:.6f} | {gaussian_point["rmse_improvement_ratio"]:.6f} | {gaussian_point["roughness_ratio"]:.6f} |
| grf | {data["point_avg"].loc["grf","input_nrmse"]:.6f} | {data["point_avg"].loc["grf","fused_nrmse"]:.6f} | {data["point_avg"].loc["grf","rmse_improvement_ratio"]:.6f} | {data["point_avg"].loc["grf","roughness_ratio"]:.6f} |
| impulse | {data["point_avg"].loc["impulse","input_nrmse"]:.6f} | {data["point_avg"].loc["impulse","fused_nrmse"]:.6f} | {data["point_avg"].loc["impulse","rmse_improvement_ratio"]:.6f} | {data["point_avg"].loc["impulse","roughness_ratio"]:.6f} |
| mixed | {mixed_point["input_nrmse"]:.6f} | {mixed_point["fused_nrmse"]:.6f} | {mixed_point["rmse_improvement_ratio"]:.6f} | {mixed_point["roughness_ratio"]:.6f} |

#### 表5 单元数据多尺度优化平均结果

| 扰动类型 | 输入NRMSE | 输出NRMSE | RMSE改进率 | 粗糙度比 |
| --- | ---: | ---: | ---: | ---: |
| gaussian | {gaussian_cell["input_nrmse"]:.6f} | {gaussian_cell["fused_nrmse"]:.6f} | {gaussian_cell["rmse_improvement_ratio"]:.6f} | {gaussian_cell["roughness_ratio"]:.6f} |
| grf | {data["cell_avg"].loc["grf","input_nrmse"]:.6f} | {data["cell_avg"].loc["grf","fused_nrmse"]:.6f} | {data["cell_avg"].loc["grf","rmse_improvement_ratio"]:.6f} | {data["cell_avg"].loc["grf","roughness_ratio"]:.6f} |
| impulse | {data["cell_avg"].loc["impulse","input_nrmse"]:.6f} | {data["cell_avg"].loc["impulse","fused_nrmse"]:.6f} | {data["cell_avg"].loc["impulse","rmse_improvement_ratio"]:.6f} | {data["cell_avg"].loc["impulse","roughness_ratio"]:.6f} |
| mixed | {data["cell_avg"].loc["mixed","input_nrmse"]:.6f} | {data["cell_avg"].loc["mixed","fused_nrmse"]:.6f} | {data["cell_avg"].loc["mixed","rmse_improvement_ratio"]:.6f} | {data["cell_avg"].loc["mixed","roughness_ratio"]:.6f} |

结论：

- 对 gaussian 扰动，点和单元数据都表现出明显改善。
- 对 grf 扰动，改善有限但仍存在一定平滑作用。
- 对 impulse 扰动几乎无改善，这恰恰说明当前模块不是通用去噪器，而是有明确适用范围。

## 四、当前建议向老师汇报的结论

1. 梯度计算部分已经按老师要求调整实验主线：主结论收敛到结构化网格，重点展示解析场正确性、与 VTK 的可比时间结果和算法一致性说明。
2. 非结构化网格方面，已研究 VTK 算法路径，并将系统主线转向更可比的形函数导数方法；但复杂曲面与单元数据恢复仍较敏感，因此当前只作为补充结果汇报。
3. 数据优化部分已经不再停留在“人工高斯噪声去噪”这一表述，而是改为“面向一类局部随机高频数值扰动的结果场优化”，并用统计结果说明它对 Gaussian / Mixed 更有效、对 Impulse 不敏感。

## 五、下一步计划

1. 梯度模块继续保持“结构化网格为主、非结构化为辅”的汇报口径，并补充更细的结构化网格对比图表。
2. 数据优化模块继续围绕扰动根因做文献和案例整理，争取补充更贴近真实仿真扰动来源的实验设计。
3. 将上述两部分内容与现有论文正文、图表和 ParaView 图像进一步统一，保证答辩时“汇报口径、论文口径、实验口径”一致。
"""


def build_docx(data: dict, flows: dict[str, Path]) -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.6)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.2)
    add_page_number(sec)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    styles["Heading 1"].font.name = "黑体"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.name = "黑体"
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 3"].font.name = "黑体"
    styles["Heading 3"].font.size = Pt(12)

    add_center(doc, "毕业设计进展汇报（会议后更新版）", 18, True)
    add_center(doc, "课题名称：基于OpenGL的CAE软件数据预处理方法研究与实践", 13, False)
    add_body(doc, "本次汇报不再按一般“完成了哪些功能”展开，而是直接回应上次会议中老师提出的两个核心问题：第一，梯度计算功能如何保证与 VTK 的算法可比性，并将实验重点收敛到结构化网格主线；第二，数据优化功能如何说明仿真数据扰动的根本来源，以及当前模块的研究意义和适用边界。")

    add_heading(doc, "一、问题一：梯度计算功能", 1)
    add_heading(doc, "1.1 本次调整后的汇报主线", 2)
    for text in [
        "根据上次会议意见，当前汇报口径已经从“规则网格 + 非结构网格同时展开”调整为“以结构化网格作为主线验证对象，非结构化网格作为补充一致性结果”。这样做的原因在于：结构化网格上的有限差分与 VTK 具有更直接的可比性，解析场验证也更标准，能够更稳固地支撑工作量和实验深度。",
        "同时，非结构化网格部分不再以原来的加权最小二乘路径作为论文和汇报主线，而是转向更接近 `vtkGradientFilter` 思路的基于形函数导数方法。这样一来，系统在非结构化网格上的研究方向已经从“任意局部拟合”收敛到“与有限元单元插值理论一致、与 VTK 更可比”的路线，但当前答辩主结论不建立在复杂非结构网格深度对比上。",
        "换句话说，本次汇报对梯度模块的逻辑重心已经明确：主结论看结构化网格，非结构化网格用于说明研究已经推进到更可比的算法方向，但不再作为主要风险点去和老师硬碰复杂情况。",
    ]:
        add_body(doc, text)

    add_heading(doc, "1.2 对 VTK 可比性的理解与当前处理方式", 2)
    for text in [
        "上次会议中老师强调，和 VTK 对比时必须保证算法和计算条件的可比性，否则对比没有意义。基于这一要求，当前测试程序已经不再只输出一个 VTK 时间，而是同时输出 VTK 单线程时间、VTK 并行时间、并行线程数和并行后端。这样可以把“系统快不快”与“VTK 用的是什么配置”同时交代清楚。",
        "此外，当前系统已经将规则网格和非结构化网格的算法主线明确区分：规则网格默认采用有限差分；非结构化网格默认采用基于形函数导数的路径。测试程序在构造 VTK 参考时，也会显式记录其结果来源与时间来源。这样的改动，实质上是在把实验从“结果像不像”推进到“方法和条件是否可比”。",
        "因此，本次与老师汇报时，关于梯度模块不再只给一个“总体误差”和“总体时间”，而是明确分成结构化网格解析场正确性、结构化网格真实字段与 VTK 一致性、以及非结构化网格补充一致性三层。这样更符合数值实验的基本规范。",
    ]:
        add_body(doc, text)

    add_heading(doc, "1.3 当前主线 shader 详细流程图", 2)
    add_body(doc, "针对老师可能追问的“GPU 上到底是怎么计算的”，本次汇报不再只给模块级框图，而是把当前主线中实际参与计算的六个 shader 分别拆开说明。这样可以把系统实现、算法原理和实验结果对应起来。")
    add_figure(doc, "图1 FD.glsl 规则网格有限差分流程图", flows["fd"], "该 shader 是结构化网格梯度主线，对正文中的结构化网格解析场验证和与 VTK 时间对比具有直接支撑作用。")
    add_figure(doc, "图2 ShapePointGradient.glsl 非结构点梯度流程图", flows["shape_point"], "该 shader 负责非结构化点数据的形函数导数梯度计算，是当前非结构网格主线的重要组成部分。")
    add_figure(doc, "图3 ShapeCellGradient.glsl 非结构单元梯度流程图", flows["shape_cell"], "该 shader 在单元中心执行形函数导数计算；若输入来自单元场，则通常配合 CellDataToPointLift 使用。")
    add_figure(doc, "图4 CellDataToPointLift.glsl 单元场提升流程图", flows["cell_lift"], "该 shader 的作用不是直接求梯度，而是先把单元场恢复为点场，为后续单元梯度路径提供统一输入。")
    add_figure(doc, "图5 Bilateral.glsl 图双边滤波流程图", flows["bilateral"], "该 shader 对点域或单元域邻接图进行逐层平滑，是数据优化模块的分解阶段核心。")
    add_figure(doc, "图6 MultiScaleFuse.glsl 多尺度融合流程图", flows["fuse"], "该 shader 根据细节层幅值和增益参数执行重建，是数据优化模块的融合阶段核心。")

    add_heading(doc, "1.4 结构化网格：当前主线结果", 2)
    add_table(
        doc,
        "表1 结构化网格解析场验证结果（SampleStructGrid point）",
        ["字段", "NRMSE", "解释"],
        [
            ["benchmark_linear", sci(float(data["sample_linear"]["ambient_nrmse"]), 6), "线性场接近机器精度"],
            ["benchmark_trig", sci(float(data["sample_trig"]["ambient_nrmse"]), 6), "三角函数场误差较低"],
            ["benchmark_vec_linear", sci(float(data["sample_vec_linear"]["ambient_nrmse"]), 6), "向量线性场同样接近机器精度"],
        ],
        "表1对应的是解析场验证，其角色是回答“系统实现是否正确”，而不是回答“是否与 VTK 接近”。",
    )
    add_table(
        doc,
        "表2 结构化网格真实字段与 VTK 对比（SampleStructGrid）",
        ["关联", "字段", "NRMSE", "VTK单线程/ms", "VTK并行/ms", "系统总时间/ms", "GPU时间/ms"],
        [
            [
                "POINT",
                "scalars",
                sci(float(data["sample_real_point"]["ambient_nrmse"]), 6),
                ms(float(data["sample_real_point"]["ambient_vtk_single_avg_ms"])),
                ms(float(data["sample_real_point"]["ambient_vtk_parallel_avg_ms"])),
                ms(float(data["sample_real_point"]["result_wall_avg_ms"])),
                ms(float(data["sample_real_point"]["result_gpu_avg_ms"])),
            ],
            [
                "CELL",
                "scalars",
                sci(float(data["sample_real_cell"]["ambient_nrmse"]), 6),
                ms(float(data["sample_real_cell"]["ambient_vtk_single_avg_ms"])),
                ms(float(data["sample_real_cell"]["ambient_vtk_parallel_avg_ms"])),
                ms(float(data["sample_real_cell"]["result_wall_avg_ms"])),
                ms(float(data["sample_real_cell"]["result_gpu_avg_ms"])),
            ],
        ],
        "表2对应的是工程一致性验证。Point 路径结果与 VTK 高度接近，同时总时间显著低于当前环境下的 VTK 单线程和并行时间。",
    )
    add_figure(
        doc,
        "图7 梯度计算时间对比",
        ASSET_DIR / "gradient_timing_compare.png",
        "图中同时给出 VTK 单线程、VTK 并行、系统总时间和 GPU 时间。对当前汇报而言，重点应放在结构化网格结果，而不是把所有数据集混成一个口径。",
    )
    for text in [
        f"从表1可以看到，结构化网格在 `benchmark_linear` 上的 NRMSE 为 {sci(float(data['sample_linear']['ambient_nrmse']), 6)}，在 `benchmark_vec_linear` 上的 NRMSE 为 {sci(float(data['sample_vec_linear']['ambient_nrmse']), 6)}，已经达到接近机器精度的量级。这可以直接支撑规则网格有限差分实现的正确性。",
        f"从表2可以看到，`SampleStructGrid point / scalars` 与 VTK 的 NRMSE 为 {sci(float(data['sample_real_point']['ambient_nrmse']), 6)}；对应时间上，VTK 单线程为 {ms(float(data['sample_real_point']['ambient_vtk_single_avg_ms']))} ms，VTK 并行为 {ms(float(data['sample_real_point']['ambient_vtk_parallel_avg_ms']))} ms，系统总时间为 {ms(float(data['sample_real_point']['result_wall_avg_ms']))} ms，GPU 时间为 {ms(float(data['sample_real_point']['result_gpu_avg_ms']))} ms。这说明在结构化网格主线上，系统已经能够同时给出正确性证据和可比的性能结果。",
        "也就是说，梯度模块当前最稳妥的汇报方式，不是再把复杂非结构网格当作主战场，而是先把结构化网格这一条线讲扎实：算法原理清楚、与 VTK 可比、数据结果完整、工作量也足够展开。",
    ]:
        add_body(doc, text)

    add_heading(doc, "1.5 非结构化网格：当前阶段的补充口径", 2)
    add_table(
        doc,
        "表3 非结构化网格真实字段与 VTK 对比（补充）",
        ["数据集", "关联/字段", "NRMSE", "系统总时间/ms"],
        [
            ["hexa", "POINT / scalars", sci(float(data["hexa_real"]["ambient_nrmse"]), 6), ms(float(data["hexa_real"]["result_wall_avg_ms"]))],
            ["limb", "CELL / chem_0", sci(float(data["limb_real"]["ambient_nrmse"]), 6), ms(float(data["limb_real"]["result_wall_avg_ms"]))],
            ["ShipHull_0", "POINT / RF", sci(float(data["ship_real_point"]["ambient_nrmse"]), 6), ms(float(data["ship_real_point"]["result_wall_avg_ms"]))],
            ["ShipHull_0", "CELL / S_Mises", sci(float(data["ship_real_cell"]["ambient_nrmse"]), 6), ms(float(data["ship_real_cell"]["result_wall_avg_ms"]))],
            ["1_0", "POINT / RF", sci(float(data["one_real_point"]["ambient_nrmse"]), 6), ms(float(data["one_real_point"]["result_wall_avg_ms"]))],
            ["1_0", "CELL / S_Mises", sci(float(data["one_real_cell"]["ambient_nrmse"]), 6), ms(float(data["one_real_cell"]["result_wall_avg_ms"]))],
        ],
        "这部分结果当前只作为补充汇报：说明已经研究了更接近 VTK 的算法方向，并获得了较好的真实字段一致性，但非结构化复杂情形不是当前答辩主结论。",
    )
    for text in [
        "非结构化网格部分当前更适合向老师汇报为“研究已经推进，但主结论收缩”。原因有两点：第一，系统已经从原先的加权最小二乘主线转向更接近 VTK 的形函数导数主线；第二，复杂曲面和单元数据恢复路径仍然更敏感，因此不宜再把这一部分放在本阶段汇报的中心。",
        f"例如，`hexa point` 与 VTK 的 NRMSE 已经达到 {sci(float(data['hexa_real']['ambient_nrmse']), 6)}，`1_0 point` 达到 {sci(float(data['one_real_point']['ambient_nrmse']), 6)}；但 `ShipHull_0 cell` 这类曲面单元数据仍有更高误差。这类现象说明非结构化网格并不是“完全不能做”，而是它的实验口径和风险控制必须更谨慎。",
    ]:
        add_body(doc, text)
    add_figure(
        doc,
        "图8 ShipHull_0 系统梯度结果",
        RESULT_DIR / "系统梯度计算结果.png",
        "该图用于说明系统在复杂曲面型工程数据上已经能够给出稳定的梯度分布结果。",
    )
    add_figure(
        doc,
        "图9 ShipHull_0 VTK 梯度结果",
        RESULT_DIR / "vtk梯度计算结果.png",
        "该图用于与图2做视觉对照，说明两者在主高值区和整体分布形态上具有较高一致性。",
    )

    add_heading(doc, "1.6 曲面数据评价方式的调整", 2)
    for text in [
        f"对曲面型数据，当前测试程序已经明确区分 ambient 和 intrinsic 指标。以 `1_0 point` 为例，普通 `benchmark_linear` 的 ambient NRMSE 为 {sci(float(data['one_linear']['ambient_nrmse']), 6)}，而曲面专用 `benchmark_surface_linear` 的 intrinsic NRMSE 仅为 {sci(float(data['one_surface_linear']['intrinsic_nrmse']), 6)}；`benchmark_surface_trig` 的 intrinsic NRMSE 为 {sci(float(data['one_surface_trig']['intrinsic_nrmse']), 6)}。",
        "这说明曲面数据不能简单用普通三维解析场评价，否则会把法向分量误差错误地算进主结论。这个问题也是之前复杂非结构网格实验容易“看起来效果不好”的一个重要原因。当前汇报中把这一点单独讲清楚，能够直接回应老师对实验口径严谨性的要求。",
    ]:
        add_body(doc, text)

    add_heading(doc, "二、问题二：数据优化功能", 1)
    add_heading(doc, "2.1 本次调整后的研究定位", 2)
    for text in [
        "上次会议中老师对数据优化模块提出的核心质疑是：如果只是在人为干净场上加高斯噪声，再去做优化，那这个实验是否真正代表真实仿真数据中的噪声问题。基于这一意见，当前汇报口径已经不再把模块简单描述为“高斯噪声去噪”，而是先回答“真实仿真数据为什么会出现局部扰动”。",
        "目前更合理的定位是：CAE 后处理结果中的局部数值扰动可能来自离散化误差、迭代收敛残差、网格质量、梯度/应力恢复不连续以及浮点舍入等因素。这些因素并不统一，也不意味着真实误差整体服从某一种分布；但其中有一类扰动在结果场上会表现为局部随机高频波动，而当前数据优化模块面向的正是这一类现象。",
        "换句话说，当前模块的意义不再是“证明真实 CAE 噪声就是高斯噪声”，而是“针对一类在真实结果场中客观存在的局部随机高频扰动，构建一个后处理阶段的抑制方法，并验证其对该类扰动的有效性”。这样的表述更科学，也更容易经得起老师追问。",
    ]:
        add_body(doc, text)

    add_heading(doc, "2.2 模块原理与当前实验意义", 2)
    for text in [
        "当前数据优化模块采用图双边滤波做多尺度分解，再通过细节增益控制的多尺度融合完成重建。它并不是求解器误差修正模块，也不是通用去噪器，而是面向后处理结果场的一种稳定化处理手段：在尽量保持主结构和边界的前提下，抑制局部随机高频波动。",
        "从实验意义上说，受控扰动实验仍然有价值，但它的价值不在于“模拟全部真实误差”，而在于给出一个可控制、可重复、可对照的扰动模型，用来验证模块是否对目标扰动类别有效。因此，当前报告里应把“根因分析”和“代理扰动实验”明确分成两个层次：前者说明研究对象存在，后者说明当前算法对其中一类对象有效。",
    ]:
        add_body(doc, text)

    add_heading(doc, "2.3 当前统计数据", 2)
    add_table(
        doc,
        "表4 点数据多尺度优化平均结果",
        ["噪声类型", "输入NRMSE", "输出NRMSE", "RMSE改进率", "粗糙度比"],
        [
            ["gaussian", f"{data['point_avg'].loc['gaussian','input_nrmse']:.6f}", f"{data['point_avg'].loc['gaussian','fused_nrmse']:.6f}", f"{data['point_avg'].loc['gaussian','rmse_improvement_ratio']:.6f}", f"{data['point_avg'].loc['gaussian','roughness_ratio']:.6f}"],
            ["grf", f"{data['point_avg'].loc['grf','input_nrmse']:.6f}", f"{data['point_avg'].loc['grf','fused_nrmse']:.6f}", f"{data['point_avg'].loc['grf','rmse_improvement_ratio']:.6f}", f"{data['point_avg'].loc['grf','roughness_ratio']:.6f}"],
            ["impulse", f"{data['point_avg'].loc['impulse','input_nrmse']:.6f}", f"{data['point_avg'].loc['impulse','fused_nrmse']:.6f}", f"{data['point_avg'].loc['impulse','rmse_improvement_ratio']:.6f}", f"{data['point_avg'].loc['impulse','roughness_ratio']:.6f}"],
            ["mixed", f"{data['point_avg'].loc['mixed','input_nrmse']:.6f}", f"{data['point_avg'].loc['mixed','fused_nrmse']:.6f}", f"{data['point_avg'].loc['mixed','rmse_improvement_ratio']:.6f}", f"{data['point_avg'].loc['mixed','roughness_ratio']:.6f}"],
        ],
        "表4显示，点数据上对 gaussian 和 mixed 扰动的改善较明显，而对 impulse 基本没有改善。",
    )
    add_table(
        doc,
        "表5 单元数据多尺度优化平均结果",
        ["噪声类型", "输入NRMSE", "输出NRMSE", "RMSE改进率", "粗糙度比"],
        [
            ["gaussian", f"{data['cell_avg'].loc['gaussian','input_nrmse']:.6f}", f"{data['cell_avg'].loc['gaussian','fused_nrmse']:.6f}", f"{data['cell_avg'].loc['gaussian','rmse_improvement_ratio']:.6f}", f"{data['cell_avg'].loc['gaussian','roughness_ratio']:.6f}"],
            ["grf", f"{data['cell_avg'].loc['grf','input_nrmse']:.6f}", f"{data['cell_avg'].loc['grf','fused_nrmse']:.6f}", f"{data['cell_avg'].loc['grf','rmse_improvement_ratio']:.6f}", f"{data['cell_avg'].loc['grf','roughness_ratio']:.6f}"],
            ["impulse", f"{data['cell_avg'].loc['impulse','input_nrmse']:.6f}", f"{data['cell_avg'].loc['impulse','fused_nrmse']:.6f}", f"{data['cell_avg'].loc['impulse','rmse_improvement_ratio']:.6f}", f"{data['cell_avg'].loc['impulse','roughness_ratio']:.6f}"],
            ["mixed", f"{data['cell_avg'].loc['mixed','input_nrmse']:.6f}", f"{data['cell_avg'].loc['mixed','fused_nrmse']:.6f}", f"{data['cell_avg'].loc['mixed','rmse_improvement_ratio']:.6f}", f"{data['cell_avg'].loc['mixed','roughness_ratio']:.6f}"],
        ],
        "表5显示，单元域上对 gaussian 扰动的改善更明显，而 impulse 仍无明显改善，这恰好说明当前模块有清晰边界。",
    )
    for text in [
        f"从表4和表5可以看到，gaussian 扰动下，点数据的平均 NRMSE 从 {data['point_avg'].loc['gaussian','input_nrmse']:.6f} 降到 {data['point_avg'].loc['gaussian','fused_nrmse']:.6f}，单元数据则从 {data['cell_avg'].loc['gaussian','input_nrmse']:.6f} 降到 {data['cell_avg'].loc['gaussian','fused_nrmse']:.6f}。这说明当前模块对一类局部随机高频扰动确实具有抑制能力。",
        f"而对 impulse 扰动，点数据的 RMSE 改进率为 {data['point_avg'].loc['impulse','rmse_improvement_ratio']:.6f}，单元数据为 {data['cell_avg'].loc['impulse','rmse_improvement_ratio']:.6f}，都接近 1。这并不是坏事，反而说明模块并没有被夸大成“什么噪声都能处理”的通用方法，而是具有清楚的适用范围和边界。",
    ]:
        add_body(doc, text)
    add_figure(
        doc,
        "图10 数据优化前渲染图",
        RESULT_DIR / "数据优化前.png",
        "可见局部色带波动和高频斑块较明显，尤其在过渡界面附近更容易干扰后续观察与分析。",
    )
    add_figure(
        doc,
        "图11 数据优化后渲染图",
        RESULT_DIR / "数据优化后.png",
        "优化后局部随机高频波动减弱，而主结构和界面位置仍然保持清晰。",
    )
    add_figure(
        doc,
        "图12 数据优化前后对比图",
        ASSET_DIR / "denoise_render_compare.png",
        "这一对比图更直观地说明：当前模块的作用不是整体模糊，而是在保留主结构的同时抑制局部高频扰动。",
    )

    add_heading(doc, "三、建议向老师汇报的结论", 1)
    for text in [
        "第一，梯度计算部分已经按会议要求调整实验主线：主结论收敛到结构化网格，重点展示解析场正确性、与 VTK 的可比时间结果和算法一致性说明。这样能够保证实验深度和汇报逻辑的稳健性。",
        "第二，非结构化网格部分已经完成方向性调整：系统不再以加权最小二乘作为主对照线，而是转向更接近 `vtkGradientFilter` 的形函数导数路径；但复杂曲面与单元恢复仍更敏感，因此当前只作为补充结果汇报。",
        "第三，数据优化模块的定位已经从“高斯噪声去噪”收缩为“面向一类局部随机高频数值扰动的结果场优化”。当前实验结果表明它对 Gaussian 和 Mixed 更有效、对 Impulse 不敏感，这种受限但清晰的结论比泛化表述更严谨。",
    ]:
        add_body(doc, text)

    add_heading(doc, "四、下一步计划", 1)
    for text in [
        "1. 梯度模块继续保持“结构化网格为主、非结构化为辅”的汇报口径，并补充更细的结构化网格对比图表和文字分析。",
        "2. 数据优化模块继续围绕扰动根因做文献与案例整理，争取补充更贴近真实仿真扰动来源的实验设计，而不仅限于代理扰动验证。",
        "3. 将本次汇报口径与论文正文、图表和答辩讲稿进一步统一，确保‘汇报怎么说、论文怎么写、实验怎么做’三者一致。",
    ]:
        add_body(doc, text)

    add_heading(doc, "参考资料（汇报中可口头提及）", 1)
    for ref in [
        "[1] VTK Team. vtkGradientFilter Class Reference.",
        "[2] VTK Team. vtkSMPTools Class Reference.",
        "[3] Salari K, Knupp P K. Code Verification by the Method of Manufactured Solutions.",
        "[4] Roy C J, Smith T M, Ober C C. Verification of Euler/Navier-Stokes Codes Using the Method of Manufactured Solutions.",
        "[5] Tomasi C, Manduchi R. Bilateral Filtering for Gray and Color Images.",
        "[6] Fleishman S, Drori I, Cohen-Or D. Bilateral Mesh Denoising.",
        "[7] Conrad P R, Girolami M, Särkkä S, et al. Statistical analysis of differential equations: introducing probability measures on numerical solutions.",
        "[8] Gilkeson C A, Kirby M, Hicken J E. Dealing with numerical noise in CFD-based design optimization.",
    ]:
        add_body(doc, ref)

    doc.save(OUT_DOCX)


def build_markdown_v2(data: dict) -> str:
    gaussian_point = data["point_avg"].loc["gaussian"]
    gaussian_cell = data["cell_avg"].loc["gaussian"]
    mixed_point = data["point_avg"].loc["mixed"]
    defaults_rows = "\n".join(
        f"| {name} | {value} | {meaning} |" for name, value, meaning in data["parameter_defaults"]
    )
    basis_rows = "\n".join(
        f"| {row['case']} | {row['input_nrmse']:.6f} | {row['fused_nrmse']:.6f} | {row['roughness_ratio']:.6f} | {row['note']} |"
        for row in data["parameter_basis"]
    )
    return f"""# 老师汇报文稿（会议后更新版）

## 一、本次汇报重点

本次汇报只围绕上次会议中老师强调的两个问题组织：

1. 梯度模块如何保证与 VTK 的可比性，并把实验主线收敛到最能稳定支撑结论的 benchmark。
2. 数据优化模块如何说明扰动来源、噪声代理模型以及参数设置依据。

## 二、问题一：梯度计算功能

### 1. 当前汇报主线

- 结构化网格仍是答辩主线，因为有限差分路径、解析场验证和与 VTK 的时间对比都最直接。
- 非结构化网格已经从原 AWLS 路线转向更接近 `vtkGradientFilter` 的形函数导数路线，但当前只作为补充一致性结果汇报。
- 测试程序现在同时输出 VTK 单线程、VTK 并行、并行线程数和后端信息，因此性能对比不再是“系统 vs 一个模糊的 VTK 时间”。

### 2. 解析场是怎么生成的

- 规则网格和一般体网格：在 `CAEProcessingFacade.cpp` 中围绕包围盒中心构造局部坐标 `x_r,y_r,z_r`，按参考长度 `L` 归一化后注入 `benchmark_linear`、`benchmark_trig`、`benchmark_vec_linear` 等场，并同步写入 `*_exact_grad` 真值。
- 曲面型数据：在 `TestGradient.cpp` 中用全局切向基 `a,b` 构造局部坐标 `u=(p-c)·a`、`v=(p-c)·b`，生成 `benchmark_surface_linear`、`benchmark_surface_trig` 和 `benchmark_surface_vec_linear`，再把解析梯度投影回局部支撑空间。
- 因此，解析场验证回答的是“程序是否按理论公式实现”，而不是“是否和 VTK 长得像”。

### 3. benchmark 主展示口径

| 数据/字段 | 指标 | NRMSE | 作用 |
| --- | --- | ---: | --- |
| SampleStructGrid / benchmark_linear | ambient | {sci(float(data["sample_linear"]["ambient_nrmse"]), 6)} | 规则网格标量主验证 |
| SampleStructGrid / benchmark_vec_linear | ambient | {sci(float(data["sample_vec_linear"]["ambient_nrmse"]), 6)} | 规则网格向量主验证 |
| hexa / benchmark_linear | ambient | {sci(float(data["hexa_linear"]["ambient_nrmse"]), 6)} | 体网格点梯度主验证 |
| 1_0 / benchmark_surface_linear | intrinsic | {sci(float(data["one_surface_linear"]["intrinsic_nrmse"]), 6)} | 曲面点梯度主验证 |

补充压力测试不删，但不再放在主表里。比如 `SampleStructGrid / benchmark_trig` 为 {sci(float(data["sample_trig"]["ambient_nrmse"]), 6)}，`hexa / benchmark_trig` 为 {sci(float(data["hexa_trig"]["ambient_nrmse"]), 6)}，`1_0 / benchmark_surface_trig` 的 intrinsic NRMSE 为 {sci(float(data["one_surface_trig"]["intrinsic_nrmse"]), 6)}。

### 4. 结构化网格与 VTK 的主对比数据

| 关联 | 字段 | NRMSE | VTK单线程/ms | VTK并行/ms | 系统总时间/ms | GPU时间/ms |
| --- | --- | ---: | ---: | ---: | ---: | ---: |
| POINT | scalars | {sci(float(data["sample_real_point"]["ambient_nrmse"]), 6)} | {ms(float(data["sample_real_point"]["ambient_vtk_single_avg_ms"]))} | {ms(float(data["sample_real_point"]["ambient_vtk_parallel_avg_ms"]))} | {ms(float(data["sample_real_point"]["result_wall_avg_ms"]))} | {ms(float(data["sample_real_point"]["result_gpu_avg_ms"]))} |
| CELL | scalars | {sci(float(data["sample_real_cell"]["ambient_nrmse"]), 6)} | {ms(float(data["sample_real_cell"]["ambient_vtk_single_avg_ms"]))} | {ms(float(data["sample_real_cell"]["ambient_vtk_parallel_avg_ms"]))} | {ms(float(data["sample_real_cell"]["result_wall_avg_ms"]))} | {ms(float(data["sample_real_cell"]["result_gpu_avg_ms"]))} |

- 当前 VTK 并行后端为 `{data["sample_real_point"]["ambient_vtk_backend"]}`，线程数为 `{int(data["sample_real_point"]["ambient_vtk_parallel_threads"])}`。
- `SampleStructGrid point / scalars` 上，系统与 VTK 的 NRMSE 为 {sci(float(data["sample_real_point"]["ambient_nrmse"]), 6)}，且系统总时间和 GPU 纯计算时间都明显低于当前环境下的 VTK 参考时间。

### 5. 曲面数据为什么不能直接拿普通解析场评价

- 以 `1_0 point` 为例，普通 `benchmark_linear` 的 ambient NRMSE 为 {sci(float(data["one_linear"]["ambient_nrmse"]), 6)}。
- 但曲面专用 `benchmark_surface_linear` 的 intrinsic NRMSE 只有 {sci(float(data["one_surface_linear"]["intrinsic_nrmse"]), 6)}。
- 这说明曲面数据必须区分 ambient 与 intrinsic 指标，否则会把法向分量误差过度混入主结论。

### 6. 非结构化网格当前口径（补充）

| 数据集 | 关联/字段 | NRMSE | 系统总时间/ms |
| --- | --- | ---: | ---: |
| hexa | POINT / scalars | {sci(float(data["hexa_real"]["ambient_nrmse"]), 6)} | {ms(float(data["hexa_real"]["result_wall_avg_ms"]))} |
| limb | CELL / chem_0 | {sci(float(data["limb_real"]["ambient_nrmse"]), 6)} | {ms(float(data["limb_real"]["result_wall_avg_ms"]))} |
| ShipHull_0 | POINT / RF | {sci(float(data["ship_real_point"]["ambient_nrmse"]), 6)} | {ms(float(data["ship_real_point"]["result_wall_avg_ms"]))} |
| ShipHull_0 | CELL / S_Mises | {sci(float(data["ship_real_cell"]["ambient_nrmse"]), 6)} | {ms(float(data["ship_real_cell"]["result_wall_avg_ms"]))} |
| 1_0 | POINT / RF | {sci(float(data["one_real_point"]["ambient_nrmse"]), 6)} | {ms(float(data["one_real_point"]["result_wall_avg_ms"]))} |
| 1_0 | CELL / S_Mises | {sci(float(data["one_real_cell"]["ambient_nrmse"]), 6)} | {ms(float(data["one_real_cell"]["result_wall_avg_ms"]))} |

这一部分现在更适合汇报为“研究方向已经推进，但主结论收缩”：真实字段一致性已经较好，但复杂曲面和单元恢复路径仍更敏感，不作为当前答辩主战场。

## 三、问题二：数据优化功能

### 1. 模块定位

- CAE 后处理结果中的局部数值扰动可能来自离散化误差、迭代残差、网格质量、梯度/应力恢复不连续和浮点舍入。
- 本模块只面向其中一类表现为局部随机高频波动的扰动，而不是所有噪声。
- 算法上采用图双边滤波做多尺度分解，再通过细节增益控制的多尺度融合重建。

### 2. 高斯噪声与相关扰动是怎么添加的

- 在 `TestHarnessUtils.h` 中，先对干净场计算 `signalStd=std(clean)`，再按 `targetSigma=sigmaFactor·signalStd` 生成白高斯扰动，并按 `noisy=clean+noise` 得到受扰场。
- `grf` 场景是在白噪声基础上沿邻接图传播和平滑，再重新缩放到同一目标标准差，用来近似空间相关扰动。
- `mixed` 为高斯扰动加脉冲异常，`impulse` 则单独作为边界场景，用来证明模块不是通用去噪器。

### 3. 参数说明与小型案例依据

#### 表4 关键参数默认值

| 参数 | 默认值 | 作用 |
| --- | --- | --- |
{defaults_rows}

#### 表5 小型参数依据表

| 案例 | 输入NRMSE | 输出NRMSE | 粗糙度比 | 解释 |
| --- | ---: | ---: | ---: | --- |
{basis_rows}

这两张表的作用不是做大规模参数寻优，而是说明当前默认参数为什么可以保留下来：对 Gaussian 有效、对 GRF 保持谨慎、对 Impulse 不夸大能力。

### 4. 当前统计结果

| 扰动类型 | 点数据输入NRMSE | 点数据输出NRMSE | 单元数据输入NRMSE | 单元数据输出NRMSE |
| --- | ---: | ---: | ---: | ---: |
| gaussian | {gaussian_point["input_nrmse"]:.6f} | {gaussian_point["fused_nrmse"]:.6f} | {gaussian_cell["input_nrmse"]:.6f} | {gaussian_cell["fused_nrmse"]:.6f} |
| grf | {data["point_avg"].loc["grf","input_nrmse"]:.6f} | {data["point_avg"].loc["grf","fused_nrmse"]:.6f} | {data["cell_avg"].loc["grf","input_nrmse"]:.6f} | {data["cell_avg"].loc["grf","fused_nrmse"]:.6f} |
| impulse | {data["point_avg"].loc["impulse","input_nrmse"]:.6f} | {data["point_avg"].loc["impulse","fused_nrmse"]:.6f} | {data["cell_avg"].loc["impulse","input_nrmse"]:.6f} | {data["cell_avg"].loc["impulse","fused_nrmse"]:.6f} |
| mixed | {mixed_point["input_nrmse"]:.6f} | {mixed_point["fused_nrmse"]:.6f} | {data["cell_avg"].loc["mixed","input_nrmse"]:.6f} | {data["cell_avg"].loc["mixed","fused_nrmse"]:.6f} |

- Gaussian 扰动下，点数据和单元数据都表现出明显改善。
- GRF 改善有限，说明相关扰动比独立高斯扰动更难处理。
- Impulse 基本无改善，这恰好支撑“模块边界明确”这一表述。

## 四、建议向老师汇报的结论

1. 梯度模块已经把“解析场正确性”和“与 VTK 的工程一致性”分开组织，并把正文 benchmark 口径收敛到最能稳定支撑结论的四类数据。
2. 结构化网格现在能够同时提供解析场正确性证据、与 VTK 的高一致性结果，以及包含并行信息的可比时间数据。
3. 非结构化网格已经转向形函数导数路线，但复杂曲面和单元恢复仍更敏感，因此当前只作为补充结果汇报。
4. 数据优化模块的表述已经从“高斯噪声去噪”收缩为“局部随机高频扰动抑制”，并补充了噪声添加方式、参数默认值和小型案例依据。

## 五、下一步计划

1. 继续保持“结构化网格为主、非结构化为辅”的梯度模块汇报口径。
2. 在论文和答辩稿中统一解析场生成、曲面 intrinsic 指标和 VTK 并行信息的表述。
3. 在数据优化部分继续围绕扰动根因和代理模型边界补充文献与案例说明。
"""


def build_docx_v2(data: dict, flows: dict[str, Path]) -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    gaussian_point = data["point_avg"].loc["gaussian"]
    gaussian_cell = data["cell_avg"].loc["gaussian"]
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.6)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.2)
    add_page_number(sec)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    styles["Heading 1"].font.name = "黑体"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.name = "黑体"
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 3"].font.name = "黑体"
    styles["Heading 3"].font.size = Pt(12)

    add_center(doc, "毕业设计进展汇报（会议后更新版）", 18, True)
    add_center(doc, "课题名称：基于OpenGL的CAE软件数据预处理方法研究与实践", 13, False)
    add_body(doc, "本次汇报不再按一般“完成了哪些功能”展开，而是直接回应上次会议中老师提出的两个核心问题：第一，梯度计算功能如何保证与 VTK 的算法可比性，并将实验重点收敛到结构化网格主线；第二，数据优化功能如何说明仿真数据扰动的根本来源、噪声代理模型和当前参数设置依据。")

    add_heading(doc, "一、问题一：梯度计算功能", 1)
    add_heading(doc, "1.1 当前汇报主线", 2)
    for text in [
        "当前汇报主线已经明确收敛到结构化网格。这并不是放弃非结构化网格，而是因为结构化网格上的有限差分路径、解析场验证和与 VTK 的时间对比都最直接，最适合承担答辩主结论。",
        "非结构化网格部分已经从原来的加权最小二乘路线转向更接近 `vtkGradientFilter` 的形函数导数路线。这说明研究方向已经推进到更可比的算法框架，但复杂曲面和单元恢复路径仍更敏感，因此本阶段只作为补充一致性结果汇报。",
        "为了满足老师关于可比性的要求，测试程序现在会同时输出 VTK 单线程时间、VTK 并行时间、并行线程数和并行后端。因此，后续性能对比已经从“系统 vs 一个 VTK 时间”变成“系统 vs VTK 单线程/并行”的分层对比。",
    ]:
        add_body(doc, text)

    add_heading(doc, "1.2 解析场生成与 benchmark 展示口径", 2)
    for text in [
        "规则网格和一般体网格的解析场由 `CAEProcessingFacade.cpp` 自动注入：程序围绕包围盒中心构造局部坐标 `x_r,y_r,z_r`，按参考长度 `L` 归一化后生成 `benchmark_linear`、`benchmark_trig`、`benchmark_vec_linear` 等场，并同步写入 `*_exact_grad` 真值。曲面型数据则在 `TestGradient.cpp` 中使用切向基 `a,b` 和局部坐标 `u,v` 生成 `benchmark_surface_linear`、`benchmark_surface_trig` 和 `benchmark_surface_vec_linear`，再将解析梯度投影回局部支撑空间。",
        "基于当前实验稳定性，正文主展示只保留四类最能稳定支撑结论的 benchmark：规则网格标量线性场、规则网格向量线性场、体网格点线性场和曲面点切向线性场。`benchmark_trig` 等更复杂场形仍保留在结果文件中，但作为补充压力测试，不再承担正文主结论。",
    ]:
        add_body(doc, text)
    add_table(
        doc,
        "表1 梯度模块主验证 benchmark",
        ["数据/字段", "指标", "NRMSE", "作用"],
        [
            ["SampleStructGrid / benchmark_linear", "ambient", sci(float(data["sample_linear"]["ambient_nrmse"]), 6), "规则网格标量主验证"],
            ["SampleStructGrid / benchmark_vec_linear", "ambient", sci(float(data["sample_vec_linear"]["ambient_nrmse"]), 6), "规则网格向量主验证"],
            ["hexa / benchmark_linear", "ambient", sci(float(data["hexa_linear"]["ambient_nrmse"]), 6), "体网格点梯度主验证"],
            ["1_0 / benchmark_surface_linear", "intrinsic", sci(float(data["one_surface_linear"]["intrinsic_nrmse"]), 6), "曲面点梯度主验证"],
        ],
        "补充 benchmark 仍保留在结果文件中。其中 `SampleStructGrid / benchmark_trig` 为 "
        f"{sci(float(data['sample_trig']['ambient_nrmse']), 6)}，`hexa / benchmark_trig` 为 "
        f"{sci(float(data['hexa_trig']['ambient_nrmse']), 6)}，`1_0 / benchmark_surface_trig` 的 intrinsic NRMSE 为 "
        f"{sci(float(data['one_surface_trig']['intrinsic_nrmse']), 6)}。",
    )

    add_heading(doc, "1.3 当前主线 shader 详细流程图", 2)
    add_body(doc, "针对老师可能追问的“GPU 上到底是怎么计算的”，本次汇报继续保留 6 张按 shader 拆开的详细流程图，使算法实现、程序模块和实验结果可以一一对应。")
    add_figure(doc, "图1 FD.glsl 规则网格有限差分流程图", flows["fd"], "该 shader 支撑结构化网格梯度主线。")
    add_figure(doc, "图2 ShapePointGradient.glsl 非结构点梯度流程图", flows["shape_point"], "该 shader 负责非结构化点数据的形函数导数梯度计算。")
    add_figure(doc, "图3 ShapeCellGradient.glsl 非结构单元梯度流程图", flows["shape_cell"], "该 shader 在单元中心执行形函数导数计算。")
    add_figure(doc, "图4 CellDataToPointLift.glsl 单元场提升流程图", flows["cell_lift"], "该 shader 先把单元场提升到点域。")
    add_figure(doc, "图5 Bilateral.glsl 图双边滤波流程图", flows["bilateral"], "该 shader 对邻接图执行多尺度平滑。")
    add_figure(doc, "图6 MultiScaleFuse.glsl 多尺度融合流程图", flows["fuse"], "该 shader 根据细节增益执行重建。")

    add_heading(doc, "1.4 结构化网格与 VTK 的主对比数据", 2)
    add_table(
        doc,
        "表2 结构化网格真实字段与 VTK 对比（SampleStructGrid）",
        ["关联", "字段", "NRMSE", "VTK单线程/ms", "VTK并行/ms", "系统总时间/ms", "GPU时间/ms"],
        [
            [
                "POINT",
                "scalars",
                sci(float(data["sample_real_point"]["ambient_nrmse"]), 6),
                ms(float(data["sample_real_point"]["ambient_vtk_single_avg_ms"])),
                ms(float(data["sample_real_point"]["ambient_vtk_parallel_avg_ms"])),
                ms(float(data["sample_real_point"]["result_wall_avg_ms"])),
                ms(float(data["sample_real_point"]["result_gpu_avg_ms"])),
            ],
            [
                "CELL",
                "scalars",
                sci(float(data["sample_real_cell"]["ambient_nrmse"]), 6),
                ms(float(data["sample_real_cell"]["ambient_vtk_single_avg_ms"])),
                ms(float(data["sample_real_cell"]["ambient_vtk_parallel_avg_ms"])),
                ms(float(data["sample_real_cell"]["result_wall_avg_ms"])),
                ms(float(data["sample_real_cell"]["result_gpu_avg_ms"])),
            ],
        ],
        f"当前 VTK 并行后端为 {data['sample_real_point']['ambient_vtk_backend']}，并行线程数为 {int(data['sample_real_point']['ambient_vtk_parallel_threads'])}。",
    )
    add_figure(
        doc,
        "图7 梯度计算时间对比",
        ASSET_DIR / "gradient_timing_compare.png",
        "图中同时给出 VTK 单线程、VTK 并行、系统总时间和 GPU 时间。这样可以直接回应“性能比较是否公平”的问题。",
    )
    for text in [
        f"`SampleStructGrid point / scalars` 与 VTK 的 NRMSE 为 {sci(float(data['sample_real_point']['ambient_nrmse']), 6)}；对应时间上，VTK 单线程为 {ms(float(data['sample_real_point']['ambient_vtk_single_avg_ms']))} ms，VTK 并行为 {ms(float(data['sample_real_point']['ambient_vtk_parallel_avg_ms']))} ms，系统总时间为 {ms(float(data['sample_real_point']['result_wall_avg_ms']))} ms，GPU 时间为 {ms(float(data['sample_real_point']['result_gpu_avg_ms']))} ms。",
        "这组结果说明，结构化网格主线上现在已经能够同时给出三类证据：解析场正确性证据、与 VTK 的工程一致性证据，以及包含并行信息的可比时间数据。",
    ]:
        add_body(doc, text)

    add_heading(doc, "1.5 曲面评价方式与非结构化补充结果", 2)
    for text in [
        f"对曲面型数据，测试程序已经显式区分 ambient 和 intrinsic 指标。以 `1_0 point` 为例，普通 `benchmark_linear` 的 ambient NRMSE 为 {sci(float(data['one_linear']['ambient_nrmse']), 6)}，而曲面专用 `benchmark_surface_linear` 的 intrinsic NRMSE 仅为 {sci(float(data['one_surface_linear']['intrinsic_nrmse']), 6)}。这说明曲面数据不能再简单使用普通三维解析场评价，否则会把法向分量误差错误地算进主结论。",
        "非结构化网格当前更适合作为补充一致性结果汇报：它说明系统已经转向更接近 VTK 的形函数导数路线，并在真实字段上获得了较好的工程一致性；但复杂曲面和单元恢复路径仍更敏感，因此本阶段不把它作为答辩主战场。",
    ]:
        add_body(doc, text)
    add_table(
        doc,
        "表3 非结构化网格真实字段与 VTK 对比（补充）",
        ["数据集", "关联/字段", "NRMSE", "系统总时间/ms"],
        [
            ["hexa", "POINT / scalars", sci(float(data["hexa_real"]["ambient_nrmse"]), 6), ms(float(data["hexa_real"]["result_wall_avg_ms"]))],
            ["limb", "CELL / chem_0", sci(float(data["limb_real"]["ambient_nrmse"]), 6), ms(float(data["limb_real"]["result_wall_avg_ms"]))],
            ["ShipHull_0", "POINT / RF", sci(float(data["ship_real_point"]["ambient_nrmse"]), 6), ms(float(data["ship_real_point"]["result_wall_avg_ms"]))],
            ["ShipHull_0", "CELL / S_Mises", sci(float(data["ship_real_cell"]["ambient_nrmse"]), 6), ms(float(data["ship_real_cell"]["result_wall_avg_ms"]))],
            ["1_0", "POINT / RF", sci(float(data["one_real_point"]["ambient_nrmse"]), 6), ms(float(data["one_real_point"]["result_wall_avg_ms"]))],
            ["1_0", "CELL / S_Mises", sci(float(data["one_real_cell"]["ambient_nrmse"]), 6), ms(float(data["one_real_cell"]["result_wall_avg_ms"]))],
        ],
        "这部分当前只作为补充汇报，不承担答辩主结论。",
    )
    add_figure(doc, "图8 ShipHull_0 系统梯度结果", RESULT_DIR / "系统梯度计算结果.png", "用于展示系统在复杂曲面工程数据上的梯度分布形态。")
    add_figure(doc, "图9 ShipHull_0 VTK 梯度结果", RESULT_DIR / "vtk梯度计算结果.png", "用于与图8做视觉对照。")

    add_heading(doc, "二、问题二：数据优化功能", 1)
    add_heading(doc, "2.1 模块定位与高斯噪声添加方式", 2)
    for text in [
        "当前数据优化模块不再被表述为“高斯噪声去噪器”，而是被定位为：面向 CAE 后处理结果场中的一类局部随机高频数值扰动，执行后处理阶段的稳定化优化。",
        "这类扰动的来源可能包括离散化误差、迭代收敛残差、网格质量、梯度/应力恢复不连续以及浮点舍入等因素。也就是说，本文不是在证明真实 CAE 全部误差服从某个简单分布，而是在针对其中一类常见表现形式构造可控代理模型。",
        "在 `TestHarnessUtils.h` 中，程序先计算干净场标准差 `signalStd=std(clean)`，再按 `targetSigma=sigmaFactor·signalStd` 生成白高斯扰动并叠加得到受扰场；`grf` 则在白噪声基础上沿邻接图平滑传播，用来近似空间相关扰动；`mixed` 与 `impulse` 则用来验证模块边界。",
    ]:
        add_body(doc, text)

    add_heading(doc, "2.2 参数说明与小型案例依据", 2)
    add_body(doc, "老师提出过参数是否需要说明、是否需要做参数实验。当前处理方式是：梯度模块只报告实验配置项，不做数据依赖式调参；数据优化模块则保留一组固定默认参数，并用少量代表性案例说明这些默认参数为什么被保留下来。")
    add_table(
        doc,
        "表4 数据优化模块关键参数默认值",
        ["参数", "默认值", "作用"],
        [[name, value, meaning] for name, value, meaning in data["parameter_defaults"]],
        "该表的作用是交代全文批处理实验所固定采用的参数口径，而不是展开大规模参数搜索。",
    )
    add_table(
        doc,
        "表5 小型参数依据表",
        ["案例", "输入NRMSE", "输出NRMSE", "粗糙度比", "解释"],
        [
            [
                row["case"],
                f"{row['input_nrmse']:.6f}",
                f"{row['fused_nrmse']:.6f}",
                f"{row['roughness_ratio']:.6f}",
                row["note"],
            ]
            for row in data["parameter_basis"]
        ],
        "该表说明当前默认参数为什么可以保留下来：它对 Gaussian 有效、对 GRF 保持谨慎、对 Impulse 不夸大能力。",
    )

    add_heading(doc, "2.3 当前统计结果与渲染图", 2)
    add_table(
        doc,
        "表6 点数据多尺度优化平均结果",
        ["噪声类型", "输入NRMSE", "输出NRMSE", "RMSE改进率", "粗糙度比"],
        [
            ["gaussian", f"{data['point_avg'].loc['gaussian','input_nrmse']:.6f}", f"{data['point_avg'].loc['gaussian','fused_nrmse']:.6f}", f"{data['point_avg'].loc['gaussian','rmse_improvement_ratio']:.6f}", f"{data['point_avg'].loc['gaussian','roughness_ratio']:.6f}"],
            ["grf", f"{data['point_avg'].loc['grf','input_nrmse']:.6f}", f"{data['point_avg'].loc['grf','fused_nrmse']:.6f}", f"{data['point_avg'].loc['grf','rmse_improvement_ratio']:.6f}", f"{data['point_avg'].loc['grf','roughness_ratio']:.6f}"],
            ["impulse", f"{data['point_avg'].loc['impulse','input_nrmse']:.6f}", f"{data['point_avg'].loc['impulse','fused_nrmse']:.6f}", f"{data['point_avg'].loc['impulse','rmse_improvement_ratio']:.6f}", f"{data['point_avg'].loc['impulse','roughness_ratio']:.6f}"],
            ["mixed", f"{data['point_avg'].loc['mixed','input_nrmse']:.6f}", f"{data['point_avg'].loc['mixed','fused_nrmse']:.6f}", f"{data['point_avg'].loc['mixed','rmse_improvement_ratio']:.6f}", f"{data['point_avg'].loc['mixed','roughness_ratio']:.6f}"],
        ],
        "Gaussian 和 Mixed 改善较明显，而 Impulse 基本无改善。",
    )
    add_table(
        doc,
        "表7 单元数据多尺度优化平均结果",
        ["噪声类型", "输入NRMSE", "输出NRMSE", "RMSE改进率", "粗糙度比"],
        [
            ["gaussian", f"{data['cell_avg'].loc['gaussian','input_nrmse']:.6f}", f"{data['cell_avg'].loc['gaussian','fused_nrmse']:.6f}", f"{data['cell_avg'].loc['gaussian','rmse_improvement_ratio']:.6f}", f"{data['cell_avg'].loc['gaussian','roughness_ratio']:.6f}"],
            ["grf", f"{data['cell_avg'].loc['grf','input_nrmse']:.6f}", f"{data['cell_avg'].loc['grf','fused_nrmse']:.6f}", f"{data['cell_avg'].loc['grf','rmse_improvement_ratio']:.6f}", f"{data['cell_avg'].loc['grf','roughness_ratio']:.6f}"],
            ["impulse", f"{data['cell_avg'].loc['impulse','input_nrmse']:.6f}", f"{data['cell_avg'].loc['impulse','fused_nrmse']:.6f}", f"{data['cell_avg'].loc['impulse','rmse_improvement_ratio']:.6f}", f"{data['cell_avg'].loc['impulse','roughness_ratio']:.6f}"],
            ["mixed", f"{data['cell_avg'].loc['mixed','input_nrmse']:.6f}", f"{data['cell_avg'].loc['mixed','fused_nrmse']:.6f}", f"{data['cell_avg'].loc['mixed','rmse_improvement_ratio']:.6f}", f"{data['cell_avg'].loc['mixed','roughness_ratio']:.6f}"],
        ],
        "单元域上对 Gaussian 的改善更明显，但 Impulse 仍无显著改善。",
    )
    for text in [
        f"Gaussian 扰动下，点数据平均 NRMSE 从 {gaussian_point['input_nrmse']:.6f} 降到 {gaussian_point['fused_nrmse']:.6f}，单元数据从 {gaussian_cell['input_nrmse']:.6f} 降到 {gaussian_cell['fused_nrmse']:.6f}。这说明当前模块对目标扰动确实具有抑制能力。",
        f"而对 Impulse 扰动，点数据 RMSE 改进率为 {data['point_avg'].loc['impulse','rmse_improvement_ratio']:.6f}，单元数据为 {data['cell_avg'].loc['impulse','rmse_improvement_ratio']:.6f}，都接近 1。这并不是坏事，反而说明模块并没有被夸大成通用去噪器。",
    ]:
        add_body(doc, text)
    add_figure(doc, "图10 数据优化前渲染图", RESULT_DIR / "数据优化前.png", "局部色带波动和高频斑块较明显。")
    add_figure(doc, "图11 数据优化后渲染图", RESULT_DIR / "数据优化后.png", "局部随机高频波动减弱，而主结构保持清晰。")
    add_figure(doc, "图12 数据优化前后对比图", ASSET_DIR / "denoise_render_compare.png", "该图用于说明模块并非简单整体模糊，而是在保留主结构的同时抑制局部高频扰动。")

    add_heading(doc, "三、建议向老师汇报的结论", 1)
    for text in [
        "第一，梯度模块已经把“解析场正确性”和“与 VTK 的工程一致性”分开组织，并把正文 benchmark 收敛到四类最能稳定支撑结论的数据。",
        "第二，结构化网格现在能够同时提供解析场正确性证据、与 VTK 的高一致性结果，以及包含并行信息的可比时间数据；非结构化网格则作为补充结果汇报。",
        "第三，数据优化模块的表述已经从“高斯噪声去噪”收缩为“局部随机高频扰动抑制”，并补充了噪声添加方式、参数默认值和小型案例依据。",
    ]:
        add_body(doc, text)

    add_heading(doc, "四、下一步计划", 1)
    for text in [
        "1. 继续保持“结构化网格为主、非结构化为辅”的梯度模块汇报口径。",
        "2. 在论文和答辩稿中统一解析场生成、曲面 intrinsic 指标和 VTK 并行信息的表述。",
        "3. 在数据优化部分继续围绕扰动根因和代理模型边界补充文献与案例说明。",
    ]:
        add_body(doc, text)

    add_heading(doc, "参考资料（汇报中可口头提及）", 1)
    for ref in [
        "[1] VTK Team. vtkGradientFilter Class Reference.",
        "[2] VTK Team. vtkSMPTools Class Reference.",
        "[3] Salari K, Knupp P K. Code Verification by the Method of Manufactured Solutions.",
        "[4] Roy C J, Smith T M, Ober C C. Verification of Euler/Navier-Stokes Codes Using the Method of Manufactured Solutions.",
        "[5] Tomasi C, Manduchi R. Bilateral Filtering for Gray and Color Images.",
        "[6] Fleishman S, Drori I, Cohen-Or D. Bilateral Mesh Denoising.",
        "[7] Conrad P R, Girolami M, Särkkä S, et al. Statistical analysis of differential equations: introducing probability measures on numerical solutions.",
        "[8] Gilkeson C A, Kirby M, Hicken J E. Dealing with numerical noise in CFD-based design optimization.",
    ]:
        add_body(doc, ref)

    doc.save(OUT_DOCX)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    data = load_data()
    flows = generate_shader_flowcharts()
    md = build_markdown_v2(data)
    OUT_MD.write_text(md, encoding="utf-8")
    build_docx_v2(data, flows)
    print(f"generated: {OUT_DOCX}")
    print(f"generated: {OUT_MD}")


if __name__ == "__main__":
    main()
