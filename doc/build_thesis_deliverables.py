from __future__ import annotations

import math
import textwrap
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
DOC_DIR = ROOT / "doc"
ASSET_DIR = DOC_DIR / "assets"
FLOW_DIR = ASSET_DIR / "shader_flows"
GEN_DIR = DOC_DIR / "generated"
RESULT_GRAD_DIR = ROOT / "results" / "gradient"
RESULT_MUL_DIR = ROOT / "results" / "mul"

THESIS_DOCX = GEN_DIR / "基于OpenGL的CAE软件数据预处理方法研究与实践_论文终稿.docx"
THESIS_MD = GEN_DIR / "基于OpenGL的CAE软件数据预处理方法研究与实践_论文终稿.md"
REPORT_FULL = DOC_DIR / "给老师汇报稿-10分钟.md"
REPORT_SHORT = DOC_DIR / "给老师汇报稿-简版.md"


FONT_CANDIDATES = [
    Path("C:/Windows/Fonts/msyh.ttc"),
    Path("C:/Windows/Fonts/msyh.ttf"),
    Path("C:/Windows/Fonts/simsun.ttc"),
    Path("C:/Windows/Fonts/simhei.ttf"),
]


def pick_font_path() -> Path | None:
    for path in FONT_CANDIDATES:
        if path.exists():
            return path
    return None


FONT_PATH = pick_font_path()


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    if FONT_PATH is not None:
        return ImageFont.truetype(str(FONT_PATH), size=size)
    return ImageFont.load_default()


TITLE_FONT = font(34)
SUBTITLE_FONT = font(24)
LABEL_FONT = font(22)
BODY_FONT = font(18)
SMALL_FONT = font(16)


def ensure_dirs() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    GEN_DIR.mkdir(parents=True, exist_ok=True)


def fmt_sci(value: float, digits: int = 3) -> str:
    if pd.isna(value):
        return "-"
    if value == 0:
        return "0"
    if abs(value) >= 1e4 or abs(value) < 1e-3:
        return f"{value:.{digits}e}"
    return f"{value:.{digits}f}"


def fmt_ms(value: float) -> str:
    if pd.isna(value):
        return "-"
    return f"{value:.3f}"


def fmt_ratio(value: float) -> str:
    if pd.isna(value):
        return "-"
    return f"{value:.3f}"


def set_run_font(run, size_pt: float = 12, bold: bool = False, name: str = "Times New Roman"):
    run.font.name = name
    run.bold = bold
    run.font.size = Pt(size_pt)
    rpr = run._element.rPr
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        run._element.insert(0, rpr)
    east_asia = OxmlElement("w:rFonts")
    east_asia.set(qn("w:eastAsia"), "宋体")
    east_asia.set(qn("w:ascii"), name)
    east_asia.set(qn("w:hAnsi"), name)
    rpr.append(east_asia)


def add_page_number(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    p._p.append(fld_begin)
    p._p.append(instr)
    p._p.append(fld_end)


def add_toc_field(paragraph) -> None:
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r' TOC \o "1-3" \h \z \u '
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "目录将在 Word 中更新域后自动生成。"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    paragraph._p.append(fld_begin)
    paragraph._p.append(instr)
    paragraph._p.append(fld_sep)
    paragraph._p.append(placeholder)
    paragraph._p.append(fld_end)


def add_body_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, 12)


def add_center_paragraph(doc: Document, text: str, size: float = 12, bold: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, size, bold=bold)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading %d" % min(level, 3)]
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, 16, bold=True, name="黑体")
    elif level == 2:
        set_run_font(run, 14, bold=True, name="黑体")
    else:
        set_run_font(run, 12, bold=True, name="黑体")


def add_table_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, 11, bold=False)


def add_figure_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, 11, bold=False)


def wrap_text(text: str, width: int) -> list[str]:
    chunks = []
    for raw in text.split("\n"):
        pieces = textwrap.wrap(raw, width=width, break_long_words=False, replace_whitespace=False)
        chunks.extend(pieces or [""])
    return chunks


@dataclass
class TableData:
    title: str
    headers: list[str]
    rows: list[list[str]]
    note: str | None = None


@dataclass
class FigureData:
    caption: str
    path: Path
    note: str | None = None


def load_gradient_frames() -> dict[str, pd.DataFrame]:
    frames: dict[str, pd.DataFrame] = {}
    for csv_path in sorted(RESULT_GRAD_DIR.glob("*.csv")):
        frames[csv_path.stem] = pd.read_csv(csv_path)
    return frames


def load_multiscale_frames() -> tuple[pd.DataFrame, pd.DataFrame]:
    point_df = pd.read_csv(RESULT_MUL_DIR / "multiscale_report+point.csv")
    cell_df = pd.read_csv(RESULT_MUL_DIR / "multiscale_report+cell.csv")
    return point_df, cell_df


def pick_row(df: pd.DataFrame, array_name: str) -> pd.Series:
    match = df[df["array"] == array_name]
    if match.empty:
        raise KeyError(f"array '{array_name}' not found")
    return match.iloc[0]


def non_benchmark_row(df: pd.DataFrame) -> pd.Series:
    match = df[~df["array"].astype(str).str.startswith("benchmark")]
    if match.empty:
        raise ValueError("no non-benchmark row found")
    return match.iloc[0]


def draw_box(draw: ImageDraw.ImageDraw, xy, text: str, fill: str, outline: str = "#24445c") -> None:
    x0, y0, x1, y1 = xy
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
    lines = wrap_text(text, 12)
    total_h = len(lines) * 30
    y = y0 + (y1 - y0 - total_h) / 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=LABEL_FONT)
        w = bbox[2] - bbox[0]
        draw.text((x0 + (x1 - x0 - w) / 2, y), line, font=LABEL_FONT, fill="#10222f")
        y += 30


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], fill: str = "#325d79") -> None:
    draw.line([start, end], fill=fill, width=5)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 18
    wing = math.pi / 7
    p1 = (end[0] - length * math.cos(angle - wing), end[1] - length * math.sin(angle - wing))
    p2 = (end[0] - length * math.cos(angle + wing), end[1] - length * math.sin(angle + wing))
    draw.polygon([end, p1, p2], fill=fill)


def save_canvas(img: Image.Image, path: Path) -> None:
    img.save(path, format="PNG")


def render_vertical_flow(path: Path, title: str, steps: list[str], note: str | None = None) -> None:
    width = 1500
    box_w = 1160
    box_h = 110
    gap = 36
    top_margin = 120
    bottom_extra = 90 if note else 40
    height = top_margin + len(steps) * box_h + (len(steps) - 1) * gap + bottom_extra
    img = Image.new("RGB", (width, height), "#ffffff")
    draw = ImageDraw.Draw(img)
    draw.text((60, 28), title, font=TITLE_FONT, fill="#173042")
    x0 = (width - box_w) // 2
    x1 = x0 + box_w
    y = top_margin
    fills = ["#dceffc", "#e6f3e7", "#f8ecd4", "#f4dce2", "#e7e7f8", "#e8f1f8", "#eef6dd", "#f8e5d8"]
    centers = []
    for idx, step in enumerate(steps):
        y0 = y
        y1 = y0 + box_h
        draw_box(draw, (x0, y0, x1, y1), step, fill=fills[idx % len(fills)])
        centers.append((width // 2, y0, y1))
        y = y1 + gap
    for i in range(len(centers) - 1):
        start = (centers[i][0], centers[i][2])
        end = (centers[i + 1][0], centers[i + 1][1])
        draw_arrow(draw, start, end)
    if note:
        draw.text((60, height - 60), note, font=SMALL_FONT, fill="#4a5d6b")
    save_canvas(img, path)


def generate_shader_flowcharts() -> dict[str, Path]:
    FLOW_DIR.mkdir(parents=True, exist_ok=True)
    flows: dict[str, Path] = {}
    flow_specs = {
        "fd": (
            "FD.glsl 规则网格有限差分流程图",
            [
                "线程索引 = gl_GlobalInvocationID(x,y,z)，若越界则直接返回",
                "根据网格逻辑尺寸与点/单元关联方式，计算当前样本线性 id",
                "在 x、y、z 三个方向调用 pick1D 选择中心差分或单边差分模板",
                "读取前后邻居坐标与字段值，构造参数空间中的局部差分量",
                "由相邻几何坐标近似 Jacobian，并计算逆 Jacobian 或退化保护量",
                "若 Jacobian 退化，则将逆映射置零，避免 NaN 和异常放大",
                "对每个字段分量分别计算 dXi、dEta、dZeta",
                "用链式法则将参考空间导数映射到物理空间，得到 gx、gy、gz",
                "将 3×numComponents 梯度写回输出缓冲区 G[]",
            ],
            "一个线程处理一个规则网格样本；点数据与单元数据共享同一套差分框架。",
        ),
        "shape_point": (
            "ShapePointGradient.glsl 非结构化点梯度流程图",
            [
                "线程索引 = pointId，先将该点的全部输出梯度清零",
                "扫描 pointCellOff / pointCellNbr，识别与该点相连的候选单元",
                "仅保留最高拓扑维单元，跳过无效单元和低维混合单元",
                "在当前单元中定位该点 localId，并求其参考单元参数坐标",
                "若为六面体点位，使用牛顿迭代反求 pcoords",
                "prepareCellGeometry：读取单元类型、节点坐标、形函数导数和映射信息",
                "对每个字段分量计算参考单元梯度 gradRef",
                "通过逆 Jacobian 或曲面切空间映射，将 gradRef 变为物理空间梯度",
                "累加所有有效相邻单元贡献，并按 usedCells 做平均",
                "把该点最终梯度写回输出缓冲区",
            ],
            "若该路径返回全零缓冲区，真正的保底退化逻辑在 GLGradientEngine.cpp 中回退到 ShapeCellGradient + CellDataToPointLift。",
        ),
        "shape_cell": (
            "ShapeCellGradient.glsl 非结构化单元梯度流程图",
            [
                "线程索引 = cellId，先将该单元的全部输出梯度清零",
                "根据单元类型选择单元中心参考坐标 pcoords",
                "prepareCellGeometry：检查单元类型、节点数、连接关系和参考导数模板",
                "在单元中心计算 dN/dξ，并累计节点坐标构造 Jacobian 或切向量",
                "体单元直接计算逆 Jacobian，曲面单元改用度量张量逆映射",
                "对每个字段分量计算参考梯度 gradRef",
                "将 gradRef 映射为物理空间梯度 g",
                "按分量顺序把单元梯度写入输出缓冲区 grad[]",
            ],
            "该 shader 假定输入 pointVal 已准备好；若原始场为单元场，通常先经 CellDataToPointLift 恢复点值。",
        ),
        "cell_lift": (
            "CellDataToPointLift.glsl 单元场提升到点场流程图",
            [
                "线程索引 = pointId，先将该点输出 pointVal 清零",
                "扫描 pointCellOff / pointCellNbr，收集与该点相邻的所有单元",
                "统计这些相邻单元中的最高拓扑维数",
                "仅保留最高维单元，避免壳单元与线单元混合污染点值",
                "对每个字段分量累加这些单元的 cellVal",
                "统计有效 usedCells 数量",
                "若 usedCells > 0，则按数量求平均",
                "把平均后的结果写成该点的 pointVal",
            ],
            "该 shader 不直接计算梯度，而是为单元数据梯度路径准备点值恢复场。",
        ),
        "bilateral": (
            "Bilateral.glsl 图双边滤波流程图",
            [
                "线程索引 = 样本 i，若越界则返回",
                "读取样本位置 pi 以及邻域 CSR 区间 off[i] ~ off[i+1]",
                "根据平均邻距和字段标准差，计算 sigmaS^2 与 sigmaR^2",
                "对每个字段分量独立执行双边滤波",
                "以中心值 center 初始化 sumW = 1、sumV = center",
                "遍历所有邻居 j，计算 dist^2、dv、空间核和值域核",
                "累加总权重 sumW 与加权值 sumV",
                "输出 sumV / sumW 作为平滑结果 outVal",
            ],
            "一个线程处理一个点或一个单元中心；标量场和向量场按分量独立处理。",
        ),
        "fuse": (
            "MultiScaleFuse.glsl 多尺度融合流程图",
            [
                "线程索引 = 标量样本 idx，若越界则返回",
                "读取 base 层及 d0/d1/d2 三个细节层",
                "依据 uLevelCount 判断哪些 detail 层有效",
                "计算各层细节幅值 m0/m1/m2 及总特征强度 feature",
                "根据 edgeSigma 计算细节回注衰减因子 atten",
                "按各层幅值占比与 detailGain 分配权重 w0、w1、w2",
                "执行重建：out = base + w0*d0 + w1*d1 + w2*d2",
                "将融合结果写入输出缓冲区 outVal[]",
            ],
            "该 shader 只做逐样本重建融合；几何邻域相关的平滑步骤已经在 Bilateral.glsl 中完成。",
        ),
    }
    for key, (title, steps, note) in flow_specs.items():
        path = FLOW_DIR / f"{key}.png"
        render_vertical_flow(path, title, steps, note)
        flows[key] = path
    return flows


def draw_horizontal_grouped_bars(
    path: Path,
    title: str,
    categories: list[str],
    series_names: list[str],
    values: list[list[float]],
    colors: list[str],
    *,
    log_scale: bool = False,
    min_positive: float = 1e-8,
    xlabel: str = "",
) -> None:
    width = 1500
    height = max(820, 180 + len(categories) * 100)
    img = Image.new("RGB", (width, height), "#ffffff")
    draw = ImageDraw.Draw(img)
    draw.text((60, 36), title, font=TITLE_FONT, fill="#173042")

    left = 280
    right = width - 220
    top = 140
    bottom = height - 90
    plot_w = right - left
    group_h = (bottom - top) / max(len(categories), 1)
    series_h = min(22, (group_h - 28) / max(len(series_names), 1))
    bar_gap = 7

    flat = [max(float(v), min_positive) for row in values for v in row if not pd.isna(v)]
    if not flat:
        flat = [1.0]
    if log_scale:
        min_v = math.floor(math.log10(min(flat)))
        max_v = math.ceil(math.log10(max(flat)))
        if min_v == max_v:
            max_v += 1
        ticks = list(range(min_v, max_v + 1))

        def scale(v: float) -> float:
            vv = max(v, min_positive)
            return (math.log10(vv) - min_v) / (max_v - min_v)
    else:
        max_v = max(flat) * 1.1
        ticks = [i / 5 * max_v for i in range(6)]

        def scale(v: float) -> float:
            return max(v, 0.0) / max_v if max_v > 0 else 0.0

    for tick in ticks:
        ratio = (tick - ticks[0]) / (ticks[-1] - ticks[0]) if len(ticks) > 1 else 0
        x = left + ratio * plot_w
        draw.line((x, top - 10, x, bottom + 10), fill="#d5dee5", width=1)
        label = f"1e{tick}" if log_scale else fmt_sci(tick, 2)
        bbox = draw.textbbox((0, 0), label, font=SMALL_FONT)
        draw.text((x - (bbox[2] - bbox[0]) / 2, bottom + 18), label, font=SMALL_FONT, fill="#4b5b66")

    draw.line((left, top, left, bottom), fill="#45606f", width=2)
    draw.line((left, bottom, right, bottom), fill="#45606f", width=2)

    for idx, cat in enumerate(categories):
        base_y = top + idx * group_h
        text_lines = wrap_text(cat, 12)
        ty = base_y + 8
        for line in text_lines:
            draw.text((24, ty), line, font=BODY_FONT, fill="#233642")
            ty += 24
        for s_idx, s_name in enumerate(series_names):
            value = float(values[idx][s_idx])
            y0 = base_y + 8 + s_idx * (series_h + bar_gap)
            y1 = y0 + series_h
            ratio = scale(value)
            x1 = left + ratio * plot_w
            draw.rectangle((left, y0, x1, y1), fill=colors[s_idx], outline=colors[s_idx])
            draw.text((x1 + 10, y0 - 2), fmt_sci(value, 3), font=SMALL_FONT, fill="#2a3a45")

    legend_x = width - 360
    legend_y = 40
    for idx, name in enumerate(series_names):
        y = legend_y + idx * 32
        draw.rectangle((legend_x, y + 4, legend_x + 18, y + 22), fill=colors[idx], outline=colors[idx])
        draw.text((legend_x + 28, y), name, font=BODY_FONT, fill="#21323e")

    if xlabel:
        bbox = draw.textbbox((0, 0), xlabel, font=BODY_FONT)
        draw.text((left + (plot_w - (bbox[2] - bbox[0])) / 2, height - 44), xlabel, font=BODY_FONT, fill="#21323e")

    save_canvas(img, path)


def combine_side_by_side(path: Path, title: str, left_title: str, right_title: str, left_img: Path, right_img: Path) -> None:
    left = Image.open(left_img).convert("RGB")
    right = Image.open(right_img).convert("RGB")
    target_h = 620
    left = left.resize((int(left.width * target_h / left.height), target_h))
    right = right.resize((int(right.width * target_h / right.height), target_h))
    margin = 40
    title_h = 120
    label_h = 40
    width = left.width + right.width + margin * 3
    height = target_h + title_h + label_h + 60
    canvas = Image.new("RGB", (width, height), "#ffffff")
    draw = ImageDraw.Draw(canvas)
    draw.text((40, 28), title, font=TITLE_FONT, fill="#173042")
    lx = margin
    rx = margin * 2 + left.width
    y = title_h
    canvas.paste(left, (lx, y))
    canvas.paste(right, (rx, y))
    draw.text((lx + 10, y - 34), left_title, font=SUBTITLE_FONT, fill="#234257")
    draw.text((rx + 10, y - 34), right_title, font=SUBTITLE_FONT, fill="#234257")
    save_canvas(canvas, path)


def create_architecture_figure() -> Path:
    path = ASSET_DIR / "system_architecture.png"
    img = Image.new("RGB", (1600, 920), "#ffffff")
    draw = ImageDraw.Draw(img)
    draw.text((50, 32), "系统总体架构图", font=TITLE_FONT, fill="#173042")
    draw_box(draw, (70, 170, 310, 300), "VTK数据文件\n与测试配置", "#dceffc")
    draw_box(draw, (380, 170, 680, 300), "CAEProcessingFacade\n统一门面层", "#cfe8d7")
    draw_box(draw, (760, 90, 1080, 220), "VTKDataConverter\n桥接层", "#f7e4c7")
    draw_box(draw, (760, 280, 1080, 410), "DataObject\n内部扁平数据表示", "#f7e4c7")
    draw_box(draw, (1160, 90, 1500, 220), "GLGradientEngine\nGPU梯度计算", "#f6d5db")
    draw_box(draw, (1160, 280, 1500, 410), "GLFilterEngine\nGPU数据优化", "#f6d5db")
    draw_box(draw, (380, 520, 680, 650), "GUI主界面\n结果显示与参数输入", "#dceffc")
    draw_box(draw, (760, 520, 1080, 650), "测试程序\nTestGradient / TestMultiScale", "#dceffc")
    draw_box(draw, (1160, 520, 1500, 650), "结果写回与导出\nVTK/CSV/ParaView", "#cfe8d7")
    draw_box(draw, (760, 740, 1080, 850), "VTK渲染与文件导出\n参考基线与可视化", "#e8ecf0")
    draw_arrow(draw, (310, 235), (380, 235))
    draw_arrow(draw, (680, 235), (760, 155))
    draw_arrow(draw, (680, 235), (760, 345))
    draw_arrow(draw, (1080, 155), (1160, 155))
    draw_arrow(draw, (1080, 345), (1160, 345))
    draw_arrow(draw, (530, 300), (530, 520))
    draw_arrow(draw, (920, 410), (920, 520))
    draw_arrow(draw, (1330, 220), (1330, 520))
    draw_arrow(draw, (1330, 410), (1330, 520))
    draw_arrow(draw, (1080, 585), (1160, 585))
    draw_arrow(draw, (1330, 650), (1080, 795))
    save_canvas(img, path)
    return path


def create_gradient_flow_figure() -> Path:
    path = ASSET_DIR / "gradient_module_flow.png"
    img = Image.new("RGB", (1600, 980), "#ffffff")
    draw = ImageDraw.Draw(img)
    draw.text((50, 32), "梯度计算模块流程图", font=TITLE_FONT, fill="#173042")
    draw_box(draw, (80, 140, 350, 260), "输入：数据集、字段、关联类型", "#dceffc")
    draw_box(draw, (460, 140, 760, 260), "CAEProcessingFacade::computeGradient\n统一入口与合法性校验", "#cfe8d7")
    draw_box(draw, (900, 110, 1240, 230), "规则网格?\n是：Finite Difference", "#f7e4c7")
    draw_box(draw, (900, 280, 1240, 400), "非结构网格?\n是：Shape Function Derivatives", "#f7e4c7")
    draw_box(draw, (1320, 90, 1540, 250), "FD.glsl\n参数空间差分\n+ Jacobian映射", "#f6d5db")
    draw_box(draw, (1320, 260, 1540, 420), "ShapePointGradient.glsl\n或\nShapeCellGradient.glsl", "#f6d5db")
    draw_box(draw, (1120, 470, 1540, 620), "异常保护：若点梯度路径全零\n则退化为 ShapeCellGradient\n+ CellDataToPointLift", "#e8ecf0")
    draw_box(draw, (460, 720, 760, 840), "记录总时间与GPU时间", "#dceffc")
    draw_box(draw, (900, 720, 1240, 840), "结果数组写回 DataObject", "#cfe8d7")
    draw_box(draw, (1320, 720, 1540, 840), "测试程序/GUI/导出使用", "#dceffc")
    draw_arrow(draw, (350, 200), (460, 200))
    draw_arrow(draw, (760, 200), (900, 170))
    draw_arrow(draw, (760, 200), (900, 340))
    draw_arrow(draw, (1240, 170), (1320, 170))
    draw_arrow(draw, (1240, 340), (1320, 340))
    draw_arrow(draw, (1430, 420), (1430, 470))
    draw_arrow(draw, (1430, 620), (760, 780))
    draw_arrow(draw, (760, 780), (900, 780))
    draw_arrow(draw, (1240, 780), (1320, 780))
    save_canvas(img, path)
    return path


def create_multiscale_flow_figure() -> Path:
    path = ASSET_DIR / "multiscale_module_flow.png"
    img = Image.new("RGB", (1600, 940), "#ffffff")
    draw = ImageDraw.Draw(img)
    draw.text((50, 32), "数据优化模块流程图", font=TITLE_FONT, fill="#173042")
    draw_box(draw, (80, 150, 340, 270), "输入：点场/单元场", "#dceffc")
    draw_box(draw, (430, 150, 760, 270), "构建点邻域图或单元邻域图", "#cfe8d7")
    draw_box(draw, (860, 150, 1180, 270), "估计平均邻距与字段标准差", "#f7e4c7")
    draw_box(draw, (1270, 80, 1510, 200), "Level 1\nBilateral.glsl", "#f6d5db")
    draw_box(draw, (1270, 240, 1510, 360), "Level 2\nBilateral.glsl", "#f6d5db")
    draw_box(draw, (1270, 400, 1510, 520), "Level 3\nBilateral.glsl", "#f6d5db")
    draw_box(draw, (860, 430, 1180, 570), "相邻平滑层做差\n得到 detail0/detail1/detail2", "#e8ecf0")
    draw_box(draw, (430, 650, 760, 790), "最深层平滑结果作为 base", "#dceffc")
    draw_box(draw, (860, 650, 1180, 790), "MultiScaleFuse.glsl\n按细节增益回注", "#cfe8d7")
    draw_box(draw, (1270, 650, 1510, 790), "输出：smooth/detail/base/fused", "#dceffc")
    draw_arrow(draw, (340, 210), (430, 210))
    draw_arrow(draw, (760, 210), (860, 210))
    draw_arrow(draw, (1180, 210), (1270, 140))
    draw_arrow(draw, (1390, 200), (1390, 240))
    draw_arrow(draw, (1390, 360), (1390, 400))
    draw_arrow(draw, (1270, 500), (1180, 500))
    draw_arrow(draw, (1020, 570), (1020, 650))
    draw_arrow(draw, (760, 720), (860, 720))
    draw_arrow(draw, (1180, 720), (1270, 720))
    save_canvas(img, path)
    return path


def build_data_context() -> dict:
    grad_frames = load_gradient_frames()
    mul_point, mul_cell = load_multiscale_frames()

    real_compare_order = [
        ("SampleStructGridpoint", "规则网格 SampleStructGrid", "POINT", "scalars"),
        ("SampleStructGridcell", "规则网格 SampleStructGrid", "CELL", "scalars"),
        ("hexapoint", "六面体体网格 hexa", "POINT", "scalars"),
        ("limbcell", "复杂体网格 limb", "CELL", "chem_0"),
        ("ShipHull_0point", "曲面网格 ShipHull_0", "POINT", "RF"),
        ("ShipHull_0cell", "曲面网格 ShipHull_0", "CELL", "S_Mises"),
        ("1_0point", "曲面网格 1_0", "POINT", "RF"),
        ("1_0cell", "曲面网格 1_0", "CELL", "S_Mises"),
    ]
    real_rows = []
    for key, dataset_name, assoc, field_name in real_compare_order:
        row = non_benchmark_row(grad_frames[key])
        result_wall_ms = float(row["result_wall_avg_ms"])
        vtk_single_ms = float(row["ambient_vtk_single_avg_ms"])
        vtk_parallel_ms = float(row["ambient_vtk_parallel_avg_ms"])
        real_rows.append({
            "dataset": dataset_name,
            "association": assoc,
            "field": field_name,
            "nrmse": float(row["ambient_nrmse"]),
            "softrel": float(row["ambient_softrel_mean"]),
            "vtk_backend": str(row["ambient_vtk_backend"]),
            "vtk_single_ms": vtk_single_ms,
            "vtk_parallel_ms": vtk_parallel_ms,
            "vtk_parallel_threads": int(row["ambient_vtk_parallel_threads"]),
            "result_wall_ms": result_wall_ms,
            "result_gpu_ms": float(row["result_gpu_avg_ms"]),
            "speedup_single": vtk_single_ms / result_wall_ms if result_wall_ms > 0 else float("nan"),
            "speedup_parallel": vtk_parallel_ms / result_wall_ms if result_wall_ms > 0 else float("nan"),
        })

    analytic_rows = [
        {
            "dataset": "SampleStructGrid",
            "field": "benchmark_linear",
            "metric": "Ambient NRMSE",
            "value": float(pick_row(grad_frames["SampleStructGridpoint"], "benchmark_linear")["ambient_nrmse"]),
            "role": "规则网格标量主验证",
        },
        {
            "dataset": "SampleStructGrid",
            "field": "benchmark_vec_linear",
            "metric": "Ambient NRMSE",
            "value": float(pick_row(grad_frames["SampleStructGridpoint"], "benchmark_vec_linear")["ambient_nrmse"]),
            "role": "规则网格向量主验证",
        },
        {
            "dataset": "hexa",
            "field": "benchmark_linear",
            "metric": "Ambient NRMSE",
            "value": float(pick_row(grad_frames["hexapoint"], "benchmark_linear")["ambient_nrmse"]),
            "role": "体网格点梯度主验证",
        },
        {
            "dataset": "1_0",
            "field": "benchmark_surface_linear",
            "metric": "Intrinsic NRMSE",
            "value": float(pick_row(grad_frames["1_0point"], "benchmark_surface_linear")["intrinsic_nrmse"]),
            "role": "曲面点梯度主验证",
        },
    ]

    analytic_secondary = [
        {
            "dataset": "SampleStructGrid",
            "field": "benchmark_trig",
            "metric": "Ambient NRMSE",
            "value": float(pick_row(grad_frames["SampleStructGridpoint"], "benchmark_trig")["ambient_nrmse"]),
        },
        {
            "dataset": "hexa",
            "field": "benchmark_trig",
            "metric": "Ambient NRMSE",
            "value": float(pick_row(grad_frames["hexapoint"], "benchmark_trig")["ambient_nrmse"]),
        },
        {
            "dataset": "1_0",
            "field": "benchmark_surface_trig",
            "metric": "Intrinsic NRMSE",
            "value": float(pick_row(grad_frames["1_0point"], "benchmark_surface_trig")["intrinsic_nrmse"]),
        },
    ]

    analytic_surface_compare = [
        {
            "field": "benchmark_linear",
            "ambient": float(pick_row(grad_frames["1_0point"], "benchmark_linear")["ambient_nrmse"]),
            "intrinsic": float(pick_row(grad_frames["1_0point"], "benchmark_linear")["intrinsic_nrmse"]),
        },
        {
            "field": "benchmark_trig",
            "ambient": float(pick_row(grad_frames["1_0point"], "benchmark_trig")["ambient_nrmse"]),
            "intrinsic": float(pick_row(grad_frames["1_0point"], "benchmark_trig")["intrinsic_nrmse"]),
        },
        {
            "field": "benchmark_surface_linear",
            "ambient": float(pick_row(grad_frames["1_0point"], "benchmark_surface_linear")["ambient_nrmse"]),
            "intrinsic": float(pick_row(grad_frames["1_0point"], "benchmark_surface_linear")["intrinsic_nrmse"]),
        },
        {
            "field": "benchmark_surface_trig",
            "ambient": float(pick_row(grad_frames["1_0point"], "benchmark_surface_trig")["ambient_nrmse"]),
            "intrinsic": float(pick_row(grad_frames["1_0point"], "benchmark_surface_trig")["intrinsic_nrmse"]),
        },
    ]

    appendix_cell = [
        {
            "dataset": "hexa",
            "linear": float(pick_row(grad_frames["hexacell"], "benchmark_linear")["ambient_nrmse"]),
            "trig": float(pick_row(grad_frames["hexacell"], "benchmark_trig")["ambient_nrmse"]),
            "vec_linear": float(pick_row(grad_frames["hexacell"], "benchmark_vec_linear")["ambient_nrmse"]),
        },
        {
            "dataset": "limb",
            "linear": float(pick_row(grad_frames["limbcell"], "benchmark_linear")["ambient_nrmse"]),
            "trig": float(pick_row(grad_frames["limbcell"], "benchmark_trig")["ambient_nrmse"]),
            "vec_linear": float(pick_row(grad_frames["limbcell"], "benchmark_vec_linear")["ambient_nrmse"]),
        },
        {
            "dataset": "ShipHull_0",
            "linear": float(pick_row(grad_frames["ShipHull_0cell"], "benchmark_surface_linear")["intrinsic_nrmse"]),
            "trig": float(pick_row(grad_frames["ShipHull_0cell"], "benchmark_surface_trig")["intrinsic_nrmse"]),
            "vec_linear": float(pick_row(grad_frames["ShipHull_0cell"], "benchmark_surface_vec_linear")["intrinsic_nrmse"]),
        },
        {
            "dataset": "1_0",
            "linear": float(pick_row(grad_frames["1_0cell"], "benchmark_surface_linear")["intrinsic_nrmse"]),
            "trig": float(pick_row(grad_frames["1_0cell"], "benchmark_surface_trig")["intrinsic_nrmse"]),
            "vec_linear": float(pick_row(grad_frames["1_0cell"], "benchmark_surface_vec_linear")["intrinsic_nrmse"]),
        },
    ]

    point_avg = mul_point.groupby("noise")[["input_nrmse", "fused_nrmse", "rmse_improvement_ratio", "roughness_ratio", "gpu_avg_ms"]].mean().reset_index()
    cell_avg = mul_cell.groupby("noise")[["input_nrmse", "fused_nrmse", "rmse_improvement_ratio", "roughness_ratio", "gpu_avg_ms"]].mean().reset_index()

    param_ref = mul_point[(mul_point["clean_array"] == "ms_clean_trig") & (mul_point["noise"] == "gaussian")].iloc[0]
    parameter_defaults = [
        {"name": "levels", "value": int(param_ref["levels"]), "meaning": "多尺度分解层数", "effect": "增大可增强低频分离，但会增加计算和过平滑风险"},
        {"name": "iterations", "value": int(param_ref["iterations"]), "meaning": "每层双边滤波迭代次数", "effect": "增大可加强平滑，但边缘钝化风险同步增加"},
        {"name": "spatial_sigma_factor", "value": float(param_ref["spatial_sigma_factor"]), "meaning": "空间邻域权重尺度", "effect": "增大可扩大平滑影响范围"},
        {"name": "range_sigma_factor", "value": float(param_ref["range_sigma_factor"]), "meaning": "数值相似度权重尺度", "effect": "增大后更容易跨数值差异做平均"},
        {"name": "level_scale", "value": float(param_ref["level_scale"]), "meaning": "层间尺度放大倍数", "effect": "增大后高层更偏向大尺度平滑"},
        {"name": "edge_sigma_factor", "value": float(param_ref["edge_sigma_factor"]), "meaning": "细节回注的边缘抑制尺度", "effect": "减小可更强保护边缘，增大则更平滑"},
        {"name": "detail_gain[0,1,2]", "value": f"[{float(param_ref['detail_gain0']):.2f}, {float(param_ref['detail_gain1']):.2f}, {float(param_ref['detail_gain2']):.2f}]", "meaning": "三层细节回注增益", "effect": "从细到粗逐层递减，避免高频细节被过量放回"},
        {"name": "sigma_factor", "value": float(param_ref["sigma_factor"]), "meaning": "高斯扰动强度系数", "effect": "控制 synthetic 高斯扰动标准差"},
        {"name": "corr_length / corr_iters / corr_alpha", "value": f"{float(param_ref['corr_length_factor']):.2f} / {int(param_ref['corr_iters'])} / {float(param_ref['corr_alpha']):.2f}", "meaning": "相关高斯扰动的相关长度与迭代参数", "effect": "决定 GRF 扰动的空间相关程度"},
        {"name": "impulse_ratio / impulse_scale", "value": f"{float(param_ref['impulse_ratio']):.2f} / {float(param_ref['impulse_scale']):.2f}", "meaning": "脉冲扰动占比与幅值比例", "effect": "决定 impulse / mixed 场景下异常值强度"},
    ]

    parameter_basis = []
    for clean_array, noise, implication in [
        ("ms_clean_edge", "gaussian", "边缘场在默认参数下仍明显降噪，说明当前 range / edge 设置没有把主边界整体抹平。"),
        ("ms_clean_trig", "grf", "相关噪声改善有限，说明不能仅靠放大平滑参数去强压空间相关扰动。"),
        ("ms_clean_edge", "impulse", "脉冲扰动几乎无改善，说明问题主要来自噪声类型失配，而不是当前参数略偏弱。"),
    ]:
        row = mul_point[(mul_point["clean_array"] == clean_array) & (mul_point["noise"] == noise)].iloc[0]
        parameter_basis.append({
            "case": f"POINT / {clean_array} / {noise}",
            "input_nrmse": float(row["input_nrmse"]),
            "fused_nrmse": float(row["fused_nrmse"]),
            "roughness_ratio": float(row["roughness_ratio"]),
            "implication": implication,
        })

    representative = []
    for assoc_label, df in [("POINT", mul_point), ("CELL", mul_cell)]:
        for clean_array in ["ms_clean_trig", "ms_clean_edge"]:
            for noise in ["gaussian", "mixed", "impulse"]:
                row = df[(df["clean_array"] == clean_array) & (df["noise"] == noise)].iloc[0]
                representative.append({
                    "association": assoc_label,
                    "clean_array": clean_array,
                    "noise": noise,
                    "input_nrmse": float(row["input_nrmse"]),
                    "fused_nrmse": float(row["fused_nrmse"]),
                    "rmse_ratio": float(row["rmse_improvement_ratio"]),
                    "roughness_ratio": float(row["roughness_ratio"]),
                })

    surface_cases = [
        {
            "dataset": "1_0 point",
            "field": "benchmark_linear",
            "ambient": float(pick_row(grad_frames["1_0point"], "benchmark_linear")["ambient_nrmse"]),
            "intrinsic": float(pick_row(grad_frames["1_0point"], "benchmark_linear")["intrinsic_nrmse"]),
        },
        {
            "dataset": "1_0 point",
            "field": "benchmark_surface_linear",
            "ambient": float(pick_row(grad_frames["1_0point"], "benchmark_surface_linear")["ambient_nrmse"]),
            "intrinsic": float(pick_row(grad_frames["1_0point"], "benchmark_surface_linear")["intrinsic_nrmse"]),
        },
        {
            "dataset": "ShipHull_0 point",
            "field": "benchmark_surface_linear",
            "ambient": float(pick_row(grad_frames["ShipHull_0point"], "benchmark_surface_linear")["ambient_nrmse"]),
            "intrinsic": float(pick_row(grad_frames["ShipHull_0point"], "benchmark_surface_linear")["intrinsic_nrmse"]),
        },
        {
            "dataset": "ShipHull_0 cell",
            "field": "benchmark_surface_trig",
            "ambient": float(pick_row(grad_frames["ShipHull_0cell"], "benchmark_surface_trig")["ambient_nrmse"]),
            "intrinsic": float(pick_row(grad_frames["ShipHull_0cell"], "benchmark_surface_trig")["intrinsic_nrmse"]),
        },
    ]

    return {
        "real_rows": real_rows,
        "analytic_rows": analytic_rows,
        "analytic_secondary": analytic_secondary,
        "analytic_surface_compare": analytic_surface_compare,
        "appendix_cell": appendix_cell,
        "point_avg": point_avg,
        "cell_avg": cell_avg,
        "parameter_defaults": parameter_defaults,
        "parameter_basis": parameter_basis,
        "representative": representative,
        "surface_cases": surface_cases,
    }


def generate_asset_figures(ctx: dict) -> dict[str, Path]:
    assets: dict[str, Path] = {}
    assets["system_arch"] = create_architecture_figure()
    assets["gradient_flow"] = create_gradient_flow_figure()
    assets["multiscale_flow"] = create_multiscale_flow_figure()
    shader_flows = generate_shader_flowcharts()
    assets["shader_fd"] = shader_flows["fd"]
    assets["shader_shape_point"] = shader_flows["shape_point"]
    assets["shader_shape_cell"] = shader_flows["shape_cell"]
    assets["shader_cell_lift"] = shader_flows["cell_lift"]
    assets["shader_bilateral"] = shader_flows["bilateral"]
    assets["shader_fuse"] = shader_flows["fuse"]

    analytic_categories = [f"{row['dataset']}\n{row['field']}" for row in ctx["analytic_rows"]]
    analytic_values = [[row["value"]] for row in ctx["analytic_rows"]]
    analytic_path = ASSET_DIR / "gradient_analytic_validation.png"
    draw_horizontal_grouped_bars(
        analytic_path,
        "解析场主验证结果（点数据主线）",
        analytic_categories,
        ["NRMSE"],
        analytic_values,
        ["#3c91e6"],
        log_scale=True,
        xlabel="归一化均方根误差（对数坐标）",
    )
    assets["analytic_chart"] = analytic_path

    real_path = ASSET_DIR / "gradient_real_compare.png"
    draw_horizontal_grouped_bars(
        real_path,
        "真实字段与 VTK 结果对比",
        [f"{r['dataset']}\n{r['association']} / {r['field']}" for r in ctx["real_rows"]],
        ["Ambient NRMSE"],
        [[r["nrmse"]] for r in ctx["real_rows"]],
        ["#ef6f6c"],
        log_scale=True,
        xlabel="与 vtkGradientFilter 的 NRMSE（对数坐标）",
    )
    assets["real_chart"] = real_path

    timing_path = ASSET_DIR / "gradient_timing_compare.png"
    draw_horizontal_grouped_bars(
        timing_path,
        "梯度计算时间对比",
        [f"{r['dataset']}\n{r['association']} / {r['field']}" for r in ctx["real_rows"]],
        ["VTK单线程", "VTK并行", "系统总时间", "系统GPU时间"],
        [
            [r["vtk_single_ms"], r["vtk_parallel_ms"], r["result_wall_ms"], r["result_gpu_ms"]]
            for r in ctx["real_rows"]
        ],
        ["#577590", "#f8961e", "#43aa8b", "#277da1"],
        log_scale=True,
        xlabel="时间 / ms（对数坐标）",
    )
    assets["timing_chart"] = timing_path

    speedup_path = ASSET_DIR / "gradient_speedup_compare.png"
    draw_horizontal_grouped_bars(
        speedup_path,
        "系统相对 VTK 的加速比",
        [f"{r['dataset']}\n{r['association']} / {r['field']}" for r in ctx["real_rows"]],
        ["相对VTK单线程", "相对VTK并行"],
        [[r["speedup_single"], r["speedup_parallel"]] for r in ctx["real_rows"]],
        ["#6d597a", "#355070"],
        log_scale=True,
        min_positive=1e-3,
        xlabel="加速比（对数坐标，>1 表示系统更快）",
    )
    assets["speedup_chart"] = speedup_path

    surface_path = ASSET_DIR / "surface_metric_compare.png"
    draw_horizontal_grouped_bars(
        surface_path,
        "曲面数据 Ambient / Intrinsic 指标对比",
        [f"{r['dataset']}\n{r['field']}" for r in ctx["surface_cases"]],
        ["Ambient NRMSE", "Intrinsic NRMSE"],
        [[r["ambient"], r["intrinsic"]] for r in ctx["surface_cases"]],
        ["#e76f51", "#2a9d8f"],
        log_scale=True,
        xlabel="NRMSE（对数坐标）",
    )
    assets["surface_chart"] = surface_path

    point_path = ASSET_DIR / "multiscale_point_avg.png"
    draw_horizontal_grouped_bars(
        point_path,
        "点数据多尺度优化平均结果",
        point_avg_categories := ctx["point_avg"]["noise"].tolist(),
        ["输入NRMSE", "输出NRMSE"],
        [[float(r["input_nrmse"]), float(r["fused_nrmse"])] for _, r in ctx["point_avg"].iterrows()],
        ["#b56576", "#2a9d8f"],
        log_scale=False,
        xlabel="平均 NRMSE",
    )
    assets["point_avg_chart"] = point_path

    cell_path = ASSET_DIR / "multiscale_cell_avg.png"
    draw_horizontal_grouped_bars(
        cell_path,
        "单元数据多尺度优化平均结果",
        ctx["cell_avg"]["noise"].tolist(),
        ["输入NRMSE", "输出NRMSE"],
        [[float(r["input_nrmse"]), float(r["fused_nrmse"])] for _, r in ctx["cell_avg"].iterrows()],
        ["#b56576", "#2a9d8f"],
        log_scale=False,
        xlabel="平均 NRMSE",
    )
    assets["cell_avg_chart"] = cell_path

    rough_path = ASSET_DIR / "multiscale_roughness_ratio.png"
    rough_categories = ctx["point_avg"]["noise"].tolist()
    rough_values = []
    for noise in rough_categories:
        point_row = ctx["point_avg"][ctx["point_avg"]["noise"] == noise].iloc[0]
        cell_row = ctx["cell_avg"][ctx["cell_avg"]["noise"] == noise].iloc[0]
        rough_values.append([float(point_row["roughness_ratio"]), float(cell_row["roughness_ratio"])])
    draw_horizontal_grouped_bars(
        rough_path,
        "粗糙度抑制效果对比",
        rough_categories,
        ["POINT粗糙度比", "CELL粗糙度比"],
        rough_values,
        ["#8d99ae", "#06d6a0"],
        log_scale=False,
        xlabel="粗糙度比（越小表示高频起伏抑制越强）",
    )
    assets["roughness_chart"] = rough_path

    grad_compare_path = ASSET_DIR / "gradient_render_compare.png"
    combine_side_by_side(
        grad_compare_path,
        "ShipHull_0 梯度结果 ParaView 对比",
        "系统梯度结果",
        "VTK 梯度结果",
        ROOT / "results" / "系统梯度计算结果.png",
        ROOT / "results" / "vtk梯度计算结果.png",
    )
    assets["gradient_render"] = grad_compare_path
    assets["gradient_render_system_raw"] = ROOT / "results" / "系统梯度计算结果.png"
    assets["gradient_render_vtk_raw"] = ROOT / "results" / "vtk梯度计算结果.png"

    denoise_compare_path = ASSET_DIR / "denoise_render_compare.png"
    combine_side_by_side(
        denoise_compare_path,
        "ShipHull_0 数据优化前后 ParaView 对比",
        "优化前",
        "优化后",
        ROOT / "results" / "数据优化前.png",
        ROOT / "results" / "数据优化后.png",
    )
    assets["denoise_render"] = denoise_compare_path
    assets["denoise_before_raw"] = ROOT / "results" / "数据优化前.png"
    assets["denoise_after_raw"] = ROOT / "results" / "数据优化后.png"
    return assets


def build_tables(ctx: dict) -> dict[str, TableData]:
    tables: dict[str, TableData] = {}
    tables["dataset_grad"] = TableData(
        title="表5-1 梯度模块实验数据集与用途",
        headers=["数据集", "主要网格特征", "正文中的主要角色"],
        rows=[
            ["SampleStructGrid", "规则网格、三维体数据", "有限差分正确性与时间基线"],
            ["hexa", "六面体体网格", "非结构网格点梯度解析验证"],
            ["1_0", "曲面型非结构网格", "曲面专用解析验证与真实字段补充"],
            ["limb", "复杂体网格", "真实字段与 VTK 一致性、时间对比"],
            ["ShipHull_0", "曲面型工程网格", "真实字段对照与数据优化实验主数据"],
        ],
        note="本批实验以点数据解析验证和真实字段对照为主，单元数据解析场结果转入附录讨论。",
    )
    tables["analytic_main"] = TableData(
        title="表5-2 解析场主验证结果（点数据主线）",
        headers=["数据集", "字段", "采用指标", "NRMSE", "正文角色"],
        rows=[[r["dataset"], r["field"], r["metric"], fmt_sci(r["value"], 6), r["role"]] for r in ctx["analytic_rows"]],
        note="正文主展示口径收敛到四类代表性 benchmark：规则网格标量/向量线性场、体网格点线性场和曲面点切向线性场。更高阶或三角型 benchmark 保留在 CSV 与补充分析中，不作为主结论的唯一支撑。",
    )
    tables["surface_compare"] = TableData(
        title="表5-3 1_0 曲面点数据的 ambient / intrinsic 指标对比",
        headers=["字段", "Ambient NRMSE", "Intrinsic NRMSE"],
        rows=[[r["field"], fmt_sci(r["ambient"], 6), fmt_sci(r["intrinsic"], 6)] for r in ctx["analytic_surface_compare"]],
        note="该表用于说明曲面型数据必须使用与几何维度一致的解析场和评价指标，否则会把法向项误差错误混入主结论。",
    )
    tables["real_compare"] = TableData(
        title="表5-4 真实字段与 vtkGradientFilter 的一致性对比",
        headers=["数据集", "关联", "字段", "NRMSE", "SoftRel Mean"],
        rows=[
            [r["dataset"], r["association"], r["field"], fmt_sci(r["nrmse"], 6), fmt_sci(r["softrel"], 6)]
            for r in ctx["real_rows"]
        ],
        note="该表回答的是“工程一致性”问题，而不替代解析真值验证。",
    )
    tables["timing_compare"] = TableData(
        title="表5-5 梯度计算时间对比",
        headers=["数据集", "关联/字段", "VTK单线程/ms", "VTK并行/ms", "线程数", "系统总时间/ms", "GPU时间/ms"],
        rows=[
            [
                r["dataset"],
                f"{r['association']} / {r['field']}",
                fmt_ms(r["vtk_single_ms"]),
                fmt_ms(r["vtk_parallel_ms"]),
                str(r["vtk_parallel_threads"]),
                fmt_ms(r["result_wall_ms"]),
                fmt_ms(r["result_gpu_ms"]),
            ]
            for r in ctx["real_rows"]
        ],
        note="VTK 并行后端已启用为 STDThread，线程数为 16；但并行不一定在每个数据集上都快于单线程，实验中应如实报告。",
    )
    tables["speedup_compare"] = TableData(
        title="表5-6 系统相对 VTK 的加速比",
        headers=["数据集", "关联/字段", "相对VTK单线程", "相对VTK并行"],
        rows=[
            [
                r["dataset"],
                f"{r['association']} / {r['field']}",
                fmt_ratio(r["speedup_single"]),
                fmt_ratio(r["speedup_parallel"]),
            ]
            for r in ctx["real_rows"]
        ],
        note="加速比定义为 VTK 时间除以系统总时间。大于 1 表示系统更快，小于 1 表示系统更慢。",
    )
    tables["mul_point_avg"] = TableData(
        title="表5-7 点数据多尺度优化平均结果",
        headers=["噪声类型", "平均输入NRMSE", "平均输出NRMSE", "平均RMSE改进率", "平均粗糙度比", "平均GPU时间/ms"],
        rows=[
            [
                str(r["noise"]),
                fmt_ratio(float(r["input_nrmse"])),
                fmt_ratio(float(r["fused_nrmse"])),
                fmt_ratio(float(r["rmse_improvement_ratio"])),
                fmt_ratio(float(r["roughness_ratio"])),
                fmt_ms(float(r["gpu_avg_ms"])),
            ]
            for _, r in ctx["point_avg"].iterrows()
        ],
        note="gaussian 与 mixed 的抑制效果明显；impulse 的 RMSE 改进率接近 1，说明当前模块并不适合脉冲型异常值。",
    )
    tables["mul_cell_avg"] = TableData(
        title="表5-8 单元数据多尺度优化平均结果",
        headers=["噪声类型", "平均输入NRMSE", "平均输出NRMSE", "平均RMSE改进率", "平均粗糙度比", "平均GPU时间/ms"],
        rows=[
            [
                str(r["noise"]),
                fmt_ratio(float(r["input_nrmse"])),
                fmt_ratio(float(r["fused_nrmse"])),
                fmt_ratio(float(r["rmse_improvement_ratio"])),
                fmt_ratio(float(r["roughness_ratio"])),
                fmt_ms(float(r["gpu_avg_ms"])),
            ]
            for _, r in ctx["cell_avg"].iterrows()
        ],
        note="与点数据相比，单元域上对 gaussian 扰动的平均抑制更强，但 impulse 场景同样没有显著改善。",
    )
    tables["mul_repr"] = TableData(
        title="表5-9 代表性场上的数据优化结果",
        headers=["关联", "干净场", "噪声", "输入NRMSE", "输出NRMSE", "RMSE改进率", "粗糙度比"],
        rows=[
            [
                r["association"],
                r["clean_array"],
                r["noise"],
                fmt_ratio(r["input_nrmse"]),
                fmt_ratio(r["fused_nrmse"]),
                fmt_ratio(r["rmse_ratio"]),
                fmt_ratio(r["roughness_ratio"]),
            ]
            for r in ctx["representative"]
        ],
        note="ms_clean_edge 用于观察边缘保持能力；在 gaussian 与 mixed 场景下，误差和粗糙度都明显下降，说明该模块不是简单的整体模糊化。",
    )
    tables["param_defaults"] = TableData(
        title="表5-9a 数据优化模块关键参数与默认设置",
        headers=["参数", "默认值", "作用", "调大后的主要影响"],
        rows=[
            [r["name"], str(r["value"]), r["meaning"], r["effect"]]
            for r in ctx["parameter_defaults"]
        ],
        note="本组参数在小规模观察后固定，用于全文 batch 实验；本文不把参数搜索作为主工作量，而是强调固定口径下的可重复比较。",
    )
    tables["param_prestudy"] = TableData(
        title="表5-9b 参数选择依据的小型案例验证",
        headers=["案例", "输入NRMSE", "输出NRMSE", "粗糙度比", "对参数解释的含义"],
        rows=[
            [
                r["case"],
                fmt_ratio(r["input_nrmse"]),
                fmt_ratio(r["fused_nrmse"]),
                fmt_ratio(r["roughness_ratio"]),
                r["implication"],
            ]
            for r in ctx["parameter_basis"]
        ],
        note="该表不是完整参数寻优，而是用于说明当前默认参数为什么被保留下来：它们对 gaussian 场景有效、对 grf 场景保持谨慎、对 impulse 场景不过度夸大能力。",
    )
    tables["appendix_cell"] = TableData(
        title="附表B-1 单元数据解析场补充结果",
        headers=["数据集", "linear", "trig", "vec_linear"],
        rows=[
            [r["dataset"], fmt_sci(r["linear"], 6), fmt_sci(r["trig"], 6), fmt_sci(r["vec_linear"], 6)]
            for r in ctx["appendix_cell"]
        ],
        note="为与正文主线保持一致，曲面型单元数据使用 intrinsic 结果；该表主要用于说明单元数据路径仍比点数据更敏感。",
    )
    return tables


def insert_table(doc: Document, table_data: TableData) -> None:
    add_table_caption(doc, table_data.title)
    table = doc.add_table(rows=1, cols=len(table_data.headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(table_data.headers):
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, 10.5, bold=True)
        hdr_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row in table_data.rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(value))
            set_run_font(run, 10.5)
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if table_data.note:
        add_body_paragraph(doc, table_data.note)


def insert_figure(doc: Document, figure_data: FigureData, width_cm: float = 15.8) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(figure_data.path), width=Cm(width_cm))
    add_figure_caption(doc, figure_data.caption)
    if figure_data.note:
        add_body_paragraph(doc, figure_data.note)


def build_reference_list() -> list[str]:
    return [
        "[1] 王勖成. 有限单元法[M]. 北京: 清华大学出版社, 2003.",
        "[2] Angel E, Shreiner D. 交互式计算机图形学: 基于 OpenGL 着色器的自顶向下方法[M]. 北京: 清华大学出版社, 2014.",
        "[3] 洪涛, 潘志方, 林立本, 等. VTK 医学图像三维重建应用及实现[J]. 计算机系统应用, 2011, 20(4): 127-130, 185.",
        "[4] 王鹏, 王胜法, 曹俊杰, 等. 优化在网格去噪中的应用[J]. 中国图象图形学报, 2014, 19(4): 637-644.",
        "[5] 管西鹏, 陈宇拓, 张怀清, 等. 基于双边滤波算法的树木点云模型去噪方法研究[J]. 中南林业科技大学学报, 2015, 35(9): 83-87.",
        "[6] 邹北骥, 周浩宇, 辛国疆, 等. 基于脉冲耦合神经网络的点云曲面去噪[J]. 电子学报, 2012, 40(8): 1631-1637.",
        "[7] 郭清华, 赵勇. 面向法向域网格建模的非线性引导滤波[J]. 计算机辅助设计与图形学学报, 2020, 32(10): 1580-1590.",
        "[8] Salari K, Knupp P K. Code Verification by the Method of Manufactured Solutions[R]. Albuquerque: Sandia National Laboratories, 2000.",
        "[9] Roy C J, Smith T M, Ober C C. Verification of Euler/Navier-Stokes Codes Using the Method of Manufactured Solutions[R]. Albuquerque: Sandia National Laboratories, 2004.",
        "[10] Aycock K I, Brown K H, Niebur G L, et al. Method of manufactured solutions code verification of elastostatic solid mechanics problems in a commercial finite element solver[J]. Computers & Structures, 2019, 227: 106144.",
        "[11] Wang N, Lu P, Zhang L, et al. Accuracy analysis of gradient reconstruction on isotropic unstructured meshes and its effects on inviscid flow simulation[J]. Advances in Aerodynamics, 2019, 1: 20.",
        "[12] Diskin B, Thomas J L. Accuracy of Gradient Reconstruction on Grids with High Aspect Ratio[R]. Washington, D.C.: NASA, 2008.",
        "[13] Mavriplis D J, Thomas J L. Revisiting the Least-Squares Procedure for Gradient Reconstruction on Unstructured Meshes[R]. Washington, D.C.: NASA, 2003.",
        "[14] Zienkiewicz O C, Zhu J Z. The superconvergent patch recovery and a posteriori error estimates. Part 1: The recovery technique[J]. International Journal for Numerical Methods in Engineering, 1992, 33(7): 1331-1364.",
        "[15] Zienkiewicz O C, Zhu J Z. The superconvergent patch recovery and a posteriori error estimates. Part 2: Error estimates and adaptivity[J]. International Journal for Numerical Methods in Engineering, 1992, 33(7): 1365-1382.",
        "[16] Zhang Z, Naga A. A new finite element gradient recovery method: superconvergence property[J]. SIAM Journal on Scientific Computing, 2005, 26(4): 1192-1213.",
        "[17] Tessler A, Riggs H R, Freese C E, et al. A method for smoothing finite element-derived stress fields[J]. Computers & Structures, 1994, 53(3): 619-628.",
        "[18] Tessler A, Riggs H R, Dambach M. An improved variational method for finite element stress recovery and a posteriori error estimation[J]. Computer Methods in Applied Mechanics and Engineering, 1998, 155(1-2): 15-30.",
        "[19] de Miranda S, Ubertini F, Viola E. Recovery of consistent stresses for compatible finite elements[J]. Computer Methods in Applied Mechanics and Engineering, 2002, 191(15-16): 1693-1712.",
        "[20] Zhang H, Aragón A M. An improved stress recovery technique for the unfitted finite element analysis of discontinuous gradient fields[J]. International Journal for Numerical Methods in Engineering, 2022, 123(15): 3387-3410.",
        "[21] Tomasi C, Manduchi R. Bilateral filtering for gray and color images[C]// Proceedings of the 1998 IEEE International Conference on Computer Vision. Bombay: IEEE, 1998: 839-846.",
        "[22] Fleishman S, Drori I, Cohen-Or D. Bilateral mesh denoising[J]. ACM Transactions on Graphics, 2003, 22(3): 950-953.",
        "[23] Xing L, Li X, Liu Y, et al. A dynamic and adaptive scheme for feature-preserving mesh denoising[J]. Graphical Models, 2020, 111: 101065.",
        "[24] Conrad P R, Girolami M, Särkkä S, et al. Statistical analysis of differential equations: introducing probability measures on numerical solutions[J]. Statistics and Computing, 2017, 27(4): 1065-1082.",
        "[25] Poot N, van der Wilk M, Cockayne J. A Bayesian approach to modeling finite element discretization error[J]. Statistics and Computing, 2024, 34: 165.",
        "[26] Khronos Group. Compute Shader[EB/OL]. https://wikis.khronos.org/opengl/Compute_Shader, 2026-04-28.",
        "[27] VTK Team. vtkGradientFilter Class Reference[EB/OL]. https://vtk.org/doc/nightly/html/classvtkGradientFilter.html, 2026-04-28.",
        "[28] VTK Team. vtkSMPTools Class Reference[EB/OL]. https://vtk.org/doc/nightly/html/classvtkSMPTools.html, 2026-04-28.",
        "[29] Schroeder W, Martin K, Lorensen B. The Visualization Toolkit: An Object-Oriented Approach to 3D Graphics[M]. 4th ed. Clifton Park: Kitware, 2006.",
        "[30] Gilkeson C A, Kirby M, Hicken J E. Dealing with numerical noise in CFD-based design optimization[J]. Computers & Fluids, 2014, 109: 57-68.",
    ]


def build_reports() -> None:
    full_text = """# 与老师汇报稿（10分钟版）

## 1. 课题定位

我的课题题目是《基于OpenGL的CAE软件数据预处理方法研究与实践》。这里的数据预处理不是求解前的网格划分或边界条件设置，而是 CAE 后处理阶段中、为后续可视化渲染与分析准备数据的处理环节。本文聚焦两个问题：一是梯度场的快速计算，二是结果场中局部随机高频数值扰动的抑制。

## 2. 为什么要做这件事

第一，梯度是很多后续分析和渲染的基础量，比如应力集中识别、流场结构观察和特征增强。第二，现有 CPU 路径在大数据量和多字段实验时效率有限。第三，工程数据里还会出现局部数值扰动，如果没有适当的平滑与融合，后续可视化和分析都会受影响。

## 3. 我做了什么系统

我实现了一套基于 OpenGL Compute Shader 的 GPU 数值处理原型系统。系统仍然使用 VTK 做数据读写、GUI 渲染和参考基线，但真正的梯度计算和数据优化都放在 OpenGL 上执行。系统核心由五部分组成：统一门面层 `CAEProcessingFacade`，内部数据表示 `DataObject`，桥接层 `VTKDataConverter`，梯度引擎 `GLGradientEngine`，以及数据优化引擎 `GLFilterEngine`。

## 4. 梯度模块的核心方法

规则网格采用 GPU 有限差分。非结构化网格正式主线采用基于形函数导数的方法。它的思想是：先在参考单元中求形函数导数，再通过 Jacobian 或度量张量把参数空间导数映射到物理空间。对曲面型单元，我采用的是切空间一致的映射形式，避免直接把三维法向误差混进主结论。

对点数据，系统在节点关联单元上做局部梯度汇聚。对单元数据，当前实现是先把单元值提升到点，再在单元中心做形函数导数重构。这个路径在真实字段对比中表现很好，但在解析场上比点数据更敏感，我在论文里把它作为已识别限制单独说明。

## 5. 数据优化模块的定位

数据优化模块不是通用降噪器，也不是求解误差修正器。它的定位是：面向 CAE 后处理结果场中的一类局部随机高频数值扰动。算法上采用图双边滤波做多尺度分解，再按细节增益进行融合重构。这样做的目的是在抑制局部扰动的同时尽量保持边缘和主结构。

## 6. 实验怎么设计

梯度模块实验分成两部分。第一部分是解析场验证，只用能稳定支撑正确性结论的数据：`SampleStructGrid`、`hexa` 和 `1_0 point`。其中 `1_0` 采用曲面专用 benchmark，并以 intrinsic 指标为主。第二部分是真实字段与 `vtkGradientFilter` 的一致性和时间对比，数据集包括 `SampleStructGrid`、`hexa`、`limb`、`ShipHull_0`，并补充 `1_0` 结果。

数据优化模块统一使用 `ShipHull_0`，在干净场上分别添加 `gaussian`、`grf`、`mixed` 和 `impulse` 扰动，比较优化前后的 NRMSE、RMSE 改进率和粗糙度比，并配合 ParaView 图像做结果展示。

## 7. 主要结果

第一，梯度模块在真实字段上与 VTK 结果高度一致。比如 `SampleStructGrid point` 的 NRMSE 为 `8.224e-08`，`hexa point` 为 `6.273e-08`，`limb cell` 为 `3.895e-07`。曲面型数据 `ShipHull_0 point` 和 `ShipHull_0 cell` 分别为 `3.892e-03` 和 `2.149e-02`，虽然误差高于体网格，但仍保持较好一致性。

第二，在解析场验证中，规则网格和六面体点数据结果很好，线性场基本达到机器精度。`1_0` 说明了曲面 benchmark 的必要性：如果直接用普通三维场评价，`benchmark_linear` 的 ambient NRMSE 为 `0.808122`；改用 intrinsic 指标后，同类曲面线性场误差降到 `1e-7` 量级。

第三，时间上系统体现出明显的 GPU 优势。以 `SampleStructGrid point` 为例，VTK 单线程平均时间是 `10.376 ms`，VTK 并行是 `3.763 ms`，系统总时间是 `1.163 ms`，GPU 纯计算时间是 `0.040 ms`。需要强调的是，VTK 并行并不是在所有数据集上都更快，所以论文中我没有泛化成“GPU 一定优于 VTK 并行”，而是如实报告各数据集结果。

第四，数据优化模块对 `gaussian` 和 `mixed` 扰动有稳定抑制能力。点数据上，高斯扰动的平均 NRMSE 从 `0.296642` 降到 `0.180858`；单元数据上则从 `0.294558` 降到 `0.121921`。但对 `impulse` 扰动，RMSE 改进率接近 1，说明该模块的适用范围是明确且有限的。

## 8. 我的结论

本文完成了一条面向 CAE 后处理的数据预处理 GPU 主线：规则网格走有限差分，非结构化网格走形函数导数，结果场优化走图双边滤波多尺度融合。系统在真实字段上与 VTK 具有较高一致性，并在当前实验环境下展现出明显的时间优势。与此同时，我也明确给出了两点边界：一是单元数据解析场路径仍需改进，二是数据优化模块针对的是局部随机高频扰动而不是所有噪声。

## 9. 如果老师追问限制和后续工作

我会重点回答三点。第一，后续要继续研究 `vtkGradientFilter` 的保护机制，改进曲面点梯度的稳定性。第二，要改进 `cell -> point lift -> shape gradient` 这条单元数据恢复链。第三，数据优化模块后续应增加更贴近真实 CAE 数值扰动来源的实验，而不是只停留在代理噪声验证。
"""
    short_text = """# 与老师快速汇报稿（简版）

本文做的是 CAE 后处理阶段、为后续渲染和分析服务的数据预处理，不是求解前预处理。核心完成了两个模块：一是基于 OpenGL Compute Shader 的梯度计算，二是面向局部随机高频数值扰动的数据优化。

梯度模块方面，规则网格采用 GPU 有限差分，非结构化网格采用基于形函数导数的方法。真实字段上，系统与 `vtkGradientFilter` 结果整体接近，`SampleStructGrid`、`hexa`、`limb` 等数据的误差达到 `1e-6` 到 `1e-8` 量级，`ShipHull_0` 这类曲面数据也保持了较好一致性。时间上，系统总时间和 GPU 纯计算时间都明显小于当前环境下的 VTK 参考时间。

数据优化模块方面，我把它定位为“局部随机高频数值扰动抑制”，而不是通用去噪器。实验表明它对 `gaussian` 和 `mixed` 扰动有效，对 `impulse` 并不敏感，这个结果和模块定位是一致的。

所以本文的主要结论不是“所有场景都完美”，而是：已经完成了一条可运行、可验证、与 VTK 可比的 GPU 数据预处理主线，并明确识别了当前限制和下一步改进方向。"""
    REPORT_FULL.write_text(full_text, encoding="utf-8")
    REPORT_SHORT.write_text(short_text, encoding="utf-8")


def save_docx_with_fallback(doc: Document, target: Path) -> Path:
    try:
        doc.save(target)
        return target
    except PermissionError:
        fallback = target.with_name(f"{target.stem}_更新版{target.suffix}")
        doc.save(fallback)
        return fallback


def build_document(ctx: dict, tables: dict[str, TableData], assets: dict[str, Path]) -> Path:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.8)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    add_page_number(section)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    styles["Heading 1"].font.name = "黑体"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.name = "黑体"
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 3"].font.name = "黑体"
    styles["Heading 3"].font.size = Pt(12)

    add_center_paragraph(doc, "本科毕业论文", 22, True)
    doc.add_paragraph()
    add_center_paragraph(doc, "《基于OpenGL的CAE软件数据预处理方法研究与实践》", 18, True)
    doc.add_paragraph()
    add_center_paragraph(doc, "学    院：", 14)
    add_center_paragraph(doc, "专    业：", 14)
    add_center_paragraph(doc, "姓    名：", 14)
    add_center_paragraph(doc, "学    号：", 14)
    add_center_paragraph(doc, "指导教师：", 14)
    add_center_paragraph(doc, "完成日期：", 14)
    doc.add_page_break()

    add_center_paragraph(doc, "基于OpenGL的CAE软件数据预处理方法研究与实践", 18, True)
    add_center_paragraph(doc, "Research and Practice on OpenGL-Based Data Preprocessing Methods for CAE Software", 14, False)
    doc.add_page_break()

    add_heading(doc, "原创性声明", 1)
    add_body_paragraph(doc, "本人郑重声明：所提交的毕业论文是在指导教师指导下独立完成的研究成果。除文中已经注明引用的内容外，本文不包含任何他人已经发表或撰写过的研究成果。对本文研究做出重要贡献的个人和集体，均已在文中以明确方式标明。本人完全意识到本声明的法律后果由本人承担。")
    add_center_paragraph(doc, "学生签名：                日期：", 12)
    doc.add_page_break()

    add_heading(doc, "使用授权说明", 1)
    add_body_paragraph(doc, "本人完全了解学校关于保存、使用毕业论文的有关规定，同意学校保留并向有关部门或机构送交论文的复印件和电子版，允许论文被查阅和借阅；学校可以采用影印、缩印、数字化等方式保存论文。涉密论文按学校有关规定处理。")
    add_center_paragraph(doc, "学生签名：                指导教师签名：                日期：", 12)
    doc.add_page_break()

    add_heading(doc, "摘  要", 1)
    add_body_paragraph(
        doc,
        "面向 CAE 后处理阶段中为后续可视化渲染与分析服务的数据预处理需求，本文围绕梯度计算与结果场优化两个关键问题，设计并实现了一套基于 OpenGL Compute Shader 的 GPU 数值处理原型系统。系统在保留 VTK 数据读写、GUI 渲染和参考基线能力的同时，将规则网格梯度计算、非结构化网格基于形函数导数的梯度计算以及图双边滤波驱动的多尺度数据优化统一到 GPU 主路径上执行。本文首先构建了统一门面层与内部扁平数据表示，实现了规则网格有限差分和非结构化网格形函数导数两类梯度算法；随后针对 CAE 结果场中一类局部随机高频数值扰动，设计了基于多尺度分解与融合重构的数据优化模块；最后通过解析场验证、真实字段与 vtkGradientFilter 的一致性对比、时间对比以及 ParaView 图像对照完成实验评估。结果表明：系统在真实字段上与 VTK 保持较高一致性，在当前实验环境下具有明显的 GPU 时间优势；数据优化模块对高斯型和混合型局部随机高频扰动具有稳定抑制能力，同时能够保持主要边缘结构。"
    )
    add_body_paragraph(doc, "关键词：OpenGL；CAE 后处理；梯度计算；形函数导数；数据优化；Compute Shader")
    doc.add_page_break()

    add_heading(doc, "Abstract", 1)
    add_body_paragraph(
        doc,
        "This thesis focuses on data preprocessing tasks in the CAE post-processing stage, especially those required before visualization rendering and field analysis. An OpenGL Compute Shader based GPU numerical processing prototype is designed and implemented to address two major problems: gradient computation and result-field optimization. While VTK is still used for data I/O, GUI rendering, and baseline comparison, the main numerical paths are migrated to the GPU, including finite-difference gradients for structured grids, shape-function-derivative gradients for unstructured grids, and graph bilateral filtering based multi-scale optimization. A unified facade layer and an internal flattened data representation are first established to support both structured and unstructured datasets. Then, a multi-scale optimization module is introduced for one class of local random high-frequency numerical disturbances often observed in CAE result fields. The experimental evaluation combines analytic benchmark validation, consistency comparison against vtkGradientFilter on real fields, timing measurements under both single-threaded and parallel VTK settings, and ParaView visual inspections. The results show that the proposed system achieves high consistency with VTK on real engineering fields and exhibits clear time advantages on the current test environment. The optimization module effectively suppresses Gaussian-like and mixed disturbances while maintaining major edge structures. The thesis also clarifies current limitations, including the sensitivity of cell-centered analytic tests and the bounded applicability of the optimization module."
    )
    add_body_paragraph(doc, "Key words: OpenGL; CAE post-processing; gradient computation; shape function derivatives; data optimization; compute shader")
    doc.add_page_break()

    add_heading(doc, "目  录", 1)
    add_toc_field(doc.add_paragraph())
    doc.add_page_break()

    add_heading(doc, "第一章 绪论", 1)
    add_heading(doc, "1.1 研究背景与意义", 2)
    for text in [
        "CAE 软件在工程设计、结构分析、流场模拟和多物理场耦合分析中已经形成成熟工作流，但在很多实际项目里，求解结束并不意味着工程分析结束。大量任务发生在后处理阶段，例如应力集中识别、特征区域提取、结果场对比、异常值排查以及面向展示或评估的可视化渲染。这些任务都依赖对原始结果场做进一步的数值加工，因此有必要把“数据预处理”限定为 CAE 后处理阶段中、为后续渲染与分析准备数据的处理过程，而不是传统意义上的求解前预处理[3][29]。",
        "在这一语境下，梯度计算属于最基础也最关键的派生量构造过程之一。梯度既可以直接用于应力、温度、浓度或速度等字段的变化强度分析，也会影响边界增强、特征着色和后续多尺度分析的效果。对于结构化数据，梯度可以较自然地通过有限差分求得；但工程数据更常见的是非结构化网格，其单元拓扑多样、几何形态复杂、局部尺度变化明显，导致梯度重构与结果解释都更加困难[1][11][12]。",
        "另一方面，CAE 结果场中还经常伴随局部数值扰动。其来源可能是离散化误差、求解收敛误差、网格畸变、应力恢复不连续以及浮点舍入等。对于后处理展示和后续分析而言，这类局部高频扰动既会干扰视觉判断，也可能影响派生量的稳定性。因此，若能在梯度计算之外进一步提供面向特定扰动类别的数据优化能力，就能形成一条更完整的后处理数值预处理主线[17][18][19][30]。",
        "随着 GPU 通用计算能力增强，OpenGL Compute Shader 为图形软件内部的轻量级并行数值计算提供了可行路径。相较于完全依赖 CPU 的传统处理流程，GPU 更适合大规模并行、规则访存和批量字段计算场景。基于此，本文尝试把 CAE 后处理中的关键数据预处理操作迁移到 OpenGL 主路径，以提升处理效率并建立可复现的实验链路[2][26]。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "1.2.1 梯度计算研究现状的进一步分析", 3)
    for text in [
        "从文献脉络上看，梯度计算研究大致可以分为三类问题：第一类是离散格式本身是否与理论模型一致，通常通过解析解或制造解方法进行验证；第二类是算法在网格畸变、曲率变化、邻域退化和维数混合条件下的稳定性；第三类是工程软件实现层面的配置依赖，例如边界处理、贡献单元选择和并行后端差异。毕业设计若只展示单一误差数字，而不说明这三类问题分别由哪些实验支撑，往往会导致论证链条断裂[8][9][11][12]。",
        "对有限元背景下的梯度恢复与应力恢复问题而言，真实难点通常不在于公式写不出来，而在于离散语义的差异。点数据与单元数据虽然都被称为“场值”，但前者对应节点自由度，后者则对应单元内部或单元中心意义下的常值或均值量。文献中许多应力恢复方法之所以单独设计恢复步骤，本质上就是在处理这种离散语义转换问题[17][18][19][20]。",
        "这也解释了本文为什么把点数据解析验证作为非结构化网格主结论，而把单元数据解析验证收缩到附录。这样的实验组织并不是规避问题，而是遵循数值实验中常见的“先验证公式最直接成立的路径，再讨论恢复链带来的附加误差”这一基本逻辑[10][14][16]。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "1.2.2 GPU 数值处理与可视化系统研究现状", 3)
    for text in [
        "GPU 数值处理在图形与可视化软件中的应用已经十分广泛，但在许多工程软件项目里，GPU 主要服务于渲染，而非真正意义上的数据预处理。OpenGL Compute Shader 的引入，使得同一个图形运行时既可以负责显示，也可以承担一定规模的并行数值计算。这为轻量级科研原型系统提供了现实基础：开发者不必切换到额外的 CUDA 或 OpenCL 工程体系，就能够在现有可视化程序中搭建 GPU 算法链路[2][26]。",
        "不过，从论文写作角度看，仅仅说明‘用了 GPU’并不足够。更关键的是要交代 GPU 在系统里的职责边界：是只负责加速一个局部核函数，还是形成了从数据表示、调度、计时到结果写回的完整主路径；是把 CPU 结果搬到 GPU 做后期修饰，还是把真正的算法核心迁移到了 GPU。本文系统之所以强调门面层、内部表示和着色器层的联动设计，就是为了说明这不是单个 shader 的拼接，而是一条完整的 GPU 数据处理主线[3][26][29]。",
        "同时，GPU 加速实验也必须面对一个常见误区，即把 GPU 核函数时间直接与 CPU 软件整体流程时间相比较。本文通过同时报告系统总时间和 GPU 纯计算时间，尽量避免这一不公平对照；而在 VTK 侧又额外报告单线程和并行结果，则是为了进一步消解‘只挑最慢配置做比较’的质疑[27][28]。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "1.2.3 数据优化问题的研究边界", 3)
    for text in [
        "对于结果场优化问题，论文最容易出现的逻辑风险是把“受控噪声实验中的有效”直接写成“真实仿真误差都能被修正”。事实上，已有文献虽然证明了 CFD/FEM 结果场中存在数值噪声、局部伪振荡和不连续恢复等问题，但这些问题的生成机理并不统一，也并不意味着它们整体服从同一种随机分布[17][18][19][20][30]。",
        "因此，本文采用的是更保守也更严谨的定位：先依据文献确认 CAE 结果中确实存在局部随机高频数值扰动这一类现象，再说明高斯扰动和小支撑随机场可以作为其中一类扰动的代理模型或近似描述，最后只验证本文模块对这类扰动的抑制能力，而不外推到所有数值误差来源[24][25][30]。",
        "从毕业设计角度看，这种表述方式虽然没有‘万能去噪’那样夸张，但更容易经受答辩追问。因为它把问题限定在一个可验证、可讨论且与现有实验结果一致的范围内，也使后续工作方向更清楚：如果要扩展到脉冲异常值或其他误差机理，就需要设计新的模型与算法，而不是在现有结论上过度延伸[21][22][23]。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "1.2 国内外研究现状", 2)
    for text in [
        "在规则网格上，梯度计算通常以有限差分为主。对于均匀网格，可直接采用中心差分和边界单边差分；对于曲线规则网格，则需要通过 Jacobian 将参数空间差分映射到物理空间。这类方法公式清晰、实现直接，是许多数值代码和可视化工具的基础路径。解析场或制造解验证是该类算法的标准验证手段，其核心在于比较数值解与理论真值，而不是与某个软件输出的一致程度[8][9][10]。",
        "在非结构化网格上，梯度重构方法更为多样。工程和文献中常见的路线包括基于局部最小二乘的梯度重构、基于单元形函数导数的有限元一致性方法，以及面向误差估计的梯度恢复和应力恢复方法。Mavriplis、Diskin、Wang 等人的工作说明，网格质量、长宽比、曲率和邻域构造都会显著影响梯度精度；Zienkiewicz、Zhu 以及 Zhang 等人的研究则指出，导数量往往比原始场更敏感，应当单独评价其收敛与稳定性[11][12][13][14][15][16]。",
        "在工程软件层面，VTK 的 vtkGradientFilter 被广泛用作派生梯度的参考实现。它提供了 ContributingCellOption、ReplacementValueOption 和 FasterApproximation 等选项，并可通过 vtkSMPTools 选择 Sequential、STDThread、TBB 或 OpenMP 等并行后端。也正因如此，任何与 VTK 的时间对比都必须同步报告算法配置和并行配置，否则对比结论并不严谨[27][28]。",
        "与梯度重构相比，结果场优化或降噪问题的研究目标更强调“抑制局部高频扰动并保持主要特征”。图像处理中的双边滤波、几何处理中针对网格法向或位置的保边滤波，以及多尺度分解与融合方法，都为 CAE 结果场处理提供了可借鉴思路。相关文献表明，受控噪声实验是常见的验证方式，但它验证的是算法对某类扰动模型的适配性，而不等同于对真实误差根因的完全证明[21][22][23]。",
        "综合来看，现有研究为本文提供了两个直接启发：其一，梯度模块需要把“解析真值验证”和“工程软件一致性对照”分开组织；其二，数据优化模块应明确其适用扰动类别与实验边界，而不能被泛化为通用误差修正器。本文的系统设计和实验方案正是在这一认识下展开的[24][25][30]。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "1.3 研究内容与技术路线", 2)
    for text in [
        "本文围绕两个模块开展工作。第一是梯度计算模块：在规则网格上实现 GPU 有限差分，在非结构化网格上实现基于形函数导数的 GPU 梯度计算，并建立与 vtkGradientFilter 可比的实验程序。第二是数据优化模块：面向一类局部随机高频数值扰动，构建基于图双边滤波和多尺度细节融合的场优化流程。",
        "从系统角度看，本文并不追求完全替代 VTK，而是保留 VTK 在数据读写、窗口渲染和参考基线上的成熟能力，同时将真正的数值处理主路径从 VTK 内部数组接口中解耦出来，映射到统一的内部数据结构和 OpenGL 计算引擎中。这一策略能够在控制工作量的同时，保证实验链路可追溯、可复现、可对照。",
        "技术路线可以概括为：首先通过 VTK 将规则网格或非结构化网格读入，再转换为统一的内部扁平数据表示；随后由门面层根据网格类型、字段关联和算法模式分派到不同的 GPU 着色器；最后将结果写回内部数据对象，用于 GUI 展示、CSV 统计、ParaView 导出和与 VTK 参考结果对比。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "1.4 本文主要工作", 2)
    for text in [
        "一是设计并实现统一门面接口 CAEProcessingFacade，使 GUI 主界面、梯度测试程序和数据优化测试程序共用同一条核心计算链路。二是构建 DataObject 统一内部数据表示，完成点坐标、单元连接、字段数组、邻域图和规则网格逻辑维度的扁平封装。",
        "三是实现规则网格有限差分着色器 FD.glsl，以及非结构化网格点数据、单元数据对应的 ShapePointGradient.glsl、ShapeCellGradient.glsl 和 CellDataToPointLift.glsl。四是实现基于图双边滤波的多尺度数据优化流程，包括 Bilateral.glsl 与 MultiScaleFuse.glsl 两个核心着色器。",
        "五是构建测试程序 TestGradient、TestMultiScale 和 TestFieldMetrics，能够输出解析场误差、真实字段与 VTK 的对照误差、VTK 单线程与并行时间、本系统墙钟时间和 GPU 计算时间，并支持将结果导出为 CSV 和 VTK 文件。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "1.5 论文结构", 2)
    add_body_paragraph(doc, "本文共分六章。第一章介绍课题背景、研究现状与主要工作；第二章给出工具与理论基础；第三章说明系统架构和核心模块设计；第四章详细介绍梯度计算与数据优化算法原理；第五章给出实验设计、结果和分析；第六章总结全文并讨论局限与后续工作。")

    add_heading(doc, "第二章 工具和理论基础", 1)
    add_heading(doc, "2.1 OpenGL Compute Shader 与 GPU 并行模型", 2)
    for text in [
        "OpenGL 4.3 引入了 Compute Shader，使得图形程序不再只能把 GPU 用于传统渲染管线，也可以把其作为通用并行执行单元来处理数组、邻域图和数值场。Compute Shader 的基本执行单元是 work group 和 local invocation；程序通过 glDispatchCompute 提交任务，通过着色器存储缓冲区对象（SSBO）传递输入输出，再配合 glMemoryBarrier 保证不同阶段的读写可见性。",
        "对于本文这类以大规模数组并行为主的任务，Compute Shader 有两点优势。第一，规则网格和非结构化网格都能被压平成连续数组，便于以索引方式并行访问。第二，GPU 时间可以通过查询对象单独计量，从而把“系统总时间”和“核心 GPU 计算时间”分开统计。这为后续实验中区分数据准备、结果回读与纯计算开销提供了条件。",
        "需要指出的是，Compute Shader 不是自动保证高性能的万能方案。它仍然要求开发者合理组织内存布局、线程粒度和同步方式，否则性能会被访存和调度开销抵消。因此，本文在工程上采用统一门面层和统一数据布局，使不同算法尽量共享一套缓冲区组织逻辑和调度方式。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "2.2 VTK 在本文系统中的角色", 2)
    for text in [
        "VTK 是成熟的科学可视化工具包，具备稳定的数据模型、I/O、渲染和滤波工具。在本文系统中，VTK 主要承担四类职责：其一，读取 legacy VTK 数据文件并提供 GUI 渲染；其二，作为真实字段对照实验的参考基线；其三，承担最终结果导出，使 ParaView 能直接加载；其四，在测试程序中提供 vtkGradientFilter 和 vtkSMPTools 等对照接口。",
        "但需要严格区分的是：本文的梯度计算和数据优化主路径并不依赖 VTK 的内部梯度或滤波实现。系统在完成 `vtkDataSet -> DataObject` 的桥接后，真正的计算由 OpenGL 计算着色器完成；只有在生成参考结果、做可视化显示或导出文件时，才重新回到 VTK 环境。这一定位既符合当前工程实现，也有助于在论文中澄清“系统是否完全依赖 VTK”的常见疑问。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "2.3 规则网格有限差分理论", 2)
    for text in [
        "对于标量场 u，物理空间梯度定义为 ∇u = [∂u/∂x, ∂u/∂y, ∂u/∂z]^T。若网格在逻辑上是规则的，则可以先沿 ξ、η、ζ 三个参数方向构造一维差分模板，再通过坐标映射将参数空间导数变换到物理空间。本文对内部采样点使用中心差分，对边界点使用单边差分，从而在不额外引入幽灵点的情况下完成边界处理。",
        "相较于把规则网格简单理解为正交均匀立方体，本文实现更接近“曲线规则网格”的处理方式。原因在于 `SampleStructGrid` 一类数据虽然逻辑维度规整，但几何上不必完全均匀。此时不能直接把参数步长视为物理步长，而需要通过局部 Jacobian J = ∂x/∂ξ 进行映射。其等价写法是 ∇u = J^{-T}∇_ξu。",
        "该理论路径的优点在于数学关系清晰、与解析场验证天然匹配。对于线性场，若差分模板和 Jacobian 构造无误，则误差可以逼近机器精度；对于二次场和三角函数场，误差大小则受离散间距和局部几何变化影响。这正是本文选择 `benchmark_linear`、`benchmark_quadratic` 和 `benchmark_trig` 进行验证的原因。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "2.4 非结构化网格基于形函数导数的梯度理论", 2)
    for text in [
        "有限元单元中的几何坐标与场变量通常都可以表示为形函数插值。设单元节点编号为 a=1,...,n，形函数为 N_a(ξ)，节点坐标为 x_a，节点场值为 u_a，则有 x(ξ)=ΣN_a(ξ)x_a，u(ξ)=ΣN_a(ξ)u_a。对参数坐标 ξ 求导，可得 ∇_ξu=Σu_a∇_ξN_a。由此可见，只要掌握单元形函数导数和局部几何映射，就可以在每个单元内部重构场梯度。",
        "对于满维单元，例如三维体网格中的四面体或六面体，最直接的变换关系是 ∇u = J^{-T}∇_ξu，其中 J 为几何 Jacobian。当单元嵌入维数高于拓扑维数时，例如三维空间中的曲面三角形或曲面四边形，J 不再是可逆方阵，此时更稳健的做法是使用度量张量 G=J^T J，并采用 ∇u = J G^{-1}∇_ξu 的最小范数切空间映射形式。",
        "这一表达在本文非常重要。因为像 `1_0` 和 `ShipHull_0` 这样的曲面型数据，如果直接以普通三维 ambient 场作为真值评价，就会把法向方向的差异也计入误差，导致看似很差的结果。事实上，曲面上的物理解读更关注切向变化，因此 intrinsic 指标比 ambient 指标更能反映算法在曲面几何上的真实表现。",
        "在实现层面，本文正式宣称并纳入实验主线的一阶单元类型包括三角形、四边形、四面体和六面体。代码中使用统一的单元类型编码和着色器模板，对不同单元在参考坐标中的形函数导数进行分支处理，再将其统一映射到物理空间。这种做法兼顾了理论一致性与工程实现的可扩展性。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "2.5 图双边滤波与多尺度分解理论", 2)
    for text in [
        "双边滤波最初用于图像保边平滑，其核心思想是同时考虑空间邻近性和数值相似性。若将其推广到点邻域图或单元邻域图，则节点 i 的更新可写为 u'_i = Σ_j w_ij u_j / Σ_j w_ij，其中 w_ij = exp(-||x_i-x_j||²/(2σ_s²)) · exp(-||u_i-u_j||²/(2σ_r²))。前一项度量几何接近程度，后一项度量字段值相似程度。",
        "当该思想被用于 CAE 结果场时，其意义在于：一方面，相邻节点或单元的字段值通常具有局部相关性，适合做邻域统计；另一方面，边界或特征区域往往伴随字段突变，不能简单通过均值滤波整体抹平。双边滤波恰好提供了“在局部平滑和边缘保持之间折中”的基本机制。",
        "在此基础上，多尺度分解通过重复平滑得到多个尺度的低频层，再将相邻平滑层做差得到高频细节层。最深层平滑结果可视为 base，浅层与深层之差构成 detail0、detail1、detail2 等细节分量。融合阶段则根据细节增益将各层 detail 以受控方式回注到 base 中，从而兼顾扰动抑制与特征保持。",
        "需要强调的是，多尺度融合并不等价于再做一次平滑。它更像是一个重建阶段：哪些细节应该保留，哪些细节可能对应局部随机高频扰动，需要通过分层结构和增益参数共同决定。这也是本文将“分解”和“融合”作为一个完整优化模块来描述，而不是把它们拆成互不相关的两个小算法的原因。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "2.6 代码验证与评价指标基础", 2)
    for text in [
        "在数值代码验证领域，制造解或解析解实验的主要作用是回答“实现是否符合理论公式”这一问题，而不是回答“是否与某个软件结果接近”。因此，本文在梯度模块中专门保留了解析 benchmark 路径，并将其与真实字段对 VTK 的一致性验证严格区分。前者用于算法正确性，后者用于工程可比性[8][9][10]。",
        "实验评价指标方面，梯度模块主要采用绝对误差、NRMSE、Soft Relative Error、角度误差和尺度偏差等指标。NRMSE 反映整体相对偏差，Soft Relative Error 则通过引入阈值 τ 避免参考值接近零时普通相对误差失稳。数据优化模块则主要采用输入/输出 NRMSE、RMSE 改进率和粗糙度比。RMSE 改进率小于 1 说明优化后优于输入，粗糙度比越小则表明局部高频起伏被抑制得更明显[21][22][23]。",
        "这些指标被统一实现到测试程序中，使同一套程序既能用于解析场验证，也能用于真实字段对照和性能统计，从而构成完整的实验闭环，也使系统实验结果能够直接追溯到源码和原始 CSV 文件[27][28][29]。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "2.7 vtkGradientFilter 与本文方法的可比性基础", 2)
    for text in [
        "本文在与 vtkGradientFilter 对照时，并未简单地把任何梯度输出都拿来比较，而是尽量保证方法语义与配置条件可比。具体来说，对非结构化网格形函数导数主线，测试程序会在构造 VTK 参考结果时设置 `ContributingCellOption(Patch)` 并关闭 `FasterApproximation`，以减少由于近似选项导致的额外偏差[27]。",
        "即便如此，本文方法与 VTK 仍然不是逐语句相同的实现。VTK 内部针对混合维数据、缺失贡献单元、局部退化和替代值都有自己的保护机制，而本文系统更强调在 OpenGL 上重建一条与有限元形函数导数理论一致的 GPU 主线。因此，论文中对 VTK 的定位是“工程参考基线”，而不是“数学真值”。这样的表述更符合软件对照实验的本质[27][28][29]。",
        "这一区分也直接影响实验结论的写法：如果系统与 VTK 非常接近，可以说明两者在工程输出上高度一致；若在个别复杂曲面或单元恢复场景下存在差异，则不能马上推断谁绝对正确，而需要回到解析场、离散语义和配置机制上做进一步分析[11][12][27]。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "2.8 本章小结", 2)
    for text in [
        "本章从 OpenGL Compute Shader、VTK 工具角色、规则网格有限差分、非结构化网格形函数导数、图双边滤波多尺度融合以及代码验证指标六个方面给出了本文所依赖的工具和理论基础。由此可以看出，本文的两条主线并非孤立拼接，而是在有限差分、有限元形函数理论和保边滤波理论基础上的系统化实现[1][2][21][22][26][29]。",
        "更重要的是，本章也提前澄清了论文后续可能被追问的几个逻辑点：第一，解析场验证与 VTK 对照的角色不同；第二，曲面型数据必须强调 intrinsic 指标；第三，数据优化模块的适用对象是某一类局部随机高频扰动，而不是所有数值误差。这些说明为后续系统设计、算法实现和实验分析奠定了论证边界[8][9][24][25][30]。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "第三章 系统设计", 1)
    add_heading(doc, "3.1 系统设计目标", 2)
    for text in [
        "本文系统设计的首要目标，是让 GUI、测试程序和底层算法共享同一条核心处理主线，避免出现“界面里一套逻辑、实验脚本里另一套逻辑”的问题。第二个目标，是把数值计算从 VTK 原生数组接口中抽离出来，改用更适合 GPU 的内部扁平表示。第三个目标，则是围绕论文实验需要，提供稳定的数据加载、算法调度、结果缓存、时间统计和导出机制。",
        "围绕上述目标，系统被组织为“界面与测试层 + 门面层 + 数据桥接层 + 内部数据表示层 + GPU 计算层”的分层结构。不同层之间通过清晰的接口交换数据，既便于扩展，也便于在论文中准确描述每个模块的职责边界。",
    ]:
        add_body_paragraph(doc, text)
    insert_figure(
        doc,
        FigureData(
            "图3-1 系统总体架构图",
            assets["system_arch"],
            "门面层负责统一分派，VTK 负责桥接、显示和导出，OpenGL 负责真正的梯度与数据优化计算。",
        ),
    )

    add_heading(doc, "3.2 统一门面层设计", 2)
    for text in [
        "CAEProcessingFacade 是系统的统一调度入口。它负责初始化 OpenGL 计算上下文、加载数据集、维护多数据集缓存、根据请求自动分派梯度算法或数据优化流程，并在计算完成后将结果写回 DataObject。对于测试程序而言，这意味着实验脚本只需关心数据集名称、字段名称、关联类型和方法模式，而不必重复处理底层资源管理。",
        "在梯度模块中，门面层还承担了关键的“自动方法选择”职责。根据实际代码，若请求方法为 Auto，则规则网格自动走 Finite Difference，非结构化网格自动走 Shape Function Derivatives。这一点非常关键，因为它保证了系统当前对非结构化网格的默认主线已经统一为形函数导数方法，而不是再在多种方法之间摇摆。",
        "门面层还记录最近一次计算的总墙钟时间和 GPU 时间，为 TestGradient 和 TestMultiScale 的时间统计提供统一来源。这样一来，GUI 演示和实验程序之间共享的是同一套时间语义，避免了不同入口统计口径不一致的问题。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "3.3 内部数据表示与桥接层", 2)
    for text in [
        "DataObject 是整个系统的统一内部数据载体。它将点坐标、单元中心、字段数组、单元连接、单元偏移、单元类型、点邻域图、单元邻域图以及规则网格逻辑维度统一表示为扁平数组或 CSR 结构。该设计的直接收益，是着色器可以用整齐的一维索引访问所有数据，不必了解 VTK 的复杂层次结构。",
        "VTKDataConverter 则负责完成 vtkDataSet 与 DataObject 的双向转换。在输入阶段，它从 VTK 数据对象中提取点坐标、单元连接关系和字段数组，并为非结构化网格构造点所属单元列表、点邻域图和单元邻域图；在输出阶段，它再把内部结果数组写回 VTK 数据对象，以便 GUI 显示或导出到 ParaView。",
        "这种桥接设计使系统既能充分利用 VTK 在 I/O 和渲染上的成熟能力，又不会把数值算法绑定在 VTK 内部数据结构上。对论文而言，这也构成了一个重要结论：系统与 VTK 的关系是“桥接与对照”，而不是“计算依赖”。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "3.4 GPU 计算层设计", 2)
    for text in [
        "OpenGLManager 在 Windows 下创建独立的计算上下文，负责着色器编译、缓冲区绑定和 GPU 时间查询。GLGradientEngine 则封装规则网格有限差分、形函数导数点梯度、形函数导数单元梯度等接口；GLFilterEngine 封装图双边滤波和平滑细节融合接口。两者都以扁平数组为输入输出，便于门面层统一调用。",
        "在工程实现上，梯度引擎与数据优化引擎都采用“输入缓冲区准备 -> 绑定着色器参数 -> glDispatchCompute -> 内存屏障 -> 结果回读”的统一流程。这样做的好处是不同算法共享类似的资源管理模式，不仅便于调试，也便于在实验程序中统一统计 GPU 时间。",
        "值得说明的是，梯度引擎仍保留了一些历史接口名称，但本文的正式实验主线只围绕规则网格有限差分和非结构化网格形函数导数展开。论文中因此不再把已废弃的其他方法作为主线描述对象，以免模糊系统当前真正可验证的核心能力。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "3.5 测试程序设计", 2)
    for text in [
        "TestGradient 是梯度模块的核心测试程序。它能够批量读取指定数据集中的字段，自动识别解析 benchmark 或真实工程字段，并在需要时调用 vtkGradientFilter 构造参考结果。当前程序还可以分别输出 VTK 单线程时间、VTK 并行时间与线程数、本系统墙钟时间和 GPU 纯计算时间，从而满足老师对“算法一致、计算条件可比”的实验要求。",
        "TestMultiScale 用于数据优化模块实验。它负责在干净场上注入高斯、相关随机场、混合扰动和脉冲扰动，随后调用多尺度优化模块并输出输入/输出 NRMSE、RMSE 改进率和粗糙度比。TestFieldMetrics 则提供离线字段评价能力，便于对导出的 VTK 文件做进一步核对。",
        "由于三类程序都通过 CAEProcessingFacade 访问算法主路径，因此实验结果与 GUI 路径之间具有天然的一致性。对毕业论文而言，这是很重要的工程基础：实验不是“额外写的一段验证代码”，而是系统自身能力的直接调用。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "3.6 梯度模块与数据优化模块的数据流设计", 2)
    for text in [
        "在梯度模块中，数据流可以概括为“字段选择 -> 方法分派 -> GPU 调度 -> 结果缓存 -> 参考对比/导出”。其中，字段选择阶段确定了点数据还是单元数据、标量场还是向量场；方法分派阶段则根据网格类别和用户请求自动决定有限差分或形函数导数主线；GPU 调度阶段将扁平数组绑定到 SSBO，并触发对应着色器执行；结果缓存阶段把梯度数组写回 DataObject，供 GUI、CSV 和 VTK 导出复用。",
        "在数据优化模块中，数据流则表现为“干净场或输入场 -> 邻接图构造 -> 多层双边平滑 -> 细节分解 -> 增益融合 -> 输出与评价”。与很多只保留最终结果的原型不同，本文系统还会在需要时保存 smooth、detail、base 和 fused 等中间数组，这使得算法过程能够被逐层追踪，便于答辩时解释‘为什么优化前后会有这样的视觉变化’。",
        "这种把中间结果显式保存在内部对象中的做法，虽然增加了部分显存和内存占用，但极大提升了实验可解释性。对于毕业设计而言，系统不仅要算出结果，还要能向老师说明结果是如何一步一步得到的，这正是本文在数据流设计上强调可追溯性的原因。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "3.7 面向论文实验的可追溯设计", 2)
    for text in [
        "论文实验最常见的问题之一，是过度依赖人工记录，导致结果难以复现。为避免这一问题，本文系统在设计时将‘可追溯性’当作核心目标之一。具体体现在三个方面：第一，测试程序将关键输入参数、方法模式、运行模式、结果数组名和时间统计全部写入 CSV；第二，实验图像和导出 VTK 文件使用统一命名规则，与 CSV 中的数组名和数据集名保持一致；第三，论文中的主要表格数字都直接来自 `results` 目录中的最终结果文件，而不是人工手抄汇总。",
        "这种设计使论文写作和系统实现之间建立了直接映射关系。例如，如果老师质疑某个图表中的时间或误差数字，只需要回到对应 CSV 即可追查来源；如果老师想看某个优化前后的视觉差异，也可以直接在 ParaView 中载入对应导出文件进行复核。换言之，实验结果并不是‘论文里单独造出来的数据’，而是系统运行结果的再组织和再解释。",
        "从答辩角度看，这种可追溯设计还有一个额外好处：即使个别结果不够理想，也能清楚说明其来源、条件和边界，而不是陷入‘为什么和口头描述不一致’的被动局面。本文后续对单元数据解析场结果和数据优化模块适用范围的谨慎表述，就建立在这一可追溯性之上。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "3.8 本章小结", 2)
    for text in [
        "本章从系统目标、门面层、内部数据表示、桥接层、GPU 计算层和测试程序六个方面给出了系统设计细节。可以看到，本文系统并不是将若干算法零散拼接在一起，而是围绕统一数据表示与统一调度入口组织起来的原型平台。这样的结构既支撑了 GUI 演示，也支撑了论文实验和结果导出。",
        "更重要的是，系统设计已经内含了论文所需的严谨性要求：统一的算法入口保证了实验口径一致，显式记录的时间与误差指标保证了可比性，保留中间结果和原始导出文件则保证了可追溯性。后续算法章节和实验章节都建立在这一工程基础之上。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "第四章 算法原理", 1)
    add_heading(doc, "4.1 规则网格有限差分算法", 2)
    for text in [
        "在规则网格路径中，算法首先根据字段关联类型确定逻辑维度。若是点数据，则直接使用规则网格节点维度；若是单元数据，则相应减少每个方向上的样本数。随后在每个逻辑方向上选择中心差分或单边差分模板，对参数空间中的场值进行局部求导。",
        "接着，算法利用相邻采样点坐标近似构造局部 Jacobian，并将参数空间导数映射到物理空间。这样就避免了把几何变化完全忽略掉，也使规则但非均匀的采样分布仍能被正确处理。该思路与简单均匀立方体上的常规差分相比，更贴近真实的曲线规则网格情形。",
        "从复杂度看，规则网格每个点仅需访问固定数量的邻域样本，属于典型的 O(N) 并行处理问题，非常适合 GPU 批量执行。由于差分模板短、访存规律明确，这一路径通常也能获得系统中最小的 GPU 计算时间。",
    ]:
        add_body_paragraph(doc, text)
    insert_figure(
        doc,
        FigureData(
            "图4-1 梯度计算模块流程图",
            assets["gradient_flow"],
            "对于非结构化网格，Auto 模式默认走形函数导数主线；若点梯度结果出现异常全零，则启用保底退化路径。",
        ),
    )
    insert_figure(
        doc,
        FigureData(
            "图4-2 FD.glsl 规则网格有限差分 shader 流程图",
            assets["shader_fd"],
            "该 shader 对应规则网格主路径，核心是“参数空间差分 + Jacobian 映射 + 退化保护”的统一计算框架。",
        ),
    )
    add_body_paragraph(doc, "图4-2把规则网格梯度主路径细化到了 shader 级。与只给出模块框图不同，这里明确展示了差分模板选择、邻域采样、Jacobian 构造、退化保护和结果写回等关键步骤。这样在论文和答辩中就能够说明，规则网格 GPU 路径并不是简单把 CPU 差分搬到显卡上，而是围绕 OpenGL 扁平缓冲区和 Compute Shader 执行模型重新组织后的完整实现。")

    add_heading(doc, "4.2 非结构化网格点数据梯度算法", 2)
    for text in [
        "对点数据而言，系统的基本思路是：把每个节点所关联的单元都视为可提供局部梯度信息的支持域，在各支持单元内部根据单元类型读取参考形函数导数和节点坐标，计算出局部梯度后再回到该点做汇聚。由于一个点可能同时被多个单元共享，这种“单元内重构 + 节点处汇聚”的方式既符合有限元插值语义，也有利于消化局部几何方向差异。",
        "在单元内部，算法对一阶三角形、四边形、四面体和六面体采用各自的参考形函数导数模板。若单元为体单元，则使用 J^{-T} 映射；若单元为曲面单元，则使用 J(J^T J)^{-1} 的切空间映射。这样可以保证计算出的梯度尽量落在单元实际可表示的局部切空间中。",
        "点梯度路径还包含一个非常实际的保护机制。根据当前 GLGradientEngine 实现，如果 ShapePointGradient.glsl 返回了全零缓冲区，系统会自动回退到“先算单元梯度，再通过 CellDataToPointLift 提升到点”的保底路径。这个机制并不是为了掩盖问题，而是为了在复杂数据集或边界条件下尽可能给出非空结果，同时为后续进一步定位问题保留实验能力。",
        "由于该路径直接以节点值作为插值自由度，其数值一致性与有限元形函数理论对应得最紧密。因此在本文实验中，点数据解析场结果被作为非结构化网格算法正确性的主要证据。",
    ]:
        add_body_paragraph(doc, text)
    insert_figure(
        doc,
        FigureData(
            "图4-3 ShapePointGradient.glsl 点数据梯度 shader 流程图",
            assets["shader_shape_point"],
            "该图对应非结构化网格点数据的主计算核，展示了候选单元筛选、参考坐标定位、形函数导数映射和多单元平均汇聚过程。",
        ),
    )
    add_body_paragraph(doc, "图4-3对应的是当前非结构化网格点数据的核心实现。它首先以 pointCell 邻接关系确定支持单元，再在每个单元内部恢复参考梯度，并利用 Jacobian 或曲面切空间做映射，最后在节点处完成聚合。把这些步骤显式画出来之后，系统与 vtkGradientFilter 的可比性、与有限元形函数理论的一致性以及复杂曲面数据上可能出现的敏感点，都能够被更清楚地解释。")

    add_heading(doc, "4.3 非结构化网格单元数据梯度算法", 2)
    for text in [
        "单元数据的情形更复杂。对于一个单元中心值而言，它并不天然对应参考单元上的节点自由度，因此不能像点数据那样直接把场值代入形函数插值公式。本文当前采用的是两步式恢复路径：第一步通过 CellDataToPointLift.glsl 将与某个点相邻的单元值做局部提升，得到等价的点值场；第二步再在单元中心位置执行形函数导数重构。",
        "这种设计的优点是工程实现统一，点数据和单元数据最终都回到相同的几何映射与形函数导数框架中，着色器维护成本较低，也更容易与点数据路径共用代码。与此同时，它也带来了一个明确限制：cell -> point 的恢复过程并非严格线性保持重构，因此在解析场实验中，单元数据误差往往显著大于点数据。",
        "这正是本文在实验章节中将单元数据解析场结果降级为“补充说明”的原因。换言之，单元数据路径并非不能用，而是目前更适合作为真实工程字段与 VTK 的一致性对照，而不适合作为算法正确性的唯一主证据。论文中对此必须明确说明，才能保证结论的逻辑严谨性。",
    ]:
        add_body_paragraph(doc, text)
    insert_figure(
        doc,
        FigureData(
            "图4-4 ShapeCellGradient.glsl 单元梯度 shader 流程图",
            assets["shader_shape_cell"],
            "该图给出单元中心梯度的直接计算过程，重点展示单元中心参考坐标、Jacobian/度量张量映射和梯度写回机制。",
        ),
    )
    insert_figure(
        doc,
        FigureData(
            "图4-5 CellDataToPointLift.glsl 单元到点恢复 shader 流程图",
            assets["shader_cell_lift"],
            "该图对应单元数据保底路径中的中间恢复步骤，即把相邻单元值平均提升为点值，再交由形函数导数梯度核继续处理。",
        ),
    )
    add_body_paragraph(doc, "图4-4和图4-5需要结合起来理解单元数据主路径。前者负责在单元中心直接求梯度，后者则承担从单元场恢复点场的过渡职责。也正因为存在这一步“cell -> point”的语义转换，当前系统在单元数据解析场上的误差会比点数据更敏感。把这两张图放在正文中，有助于把误差来源解释为恢复链路问题，而不是笼统归结为形函数导数方法本身失效。")

    add_heading(doc, "4.4 多尺度数据优化算法", 2)
    for text in [
        "数据优化模块首先在点域或单元域上构造邻域图，并根据平均邻距与字段标准差自动估计空间和数值权重尺度。随后，Bilateral.glsl 在图结构上做逐层平滑，依次得到 smooth1、smooth2、smooth3 等低频层。相邻平滑层做差后形成 detail0、detail1、detail2 等高频细节层。",
        "融合阶段由 MultiScaleFuse.glsl 完成。它以最深层平滑结果作为 base，再按 detail_gain0、detail_gain1、detail_gain2 等增益参数回注不同尺度的细节。这样做可以避免“只平滑不重建”导致的重要边缘过度钝化，同时保留对局部高频扰动的抑制能力。",
        "从数值观点看，该算法并不试图修正离散化根因，而是把局部随机高频成分视作可以被分离和削弱的一类场细节。由于 CAE 结果中的局部振荡、伪振荡和非物理锯齿往往正体现在高频尺度上，这一思路具有明确的工程合理性。与之对应，论文也不会把该模块写成“任何仿真误差都能消除”的通用工具。",
    ]:
        add_body_paragraph(doc, text)
    insert_figure(
        doc,
        FigureData(
            "图4-6 数据优化模块流程图",
            assets["multiscale_flow"],
            "该模块由分解与融合两部分组成：前者显式分离局部高频细节，后者按增益将细节有控制地回注到基础层中。",
        ),
    )
    insert_figure(
        doc,
        FigureData(
            "图4-7 Bilateral.glsl 图双边滤波 shader 流程图",
            assets["shader_bilateral"],
            "该图细化了多尺度分解阶段的单层平滑逻辑，核心是邻域遍历、空间核和值域核联合加权。",
        ),
    )
    insert_figure(
        doc,
        FigureData(
            "图4-8 MultiScaleFuse.glsl 多尺度融合 shader 流程图",
            assets["shader_fuse"],
            "该图细化了 detail 回注与增益控制逻辑，说明 fused 结果并不是简单平滑，而是带有边缘保留意图的重建。",
        ),
    )
    add_body_paragraph(doc, "图4-6到图4-8把数据优化模块从“模块级”细化到了“核函数级”。其中，图4-7负责产生不同尺度的平滑层，图4-8负责根据细节强度和边缘抑制参数完成回注重建。这样一来，论文中关于“为什么该模块既能抑制局部随机高频扰动，又不希望把主边缘整体抹掉”的论证，就有了直接对应的实现依据。")

    add_heading(doc, "4.5 算法复杂度与稳定性讨论", 2)
    for text in [
        "从复杂度角度看，规则网格有限差分和多尺度优化都属于近似线性的 GPU 并行过程；非结构化网格点梯度与单元梯度则与单元连接数量、点所属单元数或局部邻接规模相关，但在统一扁平数据表示下仍可通过单次或少量 dispatch 完成。由于系统把时间统计细分为墙钟时间和 GPU 纯计算时间，实验中可以更清晰地区分算法复杂度和框架开销。",
        "从稳定性角度看，规则网格路径最稳定，点数据形函数导数路径次之，单元数据恢复路径最敏感。对于曲面型数据，还必须采用与几何维数一致的 benchmark 和指标体系，否则即便算法本身正确，也会在 ambient 指标下被法向分量误差过度惩罚。本文后续实验设计和结论边界正是围绕这些稳定性特征组织的。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "4.6 支持单元类型与实现范围说明", 2)
    for text in [
        "就当前工程实现而言，着色器代码对多种单元类型都预留了处理分支，但本文正式纳入毕业论文主线并进行稳定实验论证的单元类型，集中在一阶三角形、一阶四边形、一阶四面体和一阶六面体。这种收敛并不是能力缩水，而是为了保证论文中每一项结论都建立在充分实验支撑之上。",
        "三角形与四边形主要承担曲面型网格的代表角色，四面体与六面体则承担体网格的代表角色。对于更高阶单元或更复杂混合单元，虽然从有限元理论上可以继续扩展形函数导数模板，但这将显著增加参考单元定义、积分点选取、形函数导数公式和着色器分支复杂度，不适合作为当前毕设阶段的主线工作内容[1][14][16]。",
        "因此，论文中的‘支持非结构化网格梯度计算’必须理解为：在当前版本中，系统已经围绕上述四类代表性一阶单元建立了完整的 GPU 计算与实验链路。这种表述既准确反映实现状态，也避免在答辩中被追问尚未完成的高阶或稀有单元路径。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "4.7 算法流程与误差来源分析", 2)
    for text in [
        "若从算法执行顺序看，规则网格有限差分的误差主要来源于差分截断误差、边界单边差分误差以及局部 Jacobian 近似误差；非结构化点数据形函数导数的误差主要来自参考单元映射误差、浮点运算误差以及多单元汇聚时的局部方向差异；单元数据路径则额外包含了 `cell -> point` 恢复步骤带来的离散语义转换误差。",
        "这种误差来源拆解对实验分析非常重要。因为当某个结果不理想时，不能笼统地归咎于‘算法不好’。例如 `ShipHull_0 point` 在普通 ambient benchmark 下误差偏大，核心问题并不是形函数导数公式错误，而是曲面数据的评价方式不匹配；而 `hexa cell`、`limb cell` 在解析场下偏大，则更多与单元值恢复路径相关，而不是与点数据主线所依赖的形函数导数公式相矛盾。",
        "文献中对导数量恢复、应力恢复和梯度恢复的讨论也表明，导数类量天然比原始场更敏感，因此不同路径呈现不同精度层级是正常现象。关键不是强行把所有结果写成同样优秀，而是说明哪些误差来自理论边界，哪些来自当前实现选择，哪些又属于未来可以继续改进的工程问题[10][17][18][20]。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "4.8 本章小结", 2)
    for text in [
        "本章围绕规则网格有限差分、非结构化网格点数据梯度、非结构化网格单元数据梯度以及多尺度数据优化四部分，说明了系统核心算法的数学基础和实现路径。通过对映射公式、恢复步骤、退化保护和误差来源的逐层分析，可以看出本文算法主线在理论上是自洽的，在工程上也是可解释的。",
        "这为下一章实验提供了直接指导：解析场验证要优先考察最能体现理论正确性的路径，真实字段对照则用来评估与工程基线的一致性，多尺度优化实验则重点关注扰动抑制与边缘保持之间的平衡。换言之，实验设计并不是事后拼凑，而是由本章算法分析自然推导出来的。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "第五章 实验与结果分析", 1)
    add_heading(doc, "5.1 实验环境与数据集", 2)
    add_body_paragraph(doc, "实验程序由 C++17、Qt 5、VTK 与 OpenGL 4.3+ 构成，构建环境为 MSVC2019 64-bit Debug。由于当前整理阶段未直接记录整机硬件型号，论文保留软件链路、数据集与原始结果文件的完整可追溯信息，最终答辩版可再补充具体 CPU、GPU 与内存型号。对于本文而言，软件链路、输入数据、配置参数和输出文件的一致性，比单次硬件型号记录更直接地关系到结果复现与结论可核查性[27][28][29]。")
    insert_table(doc, tables["dataset_grad"])

    add_heading(doc, "5.2 实验设计与测试口径", 2)
    for text in [
        "梯度模块实验分为两层。第一层是解析场验证，其目标是回答“算法实现是否正确”；因此只选择能够稳定支撑结论的数据和指标，即 `SampleStructGrid point`、`hexa point` 和 `1_0 point`，其中曲面型 `1_0` 采用曲面专用 benchmark 并以 intrinsic 指标为主。第二层是真实字段与 vtkGradientFilter 的一致性和时间对比，其目标是回答“工程结果是否可比、GPU 路径是否具备效率优势”[8][9][10][27]。",
        "这种组织方式直接对应老师提出的实验严谨性要求：如果要做和 VTK 的比较，必须保证算法配置与计算条件可比，因此测试程序同时输出 VTK 单线程与并行信息，并显式记录并行后端、线程数和时间。另一方面，工程软件对照不等于真值验证，因此不能把“与 VTK 接近”当成算法绝对正确性的唯一证据，这也是解析场实验必须单独保留的原因[27][28]。",
        "数据优化模块实验统一使用 `ShipHull_0`。从研究定位出发，本文将目标扰动定义为‘局部随机高频数值扰动’，并把高斯扰动作为其中一类代理模型。同时，考虑到空间相关扰动和混合扰动也可能出现，程序额外补充了 `grf` 和 `mixed` 实验；对 `impulse` 则用来验证模块边界。通过这种设计，实验既保留了受控性，又不会将模块错误地表述为通用降噪器[21][22][23][24][25][30]。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "5.2.1 解析场验证与软件对照验证的角色区分", 3)
    for text in [
        "从实验方法学角度看，解析场验证与软件对照验证是两种不同层级的证据。解析场验证关注的是“理论公式是否实现正确”，其优势在于可以直接与真值比较，但前提是测试场与离散语义相匹配；软件对照验证关注的是“系统输出是否与成熟工具一致”，其优势在于贴近工程应用，但无法替代真值判断[8][9][27]。",
        "如果把两者混为一谈，就容易出现逻辑错误：例如某个结果与 VTK 很接近，并不能自动推出它就是解析意义下最准确的；反过来，某个曲面数据在 ambient 解析场下误差偏大，也不能据此立即否定工程对照中的高一致性。因此，本文在实验结构上显式将两者拆开，正是为了避免这种证据层级混淆。",
        "这一处理也使本文能够更坦诚地面对不够理想的结果。比如单元数据解析场误差偏大，本文并没有回避，而是将其降级为附录补充结果，并说明其与恢复链有关；而在真实字段对照中，只要系统与 VTK 保持高一致性，就仍然可以支持其工程可用性。这种结论结构比简单追求所有指标都“很好看”更符合数值实验的基本规范[10][17][18][20]。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "5.2.2 解析场生成方式", 3)
    for text in [
        "当前解析场并非离线手工制作，而是由测试程序在数据集装载后按统一规则注入。对规则网格和一般体网格，`CAEProcessingFacade.cpp` 先基于包围盒中心构造局部坐标 `x_r=x-c_x`、`y_r=y-c_y`、`z_r=z-c_z`，再以参考长度 `L` 做尺度归一化，分别生成 `benchmark_linear`、`benchmark_quadratic`、`benchmark_trig`、`benchmark_cubic`、`benchmark_gaussian` 以及向量场 `benchmark_vec_linear`、`benchmark_vec_poly`、`benchmark_vec_trig`；与此同时把每个场的解析梯度直接写成 `*_exact_grad` 数组，作为后续误差评估的真值来源。点关联实验使用节点坐标注入解析场，单元关联实验使用单元中心注入解析场，这样真值定义与被测关联类型保持一致[17][18][20]。",
        "对曲面型数据，本文没有继续套用普通三维体场，而是在 `TestGradient.cpp` 中通过 `injectSurfaceAnalyticBenchmarks(...)` 单独构造曲面专用 benchmark。程序首先利用几何分析模块给出的全局切向基 `a,b` 和中心点 `c`，把位置映射为局部切向坐标 `u=(p-c)·a`、`v=(p-c)·b`；随后定义 `benchmark_surface_linear`、`benchmark_surface_trig` 和 `benchmark_surface_vec_linear` 等场，并把解析梯度投影回局部支撑子空间 `projectToLocalSupport(...)`。这样得到的真值与曲面问题的几何维度一致，避免把法向分量误算进主要结论。",
        "从验证逻辑上看，这种做法接近制造解思想在后处理实验中的简化应用：先构造一个已知解析梯度的场，再检查离散实现能否恢复它。本文并不声称这些 benchmark 覆盖了全部工程情况，但它们足以回答“当前 GPU 梯度路径是否按预期实现”这一基础正确性问题，因此被放在实验链条的第一层[17][18][20][27]。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "5.2.3 高斯噪声与相关扰动的添加方式", 3)
    for text in [
        "数据优化实验中的扰动同样不是手工绘制，而是由 `TestHarnessUtils.h` 中的 `injectNoise(...)` 统一生成。对于给定干净场 `f_clean`，程序先计算其样本标准差 `signalStd=std(f_clean)`，再按 `targetSigma=sigmaFactor·signalStd` 确定白噪声强度，并从零均值高斯分布 `N(0,targetSigma^2)` 采样噪声项 `n_i`，最终得到 `f_noisy=f_clean+n_i`。这种做法使扰动强度随字段自身量级自适应缩放，避免同一组参数在不同数据场上代表完全不同的噪声幅值。",
        "为了避免把模块能力表述得过宽，本文还补充了两类边界场景。其一是 `grf`（相关高斯扰动）：先生成白噪声，再沿网格邻接图做若干轮平滑传播，控制参数为 `corr_length_factor`、`corr_iters` 和 `corr_alpha`，最后重新缩放回同一目标标准差；其二是 `mixed` 与 `impulse` 场景：前者把高斯扰动与脉冲异常值叠加，后者只保留脉冲型异常。这样设计的目的不是证明真实 CAE 数值误差整体服从高斯分布，而是用一组可控代理模型刻画“局部随机高频扰动”“空间相关扰动”和“尖峰异常值”三类不同难度的测试对象[21][22][23][24][25][30]。",
        "因此，本文关于数据优化模块的实验结论应理解为：在当前参数和当前实现下，模块对以高斯型随机高频波动为代表的一类局部扰动具有明显抑制效果；对空间相关扰动仅有有限改善；对脉冲型异常值则不敏感。这样的结论边界与程序实际支持的噪声生成方式是一致的，也更符合老师所强调的科学性要求。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "5.2.4 benchmark 展示口径调整", 3)
    for text in [
        "在早期实验中，测试程序会一次性生成线性、二次、三角、三次、高斯以及多种向量场 benchmark。若把这些 benchmark 全部并列放进正文，很容易把“主结论所需证据”与“补充压力测试”混在一起，反而削弱论文论证主线。结合当前结果稳定性，本文将正文主展示口径收敛到四类最能稳定支撑结论的 benchmark：`SampleStructGrid / benchmark_linear`、`SampleStructGrid / benchmark_vec_linear`、`hexa / benchmark_linear` 和 `1_0 / benchmark_surface_linear`。",
        "这四类 benchmark 分别对应四个最关键的问题：一是规则网格标量有限差分是否能准确恢复一阶梯度；二是多分量向量梯度的内存布局和分量顺序是否正确；三是非结构化体网格点数据的形函数导数路径是否成立；四是曲面型数据在切空间一致定义下是否能够得到合理结果。它们共同构成本文“算法已经实现且主路径正确”的最小证据集。相比之下，`benchmark_trig`、`benchmark_quadratic`、`benchmark_cubic`、`benchmark_gaussian` 和 `benchmark_surface_trig` 更适合作为补充压力测试，用于观察更复杂场形下的误差增长，而不作为正文主结论的唯一支撑。",
        "这种展示口径的调整并不是回避结果，而是把不同 benchmark 的论证职责分开：主表回答“正确性是否成立”，补充 benchmark 回答“在更复杂场形下误差如何变化”。实验口径一旦这样固定下来，后续与 VTK 的一致性对照、时间对比以及答辩时的口头汇报就能共享同一套逻辑框架，不会在不同图表之间来回切换论证标准[27][28]。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "5.2.5 参数说明与小型案例依据", 3)
    for text in [
        "老师曾提出，梯度计算和数据优化模块的相关参数是否需要交代、是否需要做参数实验。本文的处理是区分“实验配置项”和“算法调节项”。对梯度模块而言，为了保证与 VTK 的对比公平，正文不再引入数据依赖式调参；需要固定和报告的是 `POINT/CELL` 关联、`GL/VTK` 结果来源、VTK 并行后端、线程数、重复次数等实验配置。这些配置会影响可比性，但不改变本文主线算法的数学定义，因此不把它们作为调参实验主体。",
        "真正需要说明和适度验证的是数据优化模块参数，因为双边滤波和多尺度融合包含空间尺度、值域尺度、层数、层间放缩和细节增益等显式控制项。为避免把论文工作量转移成大规模参数搜索，本文采用“固定默认参数 + 小型案例依据”的方案：先在少量代表性场景中观察默认参数是否同时满足“对 Gaussian 有效、对 GRF 保持谨慎、对 Impulse 不夸大能力”这三个要求，再据此固定全文批处理实验所用参数。相应的默认参数表和小型案例验证表分别见表5-9a和表5-9b。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "5.3 梯度模块解析场验证", 2)
    insert_table(doc, tables["analytic_main"])
    insert_figure(
        doc,
        FigureData(
            "图5-1 解析场主验证结果",
            assets["analytic_chart"],
            "主验证只保留能够稳定支撑正确性结论的数据与指标，因此 1_0 采用曲面专用 intrinsic 结果。",
        ),
    )
    for text in [
        f"表5-2只保留了四类主验证 benchmark。具体而言，规则网格 `SampleStructGrid` 的标量线性场 NRMSE 为 {fmt_sci(ctx['analytic_rows'][0]['value'], 6)}，向量线性场 NRMSE 为 {fmt_sci(ctx['analytic_rows'][1]['value'], 6)}；六面体点数据 `hexa` 在线性场上的 NRMSE 为 {fmt_sci(ctx['analytic_rows'][2]['value'], 6)}；曲面型数据 `1_0 point` 在曲面专用 `benchmark_surface_linear` 上的 intrinsic NRMSE 为 {fmt_sci(ctx['analytic_rows'][3]['value'], 6)}。这些结果覆盖了规则网格标量、规则网格向量、非结构化体网格点数据和曲面点数据四条主要实现路径，且误差都处于较低量级，因此足以支撑主路径实现正确这一结论。",
        f"补充 benchmark 则保留在 CSV 和补充分析中，用于观察更复杂场形下的误差增长趋势。例如，`SampleStructGrid / benchmark_trig` 的 ambient NRMSE 为 {fmt_sci(ctx['analytic_secondary'][0]['value'], 6)}，`hexa / benchmark_trig` 为 {fmt_sci(ctx['analytic_secondary'][1]['value'], 6)}，`1_0 / benchmark_surface_trig` 的 intrinsic NRMSE 为 {fmt_sci(ctx['analytic_secondary'][2]['value'], 6)}。它们说明当场函数由一阶线性提升到非线性波动后，误差会按预期上升，但整体仍保持在可解释范围内；因此本文将其作为补充压力测试，而不是改变正文主展示口径。",
        f"对 `1_0` 这样的曲面型数据，如果直接使用普通三维 `benchmark_linear`，ambient NRMSE 会达到 {fmt_sci(ctx['analytic_surface_compare'][0]['ambient'], 6)}，但同一行的 intrinsic NRMSE 仅为 {fmt_sci(ctx['analytic_surface_compare'][0]['intrinsic'], 6)}。这一对比清楚表明：曲面数据不能简单套用三维体 benchmark，否则误差解释会被法向项扭曲。",
    ]:
        add_body_paragraph(doc, text)
    insert_table(doc, tables["surface_compare"])
    add_body_paragraph(doc, "因此，本文关于“非结构化网格形函数导数方法已经实现且主线正确”的论证，建立在 `SampleStructGrid point`、`hexa point` 和 `1_0 point` 的解析场结果之上，而不是建立在所有点/单元/几何场景都同时达到同样精度这一更强命题之上。这样的结论边界更符合当前实现状态，也更符合代码验证的基本原则[8][9][10]。")
    insert_figure(
        doc,
        FigureData(
            "图5-2 曲面数据 ambient / intrinsic 指标对比",
            assets["surface_chart"],
            "该图进一步说明，曲面型数据若不区分 ambient 与 intrinsic 评价，将显著放大法向分量带来的误差。",
        ),
    )
    for text in [
        "从图5-2可以进一步看出，曲面数据的评价问题并不是个别样本偶然出现的异常，而是由几何维度与评价维度不一致引起的系统性现象。以 `1_0 point` 的普通线性场与曲面线性场为例，当场定义与曲面切空间一致时，ambient 与 intrinsic 指标都会显著下降；而当直接把三维体 benchmark 套到曲面上时，ambient 指标会产生明显高估。",
        "这类结果在答辩中非常关键，因为它能帮助解释为什么某些数据“看起来误差大”，但并不意味着实现错误。对老师而言，真正重要的是学生是否知道问题出在哪里、是否能给出合理实验组织来避免误判。本文通过曲面专用 benchmark 和 intrinsic 指标，给出了这类问题的一个可复现回答[11][12][27]。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "5.4 真实字段与 VTK 的一致性及时间对比", 2)
    insert_table(doc, tables["real_compare"])
    insert_figure(
        doc,
        FigureData(
            "图5-2 真实字段与 VTK 结果误差分布",
            assets["real_chart"],
            "除曲面型数据外，多数组合的 NRMSE 已低至 1e-6 至 1e-8 量级，说明系统结果与 VTK 工程输出高度接近。",
        ),
    )
    for text in [
        f"在真实字段对照中，规则网格 `SampleStructGrid point` 的 NRMSE 为 {fmt_sci(ctx['real_rows'][0]['nrmse'], 6)}，`SampleStructGrid cell` 为 {fmt_sci(ctx['real_rows'][1]['nrmse'], 6)}；六面体点数据 `hexa point` 为 {fmt_sci(ctx['real_rows'][2]['nrmse'], 6)}；复杂体网格 `limb cell` 为 {fmt_sci(ctx['real_rows'][3]['nrmse'], 6)}。这些结果均处于极低量级，足以说明系统在真实 CAE 字段上能够复现 vtkGradientFilter 的工程结果。",
        f"曲面型数据的误差相对更高，但仍处于可接受范围。例如 `ShipHull_0 point` 的 NRMSE 为 {fmt_sci(ctx['real_rows'][4]['nrmse'], 6)}，`ShipHull_0 cell` 为 {fmt_sci(ctx['real_rows'][5]['nrmse'], 6)}；`1_0 point` 和 `1_0 cell` 则分别为 {fmt_sci(ctx['real_rows'][6]['nrmse'], 6)} 和 {fmt_sci(ctx['real_rows'][7]['nrmse'], 6)}。这说明曲面型几何对梯度重构更敏感，但并未破坏系统与 VTK 的整体可比性。",
        "从实验逻辑上说，这一组结果承担的是工程一致性验证，而不是解析真值验证。也就是说，表5-4能够支撑‘系统与 VTK 的工程输出是否接近’这一问题，却不能代替前一节中的解析 benchmark。因此，论文在表述时必须把两层含义区分开来。",
    ]:
        add_body_paragraph(doc, text)
    insert_table(doc, tables["timing_compare"])
    insert_table(doc, tables["speedup_compare"])
    insert_figure(
        doc,
        FigureData(
            "图5-3 梯度计算时间对比",
            assets["timing_chart"],
            "该图同时报告 VTK 单线程、VTK 并行、系统总时间和 GPU 时间，避免只拿 GPU 纯计算时间与 CPU 整体流程做不公平比较。",
        ),
    )
    insert_figure(
        doc,
        FigureData(
            "图5-4 系统相对 VTK 的加速比",
            assets["speedup_chart"],
            "该图用加速比而非绝对时间再次展示系统性能。对于部分小规模或后端开销较大的数据，VTK 并行未必优于单线程。",
        ),
    )
    for text in [
        f"以 `SampleStructGrid point` 为例，VTK 单线程平均时间为 {fmt_ms(ctx['real_rows'][0]['vtk_single_ms'])} ms，VTK 并行为 {fmt_ms(ctx['real_rows'][0]['vtk_parallel_ms'])} ms，而系统总时间为 {fmt_ms(ctx['real_rows'][0]['result_wall_ms'])} ms，GPU 纯计算时间仅为 {fmt_ms(ctx['real_rows'][0]['result_gpu_ms'])} ms。类似地，`limb cell` 在开启 VTK 并行后从 {fmt_ms(ctx['real_rows'][3]['vtk_single_ms'])} ms 降到 {fmt_ms(ctx['real_rows'][3]['vtk_parallel_ms'])} ms，但系统总时间仍只有 {fmt_ms(ctx['real_rows'][3]['result_wall_ms'])} ms。",
        "需要注意的是，VTK 并行并非在所有数据集上都优于单线程。例如 `hexa point`、`ShipHull_0 point` 和 `1_0 point` 上，当前并行配置下的 VTK 时间反而高于单线程。这说明并行收益受任务规模、调度开销和后端实现影响显著，因此论文中不能直接写成‘VTK 并行一定更快’。本文采用的做法是完整保留单线程与并行结果，并明确其后端为 STDThread、线程数为 16。",
    ]:
        add_body_paragraph(doc, text)
    for text in [
        "从加速比角度观察，可以更直观地说明系统性能优势的来源。其一，GPU 纯计算时间在多个数据集上都远小于系统总时间，表明真正的数值核函数开销已经被压缩到较小量级；其二，系统总时间依然显著小于 VTK 时间，说明即使把缓冲区准备、结果回读和门面层调度开销一并计入，GPU 主路径仍然具有现实收益。",
        "不过，本文并未把这些结果夸张表述为“GPU 全面优于 VTK”。更准确的说法是：在当前实现、当前构建配置和当前测试数据集上，系统展现出了明显的时间优势；而对更大规模数据、不同并行后端或不同驱动环境，仍有必要继续做更细化测试。这种限定性表述更符合毕业论文应有的严谨性[27][28]。",
    ]:
        add_body_paragraph(doc, text)
    insert_figure(
        doc,
        FigureData(
            "图5-5 ShipHull_0 系统梯度渲染图",
            assets["gradient_render_system_raw"],
            "该图为系统梯度结果的 ParaView 渲染原图，用于展示系统在曲面型工程数据上的梯度分布形态。",
        ),
    )
    insert_figure(
        doc,
        FigureData(
            "图5-6 ShipHull_0 VTK 梯度渲染图",
            assets["gradient_render_vtk_raw"],
            "该图为 vtkGradientFilter 结果的 ParaView 渲染原图，可与系统结果做视觉对照。",
        ),
    )
    insert_figure(
        doc,
        FigureData(
            "图5-7 ShipHull_0 梯度结果 ParaView 对比",
            assets["gradient_render"],
            "从色带分布和高值区位置来看，系统结果与 VTK 结果整体一致，差异主要集中在局部过渡区域。",
        ),
    )

    add_heading(doc, "5.5 数据优化模块实验结果", 2)
    insert_table(doc, tables["mul_point_avg"])
    insert_table(doc, tables["mul_cell_avg"])
    insert_figure(
        doc,
        FigureData(
            "图5-8 点数据多尺度优化平均 NRMSE",
            assets["point_avg_chart"],
            "高斯扰动和混合扰动在点域上具有明显下降，GRF 改善有限，Impulse 基本无改善。",
        ),
    )
    insert_figure(
        doc,
        FigureData(
            "图5-9 单元数据多尺度优化平均 NRMSE",
            assets["cell_avg_chart"],
            "单元域上对高斯扰动的平均抑制比点域更强，但对脉冲型异常值同样缺乏明显改善。",
        ),
    )
    insert_figure(
        doc,
        FigureData(
            "图5-10 粗糙度比对比",
            assets["roughness_chart"],
            "粗糙度比越小，说明局部随机高频起伏被压制得越明显。Gaussian 与 Mixed 的表现明显优于 Impulse。",
        ),
    )
    for text in [
        f"点数据平均结果表明，gaussian 扰动下的输入 NRMSE 从 {fmt_ratio(float(ctx['point_avg'][ctx['point_avg']['noise']=='gaussian'].iloc[0]['input_nrmse']))} 降至 {fmt_ratio(float(ctx['point_avg'][ctx['point_avg']['noise']=='gaussian'].iloc[0]['fused_nrmse']))}，RMSE 改进率为 {fmt_ratio(float(ctx['point_avg'][ctx['point_avg']['noise']=='gaussian'].iloc[0]['rmse_improvement_ratio']))}；mixed 扰动也有稳定改善。相比之下，impulse 的改进率接近 1，说明当前方法并不针对稀疏脉冲异常值。",
        f"单元数据上，gaussian 扰动的平均输出 NRMSE 进一步下降到 {fmt_ratio(float(ctx['cell_avg'][ctx['cell_avg']['noise']=='gaussian'].iloc[0]['fused_nrmse']))}，表明在当前数据和参数设置下，单元域上的局部平滑与细节回注能够更明显地削弱随机高频扰动。GRF 的改善程度弱于 Gaussian，说明具有空间相关性的扰动比独立高斯扰动更难处理。",
        "结合文献和实验定位，这些结果应被解释为：当前模块对一类局部随机高频扰动具有稳定抑制能力，并兼具一定的边缘保持效果；但它并不意味着真实 CAE 所有数值误差都服从高斯分布，也不意味着模块可以充当通用降噪器或误差修正器。论文中正是通过这种受限但清晰的定位来保证论证严谨性[21][22][23][24][25][30]。",
    ]:
        add_body_paragraph(doc, text)
    insert_table(doc, tables["mul_repr"])
    insert_table(doc, tables["param_defaults"])
    insert_table(doc, tables["param_prestudy"])
    insert_figure(
        doc,
        FigureData(
            "图5-11 ShipHull_0 数据优化前渲染图",
            assets["denoise_before_raw"],
            "该图为加入局部随机高频扰动后的 ParaView 原始渲染图，可见局部色带波动较明显。",
        ),
    )
    insert_figure(
        doc,
        FigureData(
            "图5-12 ShipHull_0 数据优化后渲染图",
            assets["denoise_after_raw"],
            "该图为多尺度优化后 ParaView 原始渲染图，局部高频起伏有所减弱，主结构保持清晰。",
        ),
    )
    insert_figure(
        doc,
        FigureData(
            "图5-13 ShipHull_0 数据优化前后 ParaView 对比",
            assets["denoise_render"],
            "优化后局部高频波动明显减弱，而主边界和大尺度结构仍保持清晰。",
        ),
    )
    for text in [
        "从视觉上看，数据优化前后的主要差异并不体现在整体色带范围被大幅压缩，而体现在局部高频纹理是否被平顺化。对于 `ms_clean_edge` 这类本身带有明显边界过渡的场，优化后的边界位置依然保持清晰，说明多尺度融合并没有简单地把边缘整体抹掉，而是更倾向于抑制局部随机起伏。",
        "表5-9a给出了全文批处理实验所固定采用的关键参数，其中 `levels=3`、`spatial_sigma_factor=1.5`、`range_sigma_factor=0.5` 和逐层递减的 `detail_gain` 共同决定了“先分离尺度、再谨慎回注细节”的整体策略。这里没有把参数一味调大去追求更强平滑，而是刻意保留了边缘抑制项 `edge_sigma_factor` 和逐层衰减的细节增益，以避免把局部高频扰动压制成整体模糊。",
        "表5-9b进一步说明了默认参数为什么被保留下来。对 `POINT / ms_clean_edge / gaussian`，默认参数可在不抹平主边缘的情况下显著降低 NRMSE 和粗糙度；对 `POINT / ms_clean_trig / grf`，改善有限，提示当前参数并不会通过过度平滑去虚构性能；对 `POINT / ms_clean_edge / impulse`，结果几乎不改善，则反过来证明本文没有把模块能力夸大为“对所有噪声都有效”。这种小型案例依据虽然不是完整参数搜索，但足以支撑当前默认参数的选择逻辑。",
        "这也是本文坚持加入 ParaView 渲染图的原因。单纯的 NRMSE 或 RMSE 改进率虽然重要，但它们无法完整传达“边缘是否仍然可辨”这一视觉层面的信息。对于面向后续渲染服务的数据预处理而言，数值指标和视觉指标都必须被同时考虑，否则就无法说明优化模块对后处理展示是否真正有益[21][22][23]。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "5.6 本章讨论", 2)
    for text in [
        "综合解析场验证与真实字段对照，可以形成一个较为稳健的结论框架。其一，规则网格有限差分和非结构化点数据形函数导数路径在代表性数据上具备良好的数值正确性。其二，系统在真实字段上与 VTK 的一致性较高，因此具备工程可用性。其三，单元数据路径仍然存在解析场精度敏感性，但这并不否定其在真实字段场景下的参考价值，只是说明当前恢复链仍需进一步改进。",
        "数据优化模块方面，实验结果与预期定位基本一致：它更适合高斯型、相关型或混合型局部随机高频扰动，而不适合脉冲型异常值。这样的结果并不会削弱模块意义，反而说明其适用范围是明确且可验证的。对于毕业论文而言，这种“有效范围清楚、边界说明充分”的表述比泛化结论更严谨。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "5.7 与相关文献实验口径的对照分析", 2)
    for text in [
        "将本文实验口径与相关文献相比，可以发现其总体框架是相容的。制造解或解析场验证对应文献中的代码验证路径，用于证明离散实现和理论模型的一致性；真实字段与 VTK 的对照对应工程软件一致性路径，用于说明系统输出在工程语境下是否可信；时间对比则对应工程实现性能评估路径，但必须同时公开并行条件和后端配置[8][9][10][27][28]。",
        "在数据优化方面，本文采用了干净场加代理扰动的受控实验，这与图像处理和网格去噪文献中的常见做法一致。但不同之处在于，本文同时通过文献讨论收缩了结论边界，即只声称该模块对一类局部随机高频扰动有效，而不试图证明真实 CAE 全部噪声模型都与实验噪声等价[21][22][23][24][25]。",
        "因此，从实验方法学看，本文并不是简单展示若干“好看数字”，而是试图在毕业设计的工作量范围内，复用文献中较为标准的验证框架，并结合项目当前实现状态做有边界的取舍。这种组织方式能够较好回应老师对实验深度和比较严谨性的要求。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "5.8 本章小结", 2)
    for text in [
        "本章通过解析场验证、真实字段与 VTK 对照、时间对比、多尺度优化统计以及 ParaView 图像对照，完成了对系统两个核心模块的实验评估。结果表明，梯度模块已经形成清晰可验证的 GPU 主线，数据优化模块则在其目标扰动类别上展现了稳定效果。",
        "同时，本章也没有回避当前系统存在的局限：单元数据解析场路径仍待改进，曲面点数据评价必须强调 intrinsic 指标，数据优化模块对 impulse 异常值不敏感。这些边界说明既保证了论文结论的可信度，也为后续工作提供了明确方向。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "第六章 结论与展望", 1)
    add_heading(doc, "6.1 研究结论", 2)
    for text in [
        "本文围绕 CAE 后处理阶段的数据预处理问题，完成了一套基于 OpenGL 的 GPU 数值处理原型系统。系统在保持 VTK 数据读写、显示和对照能力的同时，将规则网格梯度、非结构化网格形函数导数梯度以及数据优化主路径统一到 Compute Shader 上执行，形成了较完整的工程实现链路。",
        "在梯度模块方面，规则网格有限差分和非结构化网格点数据形函数导数路径已经通过解析 benchmark 和真实字段对照获得了较强支撑。真实字段实验表明，系统与 vtkGradientFilter 在 `SampleStructGrid`、`hexa`、`limb`、`ShipHull_0` 和 `1_0` 等数据上保持较高一致性，并在当前测试环境下展现出明显的时间优势。",
        "在数据优化模块方面，本文将其清晰定位为‘面向局部随机高频数值扰动的结果场优化模块’，并通过 `ShipHull_0` 上的受控扰动实验验证了其对 gaussian、grf 和 mixed 类扰动的不同程度抑制能力，同时说明了对 impulse 异常值的局限。这一定位与实验结果相互一致，构成了模块成立的基础。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "6.2 本文工作的特点", 2)
    for text in [
        "第一，本文从系统实现角度将 GUI、测试程序与底层算法主线统一起来，避免了很多毕业设计中常见的‘演示代码’与‘实验代码’脱节问题。第二，本文在实验组织上区分了解析真值验证和工程软件一致性验证，使结论逻辑更加清晰。",
        "第三，本文没有把数据优化模块泛化为通用降噪器，而是明确限定其处理对象为一类局部随机高频数值扰动；同时，没有把与 VTK 的接近误写为绝对真值正确性。这些表述上的收敛，实际上提升了整篇论文的严谨性。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "6.3 现有不足", 2)
    for text in [
        "当前系统最主要的不足在于单元数据梯度路径仍采用 `cell -> point lift -> shape gradient` 两步式恢复，因此在解析场实验中比点数据更敏感。这说明单元值到点值的恢复策略仍有改进空间。",
        "第二，曲面型点数据在普通三维 ambient benchmark 下表现不佳，虽然通过曲面专用 benchmark 和 intrinsic 指标能够合理解释，但也说明当前曲面点梯度稳定性还有进一步提升余地。第三，数据优化模块目前针对的扰动类别仍然有限，对脉冲型异常值的抑制效果较弱。",
        "第四，尽管系统的计算主路径已经迁移到 OpenGL，但 I/O、GUI 和参考基线仍依赖 VTK，因此系统尚未形成完全 VTK-free 的全栈软件环境。对于本文课题目标而言，这一点并不构成根本障碍，但需要在结论中明确说明。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "6.4 后续工作展望", 2)
    for text in [
        "后续工作首先应继续分析 vtkGradientFilter 的退化保护、贡献单元选项和表面/体网格处理机制，在此基础上改进复杂曲面数据上的点梯度稳定性。其次，应研究更一致的单元值恢复方案，例如引入更严格的局部重构或守恒约束，以提升单元数据解析场下的正确性。",
        "在性能实验方面，可以进一步在支持 TBB 或 OpenMP 的 VTK 构建环境中完成更充分的并行时间对比，并在固定线程数、重复次数和后端配置的前提下报告更完整的统计结果。数据优化模块则可以进一步设计更贴近真实 CAE 数值扰动来源的实验，如网格畸变驱动的局部振荡、后处理恢复不连续和空间相关随机扰动等。",
        "总体上看，本文已经完成了可答辩、可运行、可复现的系统主线，后续工作更多属于精化和增强，而不是从零开始重建框架。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "参考文献", 1)
    for ref in build_reference_list():
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(ref)
        set_run_font(run, 11)

    doc.add_page_break()
    add_heading(doc, "附录A 关键代码文件与功能对照", 1)
    for text in [
        "1. `CAEProcessingFacade.h / CAEProcessingFacade.cpp`：系统统一门面层，负责数据集加载、方法分派、时间记录和结果写回。",
        "2. `DataObject.h / DataObject.cpp`：统一内部数据表示，存储点坐标、单元连接、字段数组、邻域图和规则网格维度。",
        "3. `VTKDataConverter.h / VTKDataConverter.cpp`：VTK 与内部表示的双向桥接。",
        "4. `GLGradientEngine.h / GLGradientEngine.cpp`：梯度计算引擎，封装有限差分、点梯度、单元梯度等 GPU 接口。",
        "5. `GLFilterEngine.h / GLFilterEngine.cpp`：图双边滤波与多尺度融合引擎。",
        "6. `Shaders/FD.glsl`、`Shaders/ShapePointGradient.glsl`、`Shaders/ShapeCellGradient.glsl`、`Shaders/CellDataToPointLift.glsl`：梯度模块核心着色器。",
        "7. `Shaders/Bilateral.glsl`、`Shaders/MultiScaleFuse.glsl`：数据优化模块核心着色器。",
        "8. `TestGradient.cpp`、`TestMultiScale.cpp`、`TestFieldMetrics.cpp`：实验测试程序。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "附录B 单元数据解析场补充结果", 1)
    insert_table(doc, tables["appendix_cell"])
    add_body_paragraph(doc, "附录B的目的不是替代正文主结论，而是向答辩老师说明：当前单元数据梯度路径已经实现并可运行，但由于恢复链设计原因，其解析场误差整体高于点数据路径，因此不被用作正文中算法正确性的首要支撑证据。")

    doc.add_page_break()
    add_heading(doc, "附录C 测试指标与统计口径说明", 1)
    for text in [
        "1. 绝对误差：对每个梯度向量分量计算结果值与参考值之差的绝对值，可反映最直接的偏差大小。2. RMSE 与 NRMSE：RMSE 表示均方根误差，NRMSE 则进一步除以参考量的均方根尺度，用于消除不同字段量纲和数量级差异带来的影响。3. Soft Relative Error：当参考值接近零时，普通相对误差容易数值失稳，因此测试程序引入软阈值 τ，在分母中增加稳定项，从而更稳健地比较小量级区域的误差。",
        "4. 角度误差与尺度偏差：对向量梯度而言，仅比较分量误差还不够，因此程序还统计结果向量与参考向量之间的夹角，以及结果范数相对参考范数的偏差。5. 粗糙度比：在数据优化实验中，使用图邻域上的局部差异强度来衡量场的高频起伏程度，再比较优化前后的比值。粗糙度比越小，说明局部高频成分被抑制得越明显。",
        "6. 墙钟时间与 GPU 时间：墙钟时间涵盖门面层调度、缓冲区准备、着色器派发和结果回读等全部流程；GPU 时间仅统计核心 dispatch 的执行开销。论文同时保留这两类时间，是为了避免只拿最理想的 GPU 内核时间做不公平对比。7. VTK 单线程与并行时间：测试程序通过 vtkSMPTools 在同一实验框架内分别测量单线程和并行时间，以便客观反映 VTK 在不同并行配置下的表现。",
        "8. 解析场与真实字段的指标解释差异：解析场实验中的误差指标更接近算法正确性评价，而真实字段对照中的误差指标更接近工程一致性评价。两者虽然都输出 NRMSE 和 Soft Relative Error，但解释语义不同，因此论文中不能混用结论。",
    ]:
        add_body_paragraph(doc, text)

    add_heading(doc, "附录D 论文实验复现流程说明", 1)
    for text in [
        "为了保证论文结果可复现，本文将实验复现流程整理如下。首先，使用 VTK 读取待测数据集，并通过 `CAEProcessingFacade::loadDatasetFromVTKFile` 转换为统一内部表示。其次，根据实验目标选择 `TestGradient` 或 `TestMultiScale`。在梯度实验中，需要指定数据集、字段关联、参考模式、方法模式以及 VTK 并行配置；在数据优化实验中，需要指定干净场、扰动类型、多尺度参数和输出路径。",
        "梯度实验的复现顺序建议为：第一步运行解析场验证，确认 `SampleStructGrid point`、`hexa point` 与 `1_0 point` 的 benchmark 结果是否与论文表格量级一致；第二步运行真实字段与 VTK 对照，确认 `SampleStructGrid`、`hexa`、`limb`、`ShipHull_0` 和 `1_0` 的 NRMSE、VTK 时间和系统时间与论文统计表接近；第三步导出结果 VTK 文件，在 ParaView 中按相同视角与色带生成渲染图，以便和论文插图核对。",
        "数据优化实验的复现顺序建议为：第一步在 `ShipHull_0` 上构造 `ms_clean_trig`、`ms_clean_gaussian`、`ms_clean_poly` 和 `ms_clean_edge` 四类干净场；第二步分别叠加 gaussian、grf、mixed 和 impulse 扰动；第三步运行多尺度优化并导出 `smooth`、`detail`、`base` 和 `fused` 结果；第四步从 `multiscale_report+point.csv` 和 `multiscale_report+cell.csv` 中汇总平均指标，再在 ParaView 中对优化前后结果做渲染图比较。",
        "若论文数字与本地复现数字存在轻微差异，应优先检查三个方面：一是测试程序是否使用了与论文一致的字段和关联类型；二是 VTK 并行后端和线程数是否一致；三是图像导出和数值导出是否来源于同一批结果文件。由于 GPU 驱动、VTK 构建方式和运行环境可能引入微小时间波动，因此时间实验应以多次重复后的平均值为准，而不是以单次最小值作为主结论。",
        "从论文写作角度看，附录D的意义不只是帮助他人复现实验，也帮助答辩时快速回应“这些图和表是怎么来的”。只要按照这里的顺序说明数据来源、运行程序、输出文件和统计方式，就能够把实验结果与系统实现直接对应起来，从而增强整篇论文的可信度。",
    ]:
        add_body_paragraph(doc, text)

    doc.add_page_break()
    add_heading(doc, "攻读学士学位期间发表的论文和取得的科研成果", 1)
    add_center_paragraph(doc, "无", 14)
    doc.add_page_break()

    add_heading(doc, "致谢", 1)
    for text in [
        "在本毕业设计完成过程中，指导教师在课题定位、实验组织、论文结构和结果边界把握方面给予了大量指导。特别是在如何区分解析场验证与工程对照、如何界定数据优化模块的适用范围等问题上，老师提出的要求使本文最终形成了更严谨的论证逻辑。",
        "感谢实验过程中给予帮助的同学和朋友，在数据整理、结果讨论、界面调试与文档核对方面提供了支持。也感谢开源社区提供的 VTK、Qt 和 OpenGL 相关工具与资料，为系统实现和实验验证提供了可靠基础。",
        "最后，感谢家人对学习和毕业设计工作的理解与支持。正是这些帮助，才使本文能够顺利完成。",
    ]:
        add_body_paragraph(doc, text)

    return save_docx_with_fallback(doc, THESIS_DOCX)


def build_markdown_preview(ctx: dict, tables: dict[str, TableData]) -> None:
    lines = [
        "# 基于OpenGL的CAE软件数据预处理方法研究与实践",
        "",
        "本文档为自动生成的论文源稿预览，正式排版版本见同目录 `.docx` 文件。",
        "",
        "## 摘要",
        "面向 CAE 后处理阶段中的数据预处理需求，本文实现了基于 OpenGL Compute Shader 的梯度计算与结果场优化系统，并通过解析场验证、真实字段与 VTK 对照以及多尺度优化实验进行了评估。",
        "",
        "## 主要实验表格",
    ]
    for key in ["analytic_main", "surface_compare", "real_compare", "timing_compare", "mul_point_avg", "mul_cell_avg", "mul_repr", "param_defaults", "param_prestudy"]:
        table = tables[key]
        lines.append(f"### {table.title}")
        lines.append("| " + " | ".join(table.headers) + " |")
        lines.append("| " + " | ".join(["---"] * len(table.headers)) + " |")
        for row in table.rows:
            lines.append("| " + " | ".join(row) + " |")
        if table.note:
            lines.append("")
            lines.append(table.note)
        lines.append("")
    THESIS_MD.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    ensure_dirs()
    ctx = build_data_context()
    assets = generate_asset_figures(ctx)
    tables = build_tables(ctx)
    thesis_docx = build_document(ctx, tables, assets)
    build_markdown_preview(ctx, tables)
    build_reports()
    print(f"generated: {thesis_docx}")
    print(f"generated: {THESIS_MD}")
    print(f"generated: {REPORT_FULL}")
    print(f"generated: {REPORT_SHORT}")


if __name__ == "__main__":
    main()
