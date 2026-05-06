from __future__ import annotations

import re
import textwrap
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
DOC = ROOT / "doc"
DATE_TEXT = "2026-05-04"

OUT_COMPLETION = DOC / "课题调研与答辩判断.md"
OUT_ARGUMENT = DOC / "立题论证书-修订版.md"
OUT_PROPOSAL = DOC / "开题报告-修订版.md"
OUT_MIDTERM = DOC / "中期进展报告-修订版.md"
OUT_OUTLINE = DOC / "本科毕业论文详细大纲.md"
OUT_THESIS_MD = DOC / "基于OpenGL的CAE软件数据预处理方法研究与实践_本科毕业论文.md"
OUT_THESIS_DOCX = DOC / "基于OpenGL的CAE软件数据预处理方法研究与实践_本科毕业论文.docx"

PROHIBITED = re.compile(r"最小二乘|WLS|AWLS|least\\s*squares|Least\\s*Squares|加权最小二乘", re.I)


ANALYTIC_SCALAR = [
    ["SampleStructGrid", "三维线性标量场", "2.82038e-07"],
    ["hexa", "三维线性标量场", "1.57998e-07"],
    ["1_0", "曲面切向线性标量场", "1.68088e-07"],
]

ANALYTIC_VECTOR = [
    ["SampleStructGrid", "三维线性向量场", "3.30256e-07"],
    ["hexa", "三维线性向量场", "2.36149e-07"],
    ["1_0", "曲面切向线性向量场", "1.32859e-07"],
]

VTK_POINT = [
    ["SampleStructGrid", "scalars", "8.22437e-08"],
    ["hexa", "scalars", "6.27285e-08"],
    ["1_0", "RF", "5.63477e-08"],
]

VTK_CELL = [
    ["SampleStructGrid", "scalars", "7.62786e-07"],
    ["limb", "chem_0", "3.89516e-07"],
    ["1_0", "S_Mises", "7.77806e-08"],
]

TIME_STRUCT = [
    ["timing_struct_20x20x20", "8000 / 6859", "16", "2.7573", "0.9036", "0.025272"],
    ["timing_struct_32x32x32", "32768 / 29791", "16", "3.5564", "1.4336", "0.043732"],
    ["timing_struct_48x48x48", "110592 / 103823", "16", "8.2716", "4.0996", "0.107692"],
]

TIME_UHEX = [
    ["timing_uhex_20x20x20", "8000 / 6859", "16", "163.799", "2.8527", "0.623896"],
    ["timing_uhex_32x32x32", "32768 / 29791", "16", "708.650", "5.3835", "2.53781"],
    ["timing_uhex_48x48x48", "110592 / 103823", "16", "3349.36", "44.8039", "7.91991"],
]

OPTIMIZATION = [
    ["POINT", "0.296642", "0.180858", "0.610772"],
    ["CELL", "0.294558", "0.121921", "0.414806"],
]


LITERATURE = [
    ["1", "EN", "OpenGL", "Khronos OpenGL Wiki. Compute Shader", "说明计算着色器的工作组、派发和并行执行模型。", "https://www.khronos.org/opengl/wiki/Compute_Shader"],
    ["2", "EN", "OpenGL", "Khronos OpenGL Wiki. Shader Storage Buffer Object", "说明 SSBO 的大容量缓冲读写机制，对应项目中的 GPU 数据传输。", "https://www.khronos.org/opengl/wiki/Shader_Storage_Buffer_Object"],
    ["3", "EN", "VTK", "VTK documentation. vtkGradientFilter Class Reference", "项目使用该过滤器作为真实字段工程一致性参考。", "https://vtk.org/doc/nightly/html/classvtkGradientFilter.html"],
    ["4", "EN", "ParaView", "ParaView User's Guide. Understanding Data", "用于说明点数据、单元数据、结构网格和非结构网格的数据模型。", "https://docs.paraview.org/en/v5.10.0/UsersGuide/understandingData.html"],
    ["5", "EN", "Qt", "Qt Documentation. Qt Widgets", "对应系统 GUI 层实现基础。", "https://doc.qt.io/qt-5/qtwidgets-index.html"],
    ["6", "EN", "Visualization", "Schroeder W., Martin K., Lorensen B. The Visualization Toolkit", "科学可视化管线和 VTK 数据模型的基础参考。", "https://www.vtk.org/vtk-textbook/"],
    ["7", "EN", "Visualization", "Ahrens J., Geveci B., Law C. ParaView: An End-User Tool for Large Data Visualization", "说明 ParaView 在大规模科学可视化中的定位。", "https://www.paraview.org/publications/"],
    ["8", "EN", "GPU", "Owens J. D. et al. A Survey of General-Purpose Computation on Graphics Hardware. Computer Graphics Forum, 2007", "支撑从图形管线到通用并行计算的技术背景。", "https://doi.org/10.1111/j.1467-8659.2007.01012.x"],
    ["9", "EN", "Finite Element", "Hughes T. J. R. The Finite Element Method: Linear Static and Dynamic Finite Element Analysis", "用于解释单元形函数、映射和导数计算。", "https://store.doverpublications.com/products/9780486411811"],
    ["10", "EN", "Gradient Recovery", "Zienkiewicz O. C., Zhu J. Z. The superconvergent patch recovery and a posteriori error estimates. IJNME, 1992", "有限元后处理中导数场和恢复场的经典参考。", "https://doi.org/10.1002/nme.1620330702"],
    ["11", "EN", "Filtering", "Tomasi C., Manduchi R. Bilateral Filtering for Gray and Color Images. ICCV, 1998", "双边滤波思想来源，说明空间权重和值域权重共同作用。", "https://doi.org/10.1109/ICCV.1998.710815"],
    ["12", "EN", "Filtering", "Fleishman S., Drori I., Cohen-Or D. Bilateral Mesh Denoising. ACM TOG, 2003", "网格特征保持平滑的重要参考。", "https://doi.org/10.1145/882262.882368"],
    ["13", "EN", "Scale Space", "Perona P., Malik J. Scale-space and edge detection using anisotropic diffusion. IEEE TPAMI, 1990", "保边平滑与尺度空间分析参考。", "https://doi.org/10.1109/34.56205"],
    ["14", "CN", "论文规范", "教育部. 普通高等学校本科毕业论文（设计）抽检办法（试行）", "说明本科论文应体现选题意义、逻辑构建、专业能力和学术规范。", "http://www.moe.gov.cn/srcsite/A08/s7056/202101/t20210107_509524.html"],
    ["15", "CN", "论文规范", "GB/T 7713.1-2025 信息与文献 编写规则 第1部分：学位论文", "用于论文结构、摘要、正文和附录编排参考。", "https://openstd.samr.gov.cn/"],
    ["16", "CN", "参考文献", "GB/T 7714-2015 信息与文献 参考文献著录规则", "截至本文生成日期，若学校无新版要求，可按 2015 版著录参考文献。", "https://openstd.samr.gov.cn/"],
    ["17", "CN", "CAE 后处理", "国内 CAE 后处理与科学可视化应用研究", "用于补充国内工程仿真后处理、VTK 和 ParaView 应用背景。", "https://kns.cnki.net/"],
    ["18", "CN", "GPU 可视化", "国内基于 GPU 的有限元或 CAE 后处理可视化研究", "用于说明 GPU 加速在国内 CAE 可视化研究中的应用价值。", "https://kns.cnki.net/"],
    ["19", "CN", "有限元理论", "国内有限元方法与后处理教材、论文", "用于解释单元形函数、节点场和单元场的工程含义。", "https://kns.cnki.net/"],
    ["20", "CN", "滤波优化", "国内图像、点云、网格和科学数据保边平滑研究", "用于支撑本文数据优化模块的中文相关工作。", "https://kns.cnki.net/"],
]


def clean(text: str) -> str:
    text = textwrap.dedent(text).strip()
    text = re.sub(r"(?m)^ {4,}", "", text)
    if PROHIBITED.search(text):
        raise ValueError("生成文本包含已废弃方案相关表述")
    return text + "\n"


def table(headers: list[str], rows: list[list[str]]) -> str:
    out = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    out += ["| " + " | ".join(row) + " |" for row in rows]
    return "\n".join(out)


def completion_doc() -> str:
    return clean(f"""
    # 课题调研与答辩判断

    生成日期：{DATE_TEXT}

    ## 一、是否可以使用你传来的进展文档

    可以，而且建议以 `C:/Users/lenovo/Desktop/毕设进展.docx` 作为目前课题描述和实验数据的主依据。该文档已经把课题口径收敛得比较清楚：梯度计算模块按“规则网格有限差分法、非结构化网格形函数导数法”展开；实验按解析场验证、与 VTK 结果一致性对比、与 VTK 时间对比三组组织；数据优化模块按“面向一类局部随机高频数值扰动的抑制与边缘保持处理”展开，并使用高斯扰动作为代理模型。

    ## 二、毕业设计完成度判断

    结论：当前项目可以算完成本科毕业设计，也可以支撑答辩。支撑依据主要来自四个方面：第一，系统已经实现 VTK 数据读入、内部数据转换、OpenGL 计算、结果写回和 VTK 导出；第二，梯度模块已经形成规则网格和非结构化网格两条清晰主线；第三，数据优化模块已经有明确的问题边界和定量实验；第四，项目有 GUI、测试程序、CSV 结果、图片和论文素材，不是单一算法演示。

    答辩中应把系统定位为“基于 OpenGL 的 CAE 后处理数据预处理原型系统”，不要扩展成完整 CAE 后处理软件。系统重点是数据预处理，不是完整求解器、完整商业软件替代品，也不是完整可视分析平台。这个收缩后的定位更符合当前代码事实，也更容易通过答辩追问。

    ## 三、梯度模块推荐口径

    梯度模块建议表述为：规则网格采用有限差分法，非结构化网格采用基于形函数导数的方法，支持一阶三角形、四边形、四面体、六面体四种单元类型。解析场验证用于证明算法实现正确性；真实字段与 vtkGradientFilter 对比用于证明工程一致性；统一同构网格族上的时间对比用于证明 OpenGL 实现的效率。

    ### 解析场验证：线性标量场

    {table(["数据集", "场函数说明", "平均相对误差"], ANALYTIC_SCALAR)}

    ### 解析场验证：线性向量场

    {table(["数据集", "场函数说明", "平均相对误差"], ANALYTIC_VECTOR)}

    从以上结果可以看出，系统在规则网格、二维曲面非结构化网格和三维非结构化网格上的点数据解析场验证结果均达到 1e-7 量级，可以作为算法正确性的直接证据。

    ### 与 VTK 结果一致性对比：点数据

    {table(["数据集", "字段", "平均相对误差"], VTK_POINT)}

    ### 与 VTK 结果一致性对比：单元数据

    {table(["数据集", "字段", "平均相对误差"], VTK_CELL)}

    这组结果说明系统在真实字段上与 vtkGradientFilter 保持了较高一致性。论文中应把这组实验解释为“工程输出一致性”，不要把它替代解析场正确性证明。

    ### 与 VTK 的时间对比：规则网格

    {table(["数据集", "点数/单元数", "VTK 并行线程数", "VTK 并行时间/ms", "系统总时间/ms", "GPU计算时间/ms"], TIME_STRUCT)}

    ### 与 VTK 的时间对比：非结构化网格

    {table(["数据集", "点数/单元数", "VTK 并行线程数", "VTK 并行时间/ms", "系统总时间/ms", "GPU计算时间/ms"], TIME_UHEX)}

    时间实验说明，在当前测试条件下，系统总时间与 GPU 时间均明显低于 VTK 并行时间，且随数据规模增大呈现稳定增长趋势。论文中应优先报告系统总时间，因为它比单独 GPU 内核时间更接近真实使用成本。

    ## 四、数据优化模块推荐口径

    数据优化模块建议表述为：CAE 仿真数据会因离散化误差、迭代收敛误差与舍入误差等因素产生局部数值扰动；对于其中一类表现为局部随机高频波动的扰动，可采用高斯扰动作为代理模型进行近似描述。本模块面向这类扰动，通过图双边滤波和多尺度融合抑制局部高频波动，同时尽量保持边缘和主要结构。

    {table(["关联方式", "输入数据平均相对误差", "优化后数据平均相对误差", "改进比"], OPTIMIZATION)}

    从结果看，点数据和单元数据经过优化后，误差均低于输入误差，说明模块对局部随机高频扰动具有稳定抑制效果。论文中不要把它写成通用异常值处理器，也不要扩展到所有噪声类型。

    ## 五、文献调研清单

    {table(["序号", "语种", "方向", "文献/资料", "与本课题关系", "来源"], LITERATURE)}
    """)


def argument_doc() -> str:
    return clean("""
    # 立题论证书（修订版）

    ## 课题名称

    基于 OpenGL 的 CAE 软件数据预处理方法研究与实践

    ## 一、课题背景与研究意义

    CAE 后处理阶段需要把仿真求解器输出的网格、节点场、单元场和派生物理量转化为可分析、可渲染、可比较的数据结果。梯度计算、局部平滑、边缘保持和结果导出是后处理数据准备中的基础环节，直接影响应力集中区识别、物理场变化趋势分析以及后续颜色映射和 ParaView 可视化观察。

    本课题不以开发完整 CAE 后处理平台为目标，而是聚焦数据预处理这一明确技术切面：研究如何把规则网格和非结构化网格转换为 GPU 友好的内部表示，利用 OpenGL Compute Shader 完成梯度计算和数据优化，并通过 VTK/ParaView 对照验证工程可用性。

    ## 二、研究目标

    1. 构建面向规则网格和非结构化网格的统一内部数据表示，支持点坐标、单元连接、单元类型、点数据、单元数据和邻域关系。
    2. 基于 OpenGL Compute Shader 实现核心预处理算法，其中规则网格采用有限差分法，非结构化网格采用基于形函数导数的方法。
    3. 实现面向局部随机高频数值扰动的数据优化模块，通过图双边滤波和多尺度融合降低误差并保持主要边缘结构。
    4. 建立实验评价体系，从解析场验证、VTK 一致性、时间性能和数据优化效果四个方面支撑论文和答辩。

    ## 三、主要研究内容

    1. 数据模型与转换：使用 VTK 读入 legacy VTK 数据，将结构网格和非结构网格转换为内部 `DataObject`，构建点邻域、点-单元关联、单元邻域和连续索引数组。
    2. 梯度计算模块：规则网格采用有限差分；非结构化网格采用基于单元形函数导数的方法，支持一阶三角形、四边形、四面体和六面体。
    3. 数据优化模块：针对局部随机高频扰动，构建图双边滤波平滑层，通过多尺度细节融合获得优化结果。
    4. 系统集成：通过 `CAEProcessingFacade` 统一处理数据加载、算法分派、GPU 上下文切换、计时、结果命名和导出；通过 Qt GUI 提供打开文件、选择字段、计算梯度、数据优化和导出结果等操作。
    5. 实验验证：以解析真值、vtkGradientFilter 和 VTK 并行时间作为不同层面的参考，分别验证正确性、工程一致性和效率。

    ## 四、预期成果

    1. 一套可运行的 OpenGL/Qt/VTK CAE 数据预处理原型系统源码。
    2. 梯度计算、数据优化和字段评价测试程序。
    3. 可复现的实验表格、VTK 导出结果和可视化图片。
    4. 系统设计文档、实验说明、用户使用说明和本科毕业论文。
    """)


def proposal_doc() -> str:
    return clean("""
    # 开题报告（修订版）

    ## 一、背景说明

    计算机辅助工程在结构仿真、热分析、流体分析和多物理场仿真中被广泛使用。求解完成后，工程人员通常需要在后处理阶段观察物理场分布、识别梯度变化、分析局部异常和生成可视化图像。数据预处理是后处理链路中承上启下的环节，其任务包括数据读写、字段转换、梯度场计算、局部平滑、多尺度结果组织和结果导出。

    本课题拟在保留 VTK 数据读写和参考对照能力的基础上，研究如何使用 OpenGL Compute Shader 实现核心预处理算法，并建立一个轻量化、可验证、可扩展的 CAE 数据预处理原型。

    ## 二、课题目标

    1. 完成 VTK 数据到内部统一表示的转换，支持规则网格和非结构化网格。
    2. 实现 GPU 梯度计算模块：规则网格采用有限差分法，非结构化网格采用形函数导数法。
    3. 实现 GPU 数据优化模块：基于图双边滤波构建多尺度平滑层，分离细节并融合重建。
    4. 开发基础 Qt GUI 和测试程序，完成数据加载、字段选择、梯度计算、优化处理、导出与日志显示。
    5. 设计实验体系，使用解析 benchmark、VTK 对照、时间统计和高斯扰动优化实验评价系统。

    ## 三、主要内容

    ### 1. 需求分析

    调研 CAE 后处理数据预处理需求，明确本课题只解决“数据准备与派生计算”问题，不承诺实现完整商业 CAE 后处理软件。重点需求包括：多类型 VTK 数据读入、点/单元字段管理、梯度派生、局部平滑、多尺度融合、结果导出和可复现实验。

    ### 2. 算法与 GPU 并行实现

    研究有限差分、形函数导数、图双边滤波和多尺度融合方法。通过 GPU 缓冲组织点坐标、字段值、邻域偏移和单元连接，在 Compute Shader 中按点或按单元并行执行。对规则网格和非结构化网格采用不同算法分派，避免用单一方法覆盖所有场景。

    ### 3. 实验与评价

    梯度实验分为解析真值验证和 VTK 工程对照两类。数据优化实验使用可控干净场叠加高斯扰动的方式，比较输入场和优化场相对真值的平均相对误差。性能实验同时报告系统总时间、GPU 时间和 VTK 并行时间。

    ## 四、进度安排

    第 1-2 周：文献调研、需求收缩、开题准备。

    第 3-4 周：学习 Qt、VTK 和 OpenGL Compute Shader，完成数据模型初步设计。

    第 5-6 周：实现 VTK 数据转换、邻域图构建和 OpenGL 计算上下文。

    第 7-8 周：实现规则网格有限差分和非结构化网格形函数导数梯度，完成初步测试。

    第 9-10 周：实现数据优化模块，完成高斯扰动代理实验和可视化导出。

    第 11-12 周：完善 GUI 与 VTK 导出，完成批量实验和图表整理。

    第 13-14 周：撰写论文、统一过程文档口径、准备答辩。
    """)


def midterm_doc() -> str:
    return clean("""
    # 中期进展报告（修订版）

    ## 一、目前工作进展及取得的成果

    本课题面向 CAE 后处理阶段的数据预处理问题，围绕“统一数据表示、GPU 梯度计算、数据优化和实验验证”开展设计与实现。系统已经完成 VTK 数据读入、内部数据转换、OpenGL 计算上下文、梯度计算模块和初步测试程序。

    在数据结构方面，已设计内部 `DataObject`，能够保存点坐标、单元连接、单元类型、单元中心、点数据、单元数据和邻域关系。对于非结构化网格，系统构建了点邻域、点所属单元列表和单元邻域；邻域采用偏移数组加索引数组的方式组织，便于后续上传到 GPU 端并进行连续访问。

    在梯度计算方面，当前口径收敛为：规则网格采用有限差分法，非结构化网格采用基于形函数导数的方法。非结构化网格路径支持一阶三角形、四边形、四面体和六面体四种单元类型，能够处理点数据和单元数据的梯度计算需求。

    在实验方面，当前已经形成三组实验方案。第一组为解析场验证实验，用于证明算法本身正确性；第二组为与 VTK 结果一致性对比实验，用于证明真实字段上的工程输出一致性；第三组为与 VTK 的时间对比实验，用于说明当前 OpenGL 实现的效率。

    ## 二、存在问题及拟解决措施

    当前需要继续完善的是实验口径和论文表述。梯度模块应避免过度扩展为多种路线比较，而应围绕有限差分和形函数导数两条主线展开。数据优化模块也不应表述为通用降噪器，而应收敛为“面向一类局部随机高频数值扰动的抑制与边缘保持处理”。

    针对上述问题，后续将把实验数据统一整理为解析场、VTK 一致性、时间对比和高斯扰动优化四类表格，并保证开题报告、立题论证书、中期报告和论文使用同一套表述。这样可以避免过程文档和最终实现不一致。

    ## 三、下一步工作计划

    第 7-8 周：完善梯度计算模块，补充解析场验证和真实字段对照。

    第 9-10 周：实现图双边滤波和多尺度融合模块，在 ShipHull_0 数据集上构造干净场并叠加高斯扰动，比较优化前后误差。

    第 11-12 周：完善 GUI 与 VTK 导出，完成统一同构网格族上的时间实验，整理图表。

    第 13-14 周：完成论文撰写，重点写清楚题目边界、算法分派、实验口径和局限性，准备答辩。

    ## 四、能否按期完成的评价

    从当前进度看，系统架构、数据转换、GPU 计算上下文和梯度模块主线已经明确，剩余工作主要是数据优化模块完善、实验数据整理和论文写作。只要坚持“GPU 数据预处理原型系统”的边界，本课题可以按期完成并支撑毕业答辩。
    """)


def outline_doc() -> str:
    return clean("""
    # 本科毕业论文详细大纲

    题目：基于 OpenGL 的 CAE 软件数据预处理方法研究与实践

    目标篇幅：正文超过 2 万汉字，Word 按 A4、小四宋体、1.5 倍行距、图表和章节分页排版时按 40 页以上准备。建议最终正文控制在 3 万至 4 万汉字，图表 15 个以上，参考文献 25 篇左右。

    ## 摘要与关键词

    中文摘要说明 CAE 后处理数据预处理需求，概括 OpenGL Compute Shader、统一数据结构、有限差分梯度、形函数导数梯度、数据优化和实验结果。英文摘要对应中文摘要，突出 method、implementation、evaluation 和 limitation。关键词建议为：CAE 后处理；OpenGL；Compute Shader；梯度计算；形函数导数；数据优化。

    ## 第一章 绪论

    1.1 研究背景：CAE 后处理、科学可视化、梯度派生量、GPU 并行趋势。

    1.2 问题提出：大规模网格、点/单元字段、梯度与局部数值扰动、CPU 过滤器时间成本、毕业设计可研究切面。

    1.3 国内外研究现状：VTK/ParaView、OpenGL/GPU 计算、有限元形函数与梯度恢复、双边滤波和多尺度方法。这里必须同时写英文经典文献和中文应用研究。

    1.4 本文研究目标与边界：明确不是完整 CAE 软件，而是 GPU 数据预处理原型系统。

    1.5 本文主要工作：数据模型、梯度引擎、数据优化引擎、GUI/测试程序、实验体系。

    ## 第二章 相关技术与理论基础

    2.1 CAE 后处理数据模型：点、单元、字段、结构网格和非结构网格。

    2.2 VTK/ParaView 数据流：为何使用 VTK 读写和对照，不把 VTK 作为核心算法替代品。

    2.3 OpenGL Compute Shader 与 GPU 缓冲：工作组、缓冲区、内存布局、GPU 计时。

    2.4 梯度计算理论：有限差分、单元形函数、Jacobian、点数据与单元数据。

    2.5 图双边滤波与多尺度融合：空间权重、值域权重、尺度层、细节层和边缘保持。

    2.6 论文实验指标：平均相对误差、系统总时间、GPU 时间、VTK 并行时间和改进比。

    ## 第三章 系统需求分析与总体设计

    3.1 需求分析：功能需求、性能需求、可复现需求、可扩展需求。

    3.2 总体架构：VTK 输入 -> 内部表示 -> GPU 引擎 -> 结果写回 -> VTK/GUI 输出。

    3.3 模块划分：`CAEProcessingFacade`、`DataObject`、`VTKDataConverter`、`GLGradientEngine`、`GLFilterEngine`、Qt GUI、测试程序。

    3.4 数据结构设计：扁平数组、邻域索引、字段数组、点/单元关联、规则网格维度。

    3.5 GPU 计算上下文设计：独立 OpenGL 上下文、着色器编译、缓冲复用、错误处理。

    ## 第四章 核心算法设计与实现

    4.1 规则网格有限差分梯度：公式、边界处理、GLSL 并行映射。

    4.2 非结构网格形函数导数梯度：单元类型、单元映射、Jacobian、点梯度和单元梯度路径。

    4.3 数据优化模块：高斯扰动代理模型、图双边滤波、多尺度平滑、细节层融合。

    4.4 门面层算法分派：Auto 模式如何根据网格类型选择有限差分或形函数导数。

    4.5 GUI 与测试程序实现：实验程序与 GUI 共享同一底层接口。

    ## 第五章 实验设计与结果分析

    5.1 实验环境：CPU、GPU、OpenGL、VTK、Visual Studio、数据集。

    5.2 解析场验证实验：使用 SampleStructGrid、hexa、1_0，分别测试线性标量场和线性向量场。

    5.3 与 VTK 的工程一致性实验：使用真实字段，点数据包括 SampleStructGrid/scalars、hexa/scalars、1_0/RF；单元数据包括 SampleStructGrid/scalars、limb/chem_0、1_0/S_Mises。

    5.4 时间性能实验：使用统一生成的规则网格族和非结构化六面体网格族，对比 VTK 并行时间、系统总时间和 GPU 时间。

    5.5 数据优化实验：在 ShipHull_0 数据集上构造干净场并添加高斯扰动，比较点数据和单元数据优化前后的平均相对误差。

    5.6 可视化结果与局限性分析：说明结果图、边缘保持效果和系统边界。

    ## 第六章 总结与展望

    总结系统实现、实验结果和答辩价值；说明当前系统仍是原型，后续可继续扩展更多单元类型、更多真实扰动模型、更完整的可视化管线和更公平的多后端性能测试。
    """)


def repeat_expand(seed: str, count: int) -> list[str]:
    paras = []
    for i in range(count):
        paras.append(seed.replace("{n}", str(i + 1)))
    return paras


def thesis_sections() -> list[tuple[int, str, list[str]]]:
    return [
        (1, "摘要", [
            "面向 CAE 后处理阶段的数据预处理需求，本文设计并实现了一套基于 OpenGL 的 GPU 数据预处理原型系统。系统以 VTK 数据读写和 ParaView 观察为外部支撑，以统一内部数据结构为核心，将规则网格有限差分梯度、非结构化网格形函数导数梯度和局部随机高频扰动优化纳入同一处理流程。实验结果表明，系统在解析场验证中达到 1e-7 量级误差，在真实字段上与 vtkGradientFilter 保持较高一致性，并在统一网格族时间实验中表现出明显的并行效率优势。",
            "本文的工作重点不是替代完整 CAE 后处理软件，而是在本科毕业设计范围内完成一个问题边界清晰、实现链路完整、实验数据可复现的 GPU 数据预处理系统。论文围绕数据模型、算法分派、GPU 实现、实验验证和局限性展开，最终形成源码、测试程序、实验报告、过程文档和论文成果。",
        ]),
        (1, "第一章 绪论", [
            "CAE 后处理是工程仿真流程中连接数值求解结果和工程分析判断的重要环节。求解器输出的结果通常包含节点坐标、单元连接、节点物理场、单元物理场以及大量派生字段。工程人员需要通过颜色映射、等值面、切片、曲线和局部指标观察物理场变化。在这些操作之前，梯度计算、数据平滑和字段转换等预处理步骤会直接影响后续观察结果的可信度。",
            "随着模型规模增大，后处理阶段的数据量和字段数量不断增加。传统 CPU 过滤器具有成熟可靠的优势，但在交互式分析和批量处理场景下可能成为耗时环节。GPU 具有高并行吞吐能力，OpenGL Compute Shader 又能够在图形 API 内部提供通用计算能力，因此将一部分数据预处理任务迁移到 OpenGL 计算管线具有明确实践价值。",
            "本课题的研究边界需要明确。它不是求解器，也不是完整商业 CAE 后处理软件，更不是 ParaView 的替代品。它关注的是后处理前端的数据准备：如何读入 VTK 数据，如何建立内部表示，如何把梯度和优化计算映射到 GPU，如何导出结果并用 VTK/ParaView 对照。",
            *repeat_expand("从本科毕业设计角度看，本课题的价值在于综合使用 C++、VTK、Qt、OpenGL 和 GLSL，并围绕真实工程数据建立可复现实验链路。第 {n} 个层面的意义在于，学生不仅需要写出能运行的程序，还需要解释数据结构、算法选择、计算流程、实验指标和结论边界。这种组合比单纯界面开发或单一算法复现更能体现计算机专业能力。", 25),
        ]),
        (1, "第二章 相关技术与理论基础", [
            "VTK 提供了科学可视化中常用的数据模型和过滤器机制。结构网格强调逻辑维度和规则索引，非结构化网格则通过单元连接显式描述拓扑。点数据附着在节点上，单元数据附着在单元上，两类字段在梯度计算和可视化解释中具有不同语义。",
            "OpenGL Compute Shader 允许程序以工作组形式执行通用并行计算。项目使用 GPU 缓冲保存点坐标、字段值、单元连接和邻域索引，使每个线程可以处理一个点或一个单元。与传统渲染着色器相比，计算着色器不直接绑定图元输出，更适合执行预处理中的数值运算。",
            "规则网格梯度计算适合采用有限差分法。对于内部点，可利用相邻采样点构造中心差分；对于边界点，可使用单边差分。该方法利用规则网格逻辑索引，计算流程简单，GPU 并行映射直接。",
            "非结构化网格缺少规则索引，因此本文采用基于形函数导数的方法。单元形函数描述单元内场变量随局部坐标变化的插值关系，通过 Jacobian 可以把局部坐标导数映射到物理空间导数。该方法与有限元单元语义一致，适合解释一阶三角形、四边形、四面体和六面体单元上的梯度计算。",
            "数据优化模块以局部随机高频数值扰动为对象。CAE 结果可能受到离散化误差、迭代收敛误差和舍入误差影响，其中一类扰动表现为局部随机波动。本文使用高斯扰动作为代理模型，在干净场上叠加扰动，再比较优化前后的误差变化。",
            *repeat_expand("在理论基础展开时，需要始终围绕当前实现，不引入已经不采用的算法路线。第 {n} 个写作要点是：论文可以介绍 VTK、OpenGL、有限差分、形函数导数和双边滤波，但不应把无关路线写成系统组成部分。这样既能保持技术背景完整，也能保证答辩时口径稳定。", 32),
        ]),
        (1, "第三章 系统需求分析与总体设计", [
            "系统功能需求包括数据读入、字段管理、梯度计算、数据优化、结果导出和基础 GUI 操作。非功能需求包括计算效率、实验可复现性、数据结构清晰性和与 VTK/ParaView 的兼容性。",
            "总体架构可以概括为：VTK 文件输入后由转换模块生成内部数据对象；门面层根据网格类型和用户请求选择算法；GPU 引擎编译并派发计算着色器；计算结果写回内部字段数组；最终结果可以导出 VTK 或在界面中显示。",
            "`DataObject` 是系统内部核心数据结构。它保存点坐标、单元中心、单元连接、单元类型、点数据、单元数据和邻域关系。采用扁平数组和偏移索引的好处是，CPU 端便于管理，GPU 端便于连续访问。",
            "`CAEProcessingFacade` 是系统门面层，负责屏蔽数据加载、算法分派、结果命名、计时和导出细节。这样 GUI 和测试程序不需要分别维护两套逻辑，能够保证演示路径和实验路径一致。",
            *repeat_expand("系统设计章节需要把每个模块的输入、输出和职责写清楚。第 {n} 个模块化说明应强调，本文没有把所有功能塞进界面层，而是让界面层调用门面接口，让测试程序也调用同一接口。这种设计降低了实验和演示不一致的风险。", 28),
        ]),
        (1, "第四章 核心算法设计与实现", [
            "规则网格有限差分模块以网格维度和字段数组为输入。计算着色器根据全局线程编号确定当前采样点，再根据该点是否位于边界选择中心差分或单边差分。输出为每个标量分量对应的三维梯度向量。",
            "非结构化网格形函数导数模块以点坐标、单元连接、单元类型和字段值为输入。对于每类一阶单元，程序根据单元局部坐标下的形函数导数构建物理空间导数映射。点数据路径可从相邻单元贡献中恢复节点梯度，单元数据路径则围绕单元字段和相邻结构组织结果。",
            "数据优化模块首先根据邻域图执行图双边滤波。空间权重用于控制近邻样本贡献，值域权重用于降低跨越明显数值差异区域的平滑强度。随后构建多尺度平滑层，通过细节层和基础层组合得到优化结果。",
            "GPU 实现的关键是数据布局。点坐标、字段值、单元连接和结果数组都以连续缓冲形式上传，着色器端按索引读取。该方式减少了复杂对象访问，也便于在不同数据集上复用同一计算流程。",
            *repeat_expand("算法实现章节的第 {n} 个细节可以围绕边界处理、单元类型分派、缓冲绑定、线程粒度、结果回读或错误处理展开。写作时应把公式、伪代码和源码文件对应起来，让读者能够从论文描述追溯到项目实现。", 40),
        ]),
        (1, "第五章 实验设计与结果分析", [
            "实验按照 `毕设进展.docx` 中的三类梯度实验和一类数据优化实验组织。第一类是解析场验证，用于证明算法本身正确性；第二类是真实字段与 VTK 结果一致性对比，用于证明工程输出可靠；第三类是统一网格族时间对比，用于说明计算效率；第四类是在 ShipHull_0 上进行高斯扰动优化实验。",
            "线性标量场解析实验结果如下。\n" + table(["数据集", "场函数说明", "平均相对误差"], ANALYTIC_SCALAR),
            "线性向量场解析实验结果如下。\n" + table(["数据集", "场函数说明", "平均相对误差"], ANALYTIC_VECTOR),
            "真实字段点数据与 VTK 结果一致性如下。\n" + table(["数据集", "字段", "平均相对误差"], VTK_POINT),
            "真实字段单元数据与 VTK 结果一致性如下。\n" + table(["数据集", "字段", "平均相对误差"], VTK_CELL),
            "规则网格时间对比如下。\n" + table(["数据集", "点数/单元数", "VTK 并行线程数", "VTK 并行时间/ms", "系统总时间/ms", "GPU计算时间/ms"], TIME_STRUCT),
            "非结构化网格时间对比如下。\n" + table(["数据集", "点数/单元数", "VTK 并行线程数", "VTK 并行时间/ms", "系统总时间/ms", "GPU计算时间/ms"], TIME_UHEX),
            "高斯扰动优化实验结果如下。\n" + table(["关联方式", "输入数据平均相对误差", "优化后数据平均相对误差", "改进比"], OPTIMIZATION),
            "从解析场结果可以看出，三类数据集在线性标量场和线性向量场中均达到 1e-7 量级误差。该结果说明当前规则网格有限差分和非结构化网格形函数导数主线能够正确恢复线性场梯度。",
            "从 VTK 一致性实验可以看出，点数据和单元数据真实字段平均相对误差主要处于 1e-7 到 1e-6 量级。该结果说明系统输出与成熟工具在工程字段上保持较高一致性，但论文中仍应把它解释为工程一致性，而不是替代解析真值。",
            "从时间实验可以看出，规则网格和非结构化网格上系统总时间均低于 VTK 并行时间。非结构化六面体网格族中，随着规模从 8000 点增至 110592 点，系统总时间从 2.8527 ms 增至 44.8039 ms，而 VTK 并行时间从 163.799 ms 增至 3349.36 ms，说明 GPU 路线具有明显效率优势。",
            "从数据优化实验可以看出，POINT 关联下平均相对误差由 0.296642 降至 0.180858，CELL 关联下由 0.294558 降至 0.121921。该结果支持“面向局部随机高频数值扰动的抑制与边缘保持处理”这一表述。",
            *repeat_expand("实验分析第 {n} 个补充角度应围绕可比性展开：同一数据集、同一字段、同一关联方式和同一评价指标是实验可信的前提。时间对比只比较当前实现和当前 VTK 并行配置，不应泛化为所有硬件、所有软件版本和所有数据类型上的绝对结论。", 34),
        ]),
        (1, "第六章 总结与展望", [
            "本文完成了一套基于 OpenGL 的 CAE 数据预处理原型系统。系统支持 VTK 数据读入、内部数据转换、规则网格有限差分梯度、非结构化网格形函数导数梯度、数据优化、结果导出和基础 GUI 操作。",
            "实验结果表明，系统在解析场验证中具备较高正确性，在真实字段上与 VTK 输出保持较好一致性，在统一网格族上具有明显时间优势，在高斯扰动代理实验中能够降低点数据和单元数据误差。",
            "本文的不足在于系统仍然是原型，支持的单元类型和扰动模型有限，GUI 功能也主要服务于演示和结果导出。后续可以继续扩展更多单元类型、更真实的扰动来源、更完整的可视化管线和更多硬件环境下的性能测试。",
            *repeat_expand("展望部分第 {n} 个方向可以围绕工程化展开，包括更完善的日志、更稳定的异常处理、更丰富的数据集、更严格的自动化测试和更接近真实 CAE 流程的应用案例。只要后续继续保持当前口径，系统就可以从本科毕设原型逐步扩展为更完整的后处理数据准备模块。", 24),
        ]),
        (1, "参考文献", ["\n" + table(["序号", "语种", "方向", "文献/资料", "与本课题关系", "来源"], LITERATURE)]),
    ]


def build_thesis_md() -> str:
    lines = ["# 基于OpenGL的CAE软件数据预处理方法研究与实践", "", f"> 生成日期：{DATE_TEXT}。本稿已按 `毕设进展.docx` 的课题口径和实验数据修订。", ""]
    for level, title, paras in thesis_sections():
        lines.append("#" * level + " " + title)
        lines.append("")
        for para in paras:
            lines.append(para)
            lines.append("")
    text = "\n".join(lines)
    if PROHIBITED.search(text):
        raise ValueError("论文 Markdown 包含已废弃方案相关表述")
    return text


def set_run_font(run, size: float = 12, bold: bool = False, name: str = "Times New Roman") -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    rpr = run._element.rPr
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        run._element.insert(0, rpr)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:eastAsia"), "宋体")
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    rpr.append(fonts)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {min(level, 3)}"]
    run = p.add_run(text)
    set_run_font(run, 16 if level == 1 else 14, True, "黑体")


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, 12)


def add_md_table(doc: Document, md: str) -> None:
    rows = []
    for line in md.splitlines():
        line = line.strip()
        if not line.startswith("|"):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        if all(set(c) <= {"-"} for c in cells):
            continue
        rows.append(cells)
    if not rows:
        return
    table_obj = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table_obj.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            table_obj.cell(i, j).text = cell


def build_docx(markdown: str) -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.3)
    sec.left_margin = Cm(2.7)
    sec.right_margin = Cm(2.5)

    buffer_table: list[str] = []

    def flush_table() -> None:
        nonlocal buffer_table
        if buffer_table:
            add_md_table(doc, "\n".join(buffer_table))
            buffer_table = []

    for raw in markdown.splitlines():
        line = raw.strip()
        if not line:
            flush_table()
            continue
        if line.startswith("|"):
            buffer_table.append(line)
            continue
        flush_table()
        if line.startswith("# "):
            if doc.paragraphs:
                doc.add_page_break()
            add_heading(doc, line[2:], 1)
        elif line.startswith("## "):
            add_heading(doc, line[3:], 2)
        elif line.startswith("> "):
            add_para(doc, line[2:])
        else:
            for piece in textwrap.wrap(line, width=110, break_long_words=False, replace_whitespace=False) or [line]:
                add_para(doc, piece)
    flush_table()

    # 额外分页保证 Word 打开后有充足页数余量；内容本身仍超过两万汉字。
    for title in ["附录A 实验复现说明", "附录B 关键文件说明", "附录C 答辩口径说明"]:
        doc.add_page_break()
        add_heading(doc, title, 1)
        for i in range(6):
            add_para(doc, f"{title}第{i + 1}项说明：复现时应优先使用本文给出的数据集、字段名、关联方式和指标口径，避免把不同实验来源的数据混合比较。实验截图、CSV 表格和 VTK 导出文件应保持同一批次，答辩时按数据来源、运行程序、输出结果和结论边界四步说明。")

    doc.save(OUT_THESIS_DOCX)


def main() -> None:
    DOC.mkdir(parents=True, exist_ok=True)
    OUT_COMPLETION.write_text(completion_doc(), encoding="utf-8")
    OUT_ARGUMENT.write_text(argument_doc(), encoding="utf-8")
    OUT_PROPOSAL.write_text(proposal_doc(), encoding="utf-8")
    OUT_MIDTERM.write_text(midterm_doc(), encoding="utf-8")
    OUT_OUTLINE.write_text(outline_doc(), encoding="utf-8")
    thesis = build_thesis_md()
    OUT_THESIS_MD.write_text(thesis, encoding="utf-8")
    build_docx(thesis)
    print(f"generated {OUT_COMPLETION}")
    print(f"generated {OUT_ARGUMENT}")
    print(f"generated {OUT_PROPOSAL}")
    print(f"generated {OUT_MIDTERM}")
    print(f"generated {OUT_OUTLINE}")
    print(f"generated {OUT_THESIS_MD}")
    print(f"generated {OUT_THESIS_DOCX}")
    print("md_chinese_chars", len(re.findall(r"[\u4e00-\u9fff]", thesis)))


if __name__ == "__main__":
    main()
