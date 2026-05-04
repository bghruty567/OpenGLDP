from __future__ import annotations

import argparse
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def set_run_font(run, size_pt: float = 12, bold: bool = False, name: str = "Times New Roman") -> None:
    run.font.name = name
    run.bold = bold
    run.font.size = Pt(size_pt)
    rpr = run._element.rPr
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        run._element.insert(0, rpr)
    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), "宋体")
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rpr.append(rfonts)


def add_page_number(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    p._p.append(fld_begin)
    p._p.append(instr)
    p._p.append(fld_end)


def clean_inline(text: str) -> str:
    return text.strip().replace("`", "")


def split_table_row(line: str) -> list[str]:
    content = line.strip().strip("|")
    return [clean_inline(cell) for cell in content.split("|")]


def is_table_separator(line: str) -> bool:
    stripped = line.strip()
    if not stripped.startswith("|") or not stripped.endswith("|"):
        return False
    cells = [cell.strip() for cell in stripped.strip("|").split("|")]
    return all(cell and set(cell) <= {"-", ":"} for cell in cells)


def is_numbered_item(line: str) -> bool:
    return re.match(r"^\d+\.\s+", line.strip()) is not None


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {min(level, 3)}"]
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(clean_inline(text))
    if level == 1:
        set_run_font(run, 16, True, "黑体")
    elif level == 2:
        set_run_font(run, 14, True, "黑体")
    else:
        set_run_font(run, 12, True, "黑体")


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(clean_inline(text))
    set_run_font(run, 12)


def add_list_item(doc: Document, text: str, numbered: bool) -> None:
    p = doc.add_paragraph(style="List Number" if numbered else "List Bullet")
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(clean_inline(text))
    set_run_font(run, 12)


def add_code_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.2
        run = p.add_run(line.rstrip("\n"))
        set_run_font(run, 10.5, False, "Consolas")


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_cells = table.rows[0].cells
    for i, head in enumerate(rows[0]):
        p = header_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(head)
        set_run_font(run, 10.5, True)
        header_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for row in rows[1:]:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(value)
            set_run_font(run, 10.5)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def build_doc(markdown: str) -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.start_type = WD_SECTION_START.NEW_PAGE
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.6)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.2)
    add_page_number(sec)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    styles["Heading 1"].font.name = "黑体"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.name = "黑体"
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 3"].font.name = "黑体"
    styles["Heading 3"].font.size = Pt(12)

    lines = markdown.splitlines()
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("```"):
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, code_lines)
            i += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            add_heading(doc, stripped[level:].strip(), level)
            i += 1
            continue

        if stripped.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            table_rows = [split_table_row(lines[i])]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_rows.append(split_table_row(lines[i]))
                i += 1
            add_table(doc, table_rows)
            continue

        if stripped.startswith("- "):
            add_list_item(doc, stripped[2:].strip(), numbered=False)
            i += 1
            continue

        if is_numbered_item(stripped):
            add_list_item(doc, re.sub(r"^\d+\.\s+", "", stripped), numbered=True)
            i += 1
            continue

        para_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if (
                not nxt
                or nxt.startswith("#")
                or nxt.startswith("```")
                or nxt.startswith("|")
                or nxt.startswith("- ")
                or is_numbered_item(nxt)
            ):
                break
            para_lines.append(nxt)
            i += 1
        add_paragraph(doc, " ".join(para_lines))

    return doc


def save_doc_with_fallback(doc: Document, path: Path) -> Path:
    try:
        doc.save(path)
        return path
    except PermissionError:
        fallback = path.with_name(f"{path.stem}_更新版{path.suffix}")
        try:
            doc.save(fallback)
            return fallback
        except PermissionError:
            stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            fallback2 = path.with_name(f"{path.stem}_{stamp}{path.suffix}")
            doc.save(fallback2)
            return fallback2


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Export markdown report to Word docx.")
    parser.add_argument("--src", required=True, help="Source markdown path.")
    parser.add_argument("--out-docx", required=True, help="Output docx path.")
    parser.add_argument("--out-md", help="Optional copied markdown output path.")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    src = Path(args.src).resolve()
    out_docx = Path(args.out_docx).resolve()
    out_docx.parent.mkdir(parents=True, exist_ok=True)

    markdown = src.read_text(encoding="utf-8")

    if args.out_md:
        out_md = Path(args.out_md).resolve()
        out_md.parent.mkdir(parents=True, exist_ok=True)
        out_md.write_text(markdown, encoding="utf-8")

    doc = build_doc(markdown)
    saved = save_doc_with_fallback(doc, out_docx)
    print(f"source: {src}")
    if args.out_md:
        print(f"markdown: {Path(args.out_md).resolve()}")
    print(f"docx: {saved}")


if __name__ == "__main__":
    main()
